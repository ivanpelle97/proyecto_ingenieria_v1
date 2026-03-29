from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = Path(r"C:\Users\hi_iv\Downloads\PROYECTO_FINAL_CORREGIDO_MANUAL.docx")
OUTPUT_DOC = ROOT / "PROYECTO_FINAL_CORREGIDO_MANUAL_v2_prueba2_prueba3.docx"
RUN_3A = ROOT / "runs" / "demo_s2" / "train_02"
RUN_4A = ROOT / "runs" / "demo_s4" / "train_04"
SCREENSHOT_DIR = Path(r"C:\Users\hi_iv\OneDrive\Imágenes\Capturas de pantalla")
GENERATED_DIR = ROOT / "doc_assets" / "generated"


def ensure_dirs() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = bold


def style_last_paragraph(doc: Document, italic: bool = False, align_center: bool = False) -> None:
    paragraph = doc.paragraphs[-1]
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = italic


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.style = doc.styles["Normal"]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(8)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True


def add_table_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr_cells[i], header, bold=True)

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], "" if value is None else value)

    doc.add_paragraph("")


def create_placeholder_image(title: str, subtitle: str, output_path: Path) -> Path:
    width, height = 1600, 900
    image = Image.new("RGB", (width, height), (241, 243, 245))
    draw = ImageDraw.Draw(image)
    font_title = ImageFont.load_default()
    font_body = ImageFont.load_default()
    draw.rectangle((40, 40, width - 40, height - 40), outline=(120, 120, 120), width=5)
    draw.multiline_text((120, 220), title, fill=(30, 30, 30), font=font_title, spacing=12)
    draw.multiline_text((120, 420), subtitle, fill=(70, 70, 70), font=font_body, spacing=10)
    image.save(output_path)
    return output_path


def create_contact_sheet(image_paths: list[Path], output_path: Path) -> Path:
    thumb_w, thumb_h = 420, 240
    cols = 2
    rows = (len(image_paths) + cols - 1) // cols
    canvas = Image.new("RGB", (cols * thumb_w + 80, rows * (thumb_h + 50) + 80), (255, 255, 255))
    draw = ImageDraw.Draw(canvas)
    font = ImageFont.load_default()

    for index, path in enumerate(image_paths):
        row = index // cols
        col = index % cols
        x = 40 + col * thumb_w
        y = 40 + row * (thumb_h + 50)
        with Image.open(path) as image:
            image = image.convert("RGB")
            image.thumbnail((thumb_w - 20, thumb_h - 20))
            tile = Image.new("RGB", (thumb_w - 20, thumb_h - 20), (248, 249, 250))
            offset_x = (tile.width - image.width) // 2
            offset_y = (tile.height - image.height) // 2
            tile.paste(image, (offset_x, offset_y))
            canvas.paste(tile, (x + 10, y + 10))
            draw.rectangle((x, y, x + thumb_w - 1, y + thumb_h - 1), outline=(180, 180, 180), width=2)
        draw.text((x, y + thumb_h + 8), path.stem.replace("Captura de pantalla ", ""), fill=(20, 20, 20), font=font)

    canvas.save(output_path)
    return output_path


def load_run_summary(run_dir: Path) -> dict:
    experiment = json.loads((run_dir / "experiment.json").read_text(encoding="utf-8"))
    points = pd.read_csv(run_dir / "points.csv")
    samples = pd.read_csv(run_dir / "samples.csv")
    dataset = pd.read_csv(run_dir / "dataset.csv") if (run_dir / "dataset.csv").exists() else None
    metrics = pd.read_csv(run_dir / "models" / "metrics.csv") if (run_dir / "models" / "metrics.csv").exists() else None

    points["started_at"] = pd.to_datetime(points["started_at"])
    points["capture_complete_at"] = pd.to_datetime(points["capture_complete_at"])
    points["completed_at"] = pd.to_datetime(points["completed_at"])
    points["capture_duration_s"] = (points["capture_complete_at"] - points["started_at"]).dt.total_seconds()
    points["close_duration_s"] = (points["completed_at"] - points["capture_complete_at"]).dt.total_seconds()
    points["total_duration_s"] = (points["completed_at"] - points["started_at"]).dt.total_seconds()

    raw_sizes = {path.stem: path.stat().st_size for path in sorted((run_dir / "raw").glob("*.jsonl"))}
    rssi_stats = (
        samples.groupby("anchor_id")["rssi"]
        .agg(["count", "min", "max", "mean", "std"])
        .reset_index()
        .rename(columns={"count": "n"})
    )

    summary = {
        "run_dir": run_dir,
        "experiment": experiment,
        "points": points,
        "samples": samples,
        "dataset": dataset,
        "metrics": metrics,
        "raw_sizes": raw_sizes,
        "rssi_stats": rssi_stats,
    }

    if metrics is not None and not metrics.empty:
        summary["best_mae"] = metrics.sort_values("mae_eucl").iloc[0].to_dict()
        summary["best_p95"] = metrics.sort_values("p95").iloc[0].to_dict()
    else:
        summary["best_mae"] = None
        summary["best_p95"] = None

    return summary


def find_screenshots() -> list[Path]:
    return sorted(
        [
            path
            for path in SCREENSHOT_DIR.glob("*.png")
            if "2026-03-23" in path.name
        ]
    )


def build_screenshot_caption_map() -> dict[str, str]:
    return {
        "Captura de pantalla 2026-03-23 163502.png": "Figura 31.2. Configuración inicial de la campaña demo_s4 / train_04 con cuatro antenas ESP32 y definición de la geometría del recinto.",
        "Captura de pantalla 2026-03-23 163509.png": "Figura 31.3. Verificación de parámetros de campaña y consistencia entre identificadores de sesión, campaña y MAC del emisor objetivo.",
        "Captura de pantalla 2026-03-23 163519.png": "Figura 31.4. Detalle de la tabla de antenas ESP32 y de las variables geométricas cargadas antes del inicio de la medición.",
        "Captura de pantalla 2026-03-23 163526.png": "Figura 31.5. Vista operativa del plano durante la captura, con el punto P03 como punto activo de entrenamiento.",
        "Captura de pantalla 2026-03-23 180641.png": "Figura 31.6. Plano con los dieciocho puntos capturados al finalizar la fase de medición de la segunda prueba.",
        "Captura de pantalla 2026-03-23 181743.png": "Figura 31.7. Primera vista de inferencia en línea con múltiples dispositivos detectados sobre el plano del ambiente.",
        "Captura de pantalla 2026-03-23 181839.png": "Figura 31.8. Evolución de la inferencia en línea con concentración de predicciones en la región central del ambiente.",
        "Captura de pantalla 2026-03-23 182108.png": "Figura 31.9. Inferencia en línea con seguimiento del emisor objetivo sobre la nube de dispositivos detectados.",
        "Captura de pantalla 2026-03-23 182301.png": "Figura 31.10. Actualización intermedia del mapa de calor durante la segunda prueba con cuatro antenas ESP32.",
        "Captura de pantalla 2026-03-23 182601.png": "Figura 31.11. Persistencia espacial del mapa de calor durante la inferencia en línea posterior al entrenamiento.",
        "Captura de pantalla 2026-03-23 183605.png": "Figura 31.12. Inferencia en línea en fase avanzada, con superposición de posiciones estimadas para dispositivos detectados.",
        "Captura de pantalla 2026-03-23 183853.png": "Figura 31.13. Comparación visual entre puntos entrenados y posiciones inferidas durante la operación en línea.",
        "Captura de pantalla 2026-03-23 183920.png": "Figura 31.14. Estado de la inferencia con menor cantidad de dispositivos visibles y concentración espacial más acotada.",
        "Captura de pantalla 2026-03-23 183930.png": "Figura 31.15. Vista combinada del mapa de calor y de la tabla completa de puntos capturados en la campaña.",
        "Captura de pantalla 2026-03-23 184258.png": "Figura 31.16. Estado final de la inferencia en línea registrado al cierre de la segunda prueba con cuatro antenas ESP32.",
    }


def append_existing_references(doc: Document, reference_texts: list[str]) -> None:
    add_heading(doc, "33. Referencias técnicas complementarias", level=1)
    for text in reference_texts:
        if not text:
            continue
        paragraph = doc.add_paragraph(text)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, width: float = 6.5) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def build_trial2_summary_rows(run4: dict) -> list[list[object]]:
    exp = run4["experiment"]
    points = run4["points"]
    total_samples = len(run4["samples"])
    best_mae = run4["best_mae"]
    best_p95 = run4["best_p95"]
    return [
        ["Sesión", exp["session_id"]],
        ["Campaña", exp["campaign_id"]],
        ["MAC del emisor objetivo", exp["target_mac"]],
        ["Cantidad de antenas ESP32", len(exp["anchors"])],
        ["Puntos capturados", len(points)],
        ["Muestras aceptadas totales", total_samples],
        ["Muestras por antena ESP32 y por punto", exp["samples_per_anchor"]],
        ["Dimensiones del dataset", f"{run4['dataset'].shape[0]} filas x {run4['dataset'].shape[1]} columnas"],
        ["Mejor modelo por MAE", f"{best_mae['model']} (MAE = {best_mae['mae_eucl']:.4f} m)"],
        ["Mejor modelo por percentil 95", f"{best_p95['model']} (P95 = {best_p95['p95']:.4f} m)"],
        ["Tiempo medio de captura por punto", f"{points['capture_duration_s'].mean():.2f} s"],
        ["Tiempo medio total por punto", f"{points['total_duration_s'].mean():.2f} s"],
    ]


def append_trial2_results(doc: Document, run3: dict, run4: dict, screenshot_paths: list[Path]) -> None:
    exp = run4["experiment"]
    points = run4["points"]
    metrics = run4["metrics"].copy()
    metrics = metrics.sort_values("mae_eucl")

    add_heading(doc, "28. Actualización experimental de la segunda prueba con cuatro antenas ESP32", level=1)
    add_body_paragraph(
        doc,
        "La estructura original del informe dejaba la Prueba 2 con cuatro antenas ESP32 como etapa planificada. "
        "A la fecha de esta actualización ya se ejecutó una campaña real, identificada como demo_s4 / train_04, "
        "con cuatro antenas ESP32 instaladas en las cuatro esquinas del recinto y con dieciocho puntos de referencia "
        "relevados en la misma grilla geométrica general utilizada en la etapa previa. Esta incorporación convierte a la "
        "Prueba 2 en evidencia experimental efectiva y ya no solamente en una previsión metodológica."
    )
    add_body_paragraph(
        doc,
        "La campaña se diseñó para contrastar de manera directa la hipótesis de mejora geométrica. Se esperaba que una "
        "topología de observación cerrada, con un receptor en cada esquina, mejorara la observabilidad espacial y redujera "
        "la asimetría detectada durante la primera etapa de tres antenas ESP32. El resultado más relevante de esta segunda "
        "prueba es que la completitud de captura se obtuvo sin inconvenientes: en los dieciocho puntos se alcanzó el cupo "
        "de diez muestras por cada una de las cuatro antenas ESP32. Sin embargo, esa mejora operativa no se tradujo en una "
        "mejora automática de las métricas de aprendizaje automático."
    )
    add_body_paragraph(
        doc,
        "La actualización que sigue documenta las condiciones reales de la campaña, resume los resultados cuantitativos "
        "persistidos en los archivos de ejecución y deja trazabilidad explícita de las capturas gráficas obtenidas durante "
        "la operación. Esta documentación es importante porque muestra que la segunda prueba no fue descartada por falta de "
        "datos, sino por una combinación de resultados medidos y observaciones de comportamiento en línea."
    )

    add_table_title(doc, "Tabla 28.1 - Geometría real utilizada en la segunda prueba con cuatro antenas ESP32.")
    geometry_rows = [
        [anchor["anchor_id"], f"{anchor['x_m']:.2f}", f"{anchor['y_m']:.2f}", f"{anchor['z_m']:.2f}", anchor.get("model", "ESP32")]
        for anchor in exp["anchors"]
    ]
    add_table(doc, ["Antena ESP32", "x [m]", "y [m]", "z [m]", "Modelo"], geometry_rows)

    add_table_title(doc, "Tabla 28.2 - Resumen global de la campaña demo_s4 / train_04.")
    add_table(doc, ["Variable", "Valor medido"], build_trial2_summary_rows(run4))

    add_table_title(doc, "Tabla 28.3 - Resultados por punto de referencia en la segunda prueba con cuatro antenas ESP32.")
    point_rows = []
    for row in points.itertuples(index=False):
        point_rows.append(
            [
                row.point_id,
                f"{row.x_m:.1f}",
                f"{row.y_m:.1f}",
                f"{row.z_m:.1f}",
                row.started_at.strftime("%H:%M:%S"),
                f"{row.capture_duration_s:.2f}",
                f"{row.close_duration_s:.2f}",
                f"{row.total_duration_s:.2f}",
                int(row.A1_count),
                int(row.A2_count),
                int(row.A3_count),
                int(row.A4_count),
            ]
        )
    add_table(
        doc,
        ["Punto", "x [m]", "y [m]", "z [m]", "Inicio", "Captura [s]", "Cierre [s]", "Total [s]", "A1", "A2", "A3", "A4"],
        point_rows,
    )

    add_table_title(doc, "Tabla 28.4 - Volumen de muestras persistidas y tamaños de archivos RAW por antena ESP32.")
    sample_counts = run4["samples"]["anchor_id"].value_counts().to_dict()
    volume_rows = []
    for anchor_id in sorted(sample_counts):
        raw_size = run4["raw_sizes"].get(anchor_id, 0)
        volume_rows.append([anchor_id, sample_counts[anchor_id], raw_size, f"{raw_size / 1024:.1f} KiB"])
    add_table(doc, ["Antena ESP32", "Muestras aceptadas", "Tamaño RAW [bytes]", "Tamaño RAW [KiB]"], volume_rows)

    add_table_title(doc, "Tabla 28.5 - Estadísticos de RSSI aceptado por antena ESP32 en la segunda prueba.")
    rssi_rows = []
    for row in run4["rssi_stats"].itertuples(index=False):
        rssi_rows.append([row.anchor_id, int(row.n), int(row.min), int(row.max), f"{row.mean:.3f}", f"{row.std:.3f}"])
    add_table(doc, ["Antena ESP32", "n", "RSSI mín [dBm]", "RSSI máx [dBm]", "Media [dBm]", "Desvío [dBm]"], rssi_rows)

    add_table_title(doc, "Tabla 28.6 - Métricas comparativas de modelos entrenados con la campaña demo_s4 / train_04.")
    metric_rows = []
    for row in metrics.itertuples(index=False):
        metric_rows.append(
            [
                row.model,
                f"{row.mae_eucl:.4f}",
                f"{row.p50:.4f}",
                f"{row.p90:.4f}",
                f"{row.p95:.4f}",
                f"{row.rmse_eucl:.4f}",
                row.eval_mode,
            ]
        )
    add_table(doc, ["Modelo", "MAE [m]", "P50 [m]", "P90 [m]", "P95 [m]", "RMSE [m]", "Modo de evaluación"], metric_rows)

    add_body_paragraph(
        doc,
        "El mejor modelo de la segunda prueba según error absoluto medio fue extratrees_500, con un MAE de "
        f"{run4['best_mae']['mae_eucl']:.4f} m. En términos de percentil 95, el mejor comportamiento correspondió a "
        f"{run4['best_p95']['model']}, con un valor de {run4['best_p95']['p95']:.4f} m. Estos valores deben leerse en conjunto "
        "con la evidencia operativa: la campaña fue completa y consistente en cantidad de muestras, pero el rendimiento "
        "predictivo no superó a la mejor campaña con tres antenas ESP32."
    )
    add_body_paragraph(
        doc,
        "En consecuencia, la segunda prueba con cuatro antenas ESP32 no puede interpretarse como un fracaso de captura; "
        "por el contrario, constituye una medición plenamente válida. La conclusión más importante es metodológica: "
        "aumentar la cantidad de receptores no garantiza por sí mismo una reducción del error si la estrategia temporal de "
        "inferencia y la dinámica de reporte entre antenas ESP32 introducen demoras, asincronías o dependencia excesiva de "
        "ventanas temporales amplias."
    )

    add_heading(doc, "29. Análisis comparativo entre la campaña principal con tres antenas ESP32 y la segunda prueba con cuatro antenas ESP32", level=1)
    add_body_paragraph(
        doc,
        "La comparación entre demo_s2 / train_02 y demo_s4 / train_04 permite separar dos dimensiones distintas del problema. "
        "La primera es la completitud de adquisición. La segunda es la calidad final de la inferencia. En la segunda prueba se "
        "consiguió la primera sin dificultades sustantivas: dieciocho puntos completos y setecientas veinte muestras aceptadas, "
        "distribuidas de forma perfectamente balanceada entre las cuatro antenas ESP32. Sin embargo, la segunda dimensión no mejoró."
    )
    add_body_paragraph(
        doc,
        "La campaña principal con tres antenas ESP32 había alcanzado un mejor MAE y un mejor percentil 95 que la segunda campaña "
        "con cuatro receptores. Esta situación sugiere que la simple expansión geométrica no fue suficiente para compensar el efecto "
        "de ventanas de detección amplias, de la espera práctica hasta obtener información de las cuatro antenas ESP32 y de la "
        "heterogeneidad temporal propia del tráfico observado. En otras palabras, la segunda prueba amplió cobertura, pero introdujo "
        "una demanda temporal mayor sobre la consolidación de la observación."
    )
    comparison_rows = [
        [
            "Campaña principal 3 antenas ESP32",
            "demo_s2 / train_02",
            len(run3["experiment"]["anchors"]),
            len(run3["points"]),
            len(run3["samples"]),
            f"{run3['dataset'].shape[0]} x {run3['dataset'].shape[1]}",
            run3["best_mae"]["model"],
            f"{run3['best_mae']['mae_eucl']:.4f}",
            run3["best_p95"]["model"],
            f"{run3['best_p95']['p95']:.4f}",
        ],
        [
            "Segunda prueba 4 antenas ESP32",
            "demo_s4 / train_04",
            len(run4["experiment"]["anchors"]),
            len(run4["points"]),
            len(run4["samples"]),
            f"{run4['dataset'].shape[0]} x {run4['dataset'].shape[1]}",
            run4["best_mae"]["model"],
            f"{run4['best_mae']['mae_eucl']:.4f}",
            run4["best_p95"]["model"],
            f"{run4['best_p95']['p95']:.4f}",
        ],
    ]
    add_table_title(doc, "Tabla 29.1 - Comparación cuantitativa entre la campaña principal con tres antenas ESP32 y la segunda prueba con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Etapa", "Run", "Antenas", "Puntos", "Muestras", "Dataset", "Mejor MAE", "MAE [m]", "Mejor P95", "P95 [m]"],
        comparison_rows,
    )
    add_body_paragraph(
        doc,
        "La interpretación más prudente de esta evidencia es que la etapa con cuatro antenas ESP32 dejó planteado un nuevo cuello "
        "de botella: la necesidad de una inferencia en línea menos dependiente de esperar simultaneidad práctica entre receptores. "
        "Esa observación, reforzada por las capturas de pantalla obtenidas durante la campaña, motiva directamente la tercera prueba."
    )

    add_heading(doc, "30. Formulación de la Prueba 3: inferencia en línea continua con canal fijo y ventana temporal reducida", level=1)
    add_body_paragraph(
        doc,
        "A partir de la segunda prueba se definió una tercera etapa con un objetivo claramente distinto. Ya no se trata solo de "
        "incrementar la geometría de observación, sino de modificar el régimen operativo de inferencia en línea. La hipótesis de "
        "trabajo es que las ventanas de detección empleadas hasta aquí resultaron demasiado amplias y que la necesidad práctica de "
        "esperar muestras de las cuatro antenas ESP32 redujo la inmediatez de la estimación y afectó la percepción de robustez."
    )
    add_body_paragraph(
        doc,
        "La tercera prueba se apoyará en dos cambios de firmware específicos. En primer lugar, se utilizará una versión continua de "
        "firmware para las antenas ESP32, capaz de permanecer conectada al punto de acceso y de enviar mini lotes periódicos sin "
        "alternar largos ciclos de captura y reconexión. En segundo lugar, se empleará un emisor objetivo dedicado a inferencia en "
        "línea sobre canal fijo, de modo tal que la observación de las cuatro antenas ESP32 se produzca bajo una condición radioeléctrica "
        "más estable y compatible con tráfico continuo."
    )
    add_body_paragraph(
        doc,
        "La inferencia en línea de esta tercera prueba debe operar sobre una campaña ya entrenada y activada en modo de solo inferencia. "
        "El foco no estará puesto en construir un nuevo dataset supervisado, sino en evaluar estabilidad temporal de la estimación actual, "
        "disminución de tiempos muertos, continuidad de reporte y sensibilidad de la posición estimada frente a ajustes de ventana de detección, "
        "memoria por antena y mínimos de muestras por antena ESP32."
    )
    add_table_title(doc, "Tabla 30.1 - Cambios técnicos previstos para la tercera prueba respecto de la segunda prueba con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Componente", "Segunda prueba", "Tercera prueba planificada", "Objetivo técnico"],
        [
            ["Firmware de antena ESP32", "Captura por ciclos", "Captura continua con conexión persistente", "Disminuir demoras de reconexión y estabilizar el flujo RAW"],
            ["Firmware del emisor objetivo", "Emisión no dedicada a canal fijo", "Emisión continua en el mismo canal del AP", "Aumentar coincidencia temporal entre transmisión y escucha"],
            ["Modo de operación", "Campaña con entrenamiento e inferencia posterior", "Solo inferencia sobre campaña ya entrenada", "Separar evaluación en línea del costo de la campaña"],
            ["Ventanas temporales", "Amplias", "Reducidas y parametrizadas", "Reducir arrastre de historial y demoras de consolidación"],
            ["Condición de aceptación", "Disponibilidad general del mapa de calor", "Continuidad del flujo desde las cuatro antenas ESP32", "Priorizar estabilidad de operación y latencia baja"],
        ],
    )
    add_table_title(doc, "Tabla 30.2 - Plantilla de resultados prevista para la tercera prueba de inferencia en línea continua.")
    add_table(
        doc,
        ["Escenario", "Ventana de detección [s]", "Memoria por antena [s]", "Mínimo por antena", "Antenas requeridas", "Dispositivo evaluado", "Observación resultante", "Error estimado o cualitativo"],
        [
            ["T3_E01", "Completar", "Completar", "Completar", "Completar", "TAG", "Completar luego de la prueba", "Completar luego de la prueba"],
            ["T3_E02", "Completar", "Completar", "Completar", "Completar", "TAG", "Completar luego de la prueba", "Completar luego de la prueba"],
            ["T3_E03", "Completar", "Completar", "Completar", "Completar", "Multidispositivo", "Completar luego de la prueba", "Completar luego de la prueba"],
        ],
    )
    add_body_paragraph(
        doc,
        "En términos de criterio de éxito, la tercera prueba deberá demostrar tres propiedades mínimas: continuidad de reporte desde las "
        "cuatro antenas ESP32 durante la inferencia, estabilidad visual del punto estimado del emisor objetivo y reducción del tiempo "
        "necesario para disponer de una predicción utilizable. Si estas condiciones se verifican, la tercera prueba aportará una respuesta "
        "concreta al principal límite observado en la segunda campaña."
    )

    doc.add_page_break()
    add_heading(doc, "31. Registro visual de la segunda campaña con cuatro antenas ESP32", level=1)
    add_body_paragraph(
        doc,
        "Las capturas de pantalla siguientes documentan la secuencia operativa real de la segunda prueba. Su incorporación al informe es "
        "relevante porque permite asociar la evidencia cuantitativa de los archivos persistidos con el comportamiento observado en la interfaz "
        "de operación durante la campaña y durante la inferencia en línea posterior."
    )
    contact_sheet = create_contact_sheet(screenshot_paths, GENERATED_DIR / "anexo_capturas_prueba2_4_antenas.png")
    add_picture_with_caption(
        doc,
        contact_sheet,
        "Figura 31.1. Hoja de contacto con el conjunto completo de capturas registradas el 23 de marzo de 2026 durante la segunda prueba con cuatro antenas ESP32.",
    )

    caption_map = build_screenshot_caption_map()
    for screenshot in screenshot_paths:
        caption = caption_map.get(
            screenshot.name,
            f"Figura 31.X. Captura operativa correspondiente a {screenshot.name}.",
        )
        add_picture_with_caption(doc, screenshot, caption)

    doc.add_page_break()
    add_heading(doc, "32. Espacios reservados para fotografías de campo y verificación espacial", level=1)
    add_body_paragraph(
        doc,
        "Las fotografías de campo constituyen un complemento esencial para la defensa metodológica del trabajo, ya que permiten asociar cada "
        "resultado cuantitativo con una posición física concreta del emisor objetivo, con la geometría real del recinto y con el procedimiento "
        "aplicado durante la campaña. A continuación se dejan espacios reservados para insertar esas imágenes una vez seleccionadas y depuradas."
    )
    placeholders = [
        (
            "Figura 32.1. Fotografía del montaje general de la segunda prueba con cuatro antenas ESP32.",
            create_placeholder_image(
                "Insertar fotografía del montaje general",
                "Incluir vista amplia del recinto, ubicación de las cuatro antenas ESP32,\nrouter de soporte y orientación del plano empleado en la interfaz.",
                GENERATED_DIR / "placeholder_montaje_general_4_antenas.png",
            ),
        ),
        (
            "Figura 32.2. Fotografía de la posición del emisor objetivo durante una prueba de inferencia en línea.",
            create_placeholder_image(
                "Insertar fotografía del TAG durante inferencia",
                "Incluir marca visual del punto real del TAG,\nfecha, escenario y referencia a la captura de pantalla correspondiente.",
                GENERATED_DIR / "placeholder_tag_inferencia_online.png",
            ),
        ),
        (
            "Figura 32.3. Fotografía del procedimiento de medición de un punto de entrenamiento.",
            create_placeholder_image(
                "Insertar fotografía de medición de punto de entrenamiento",
                "Mostrar el TAG en el punto rotulado, la referencia sobre el piso\ny la correspondencia con la tabla de puntos del experimento.",
                GENERATED_DIR / "placeholder_punto_entrenamiento.png",
            ),
        ),
    ]
    for caption, image_path in placeholders:
        add_picture_with_caption(doc, image_path, caption)
        add_body_paragraph(
            doc,
            "Espacio reservado para incorporar la fotografía definitiva seleccionada por el autor, acompañada de la descripción puntual del "
            "escenario, la fecha, el punto de referencia y el objetivo técnico de la imagen."
        )


def main() -> None:
    ensure_dirs()
    run3 = load_run_summary(RUN_3A)
    run4 = load_run_summary(RUN_4A)
    screenshots = find_screenshots()

    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"No se encontró el documento fuente: {SOURCE_DOC}")
    if not screenshots:
        raise FileNotFoundError(f"No se encontraron capturas del 23 de marzo de 2026 en {SCREENSHOT_DIR}")

    doc = Document(str(SOURCE_DOC))
    ref_index = next(
        i for i, paragraph in enumerate(doc.paragraphs)
        if paragraph.text.strip().startswith("28. Referencias técnicas complementarias")
    )
    reference_texts = [paragraph.text.strip() for paragraph in doc.paragraphs[ref_index + 1 :] if paragraph.text.strip()]

    for paragraph in list(doc.paragraphs[ref_index:]):
        remove_paragraph(paragraph)

    append_trial2_results(doc, run3, run4, screenshots)
    append_existing_references(doc, reference_texts)
    doc.save(str(OUTPUT_DOC))
    print(f"Documento generado en: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
