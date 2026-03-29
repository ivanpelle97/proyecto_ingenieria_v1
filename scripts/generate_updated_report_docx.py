from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Iterable

import matplotlib

matplotlib.use("Agg")
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches


ROOT = Path(__file__).resolve().parents[1]
TEMPLATE_DOCX = Path(r"C:\Users\hi_iv\Downloads\PROYECTO_DE_INGENIERIA_parte2_3_living_TAG_ESP32_v9.docx")
OUTPUT_DOCX = ROOT / "PROYECTO_DE_INGENIERIA_parte2_3_v14_tesis_extendida.docx"
EXAMPLE_FIGS_DIR = ROOT / "figs_example"
GENERATED_FIGS_DIR = ROOT / "doc_assets" / "generated"

RUN_1A = ROOT / "runs" / "demo_s1" / "train_01"
RUN_1B = ROOT / "runs" / "demo_s2" / "train_02"

TRIAL_2_CAMPAIGNS = [
    ("P2_C01_base_4A", "Campana base con cuatro antenas ESP32 en las cuatro esquinas del recinto."),
    ("P2_C02_layout_4A", "Campana con modificacion controlada del mobiliario y misma geometria de antenas."),
    ("P2_C03_robustez_4A", "Campana de repetibilidad para verificar estabilidad temporal y sensibilidad a cambios menores."),
]


def clear_document(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag == qn("w:sectPr"):
            continue
        body.remove(child)


def add_heading(doc: Document, text: str, level: int) -> None:
    doc.add_paragraph(text, style=f"Heading {level}")


def add_para(
    doc: Document,
    text: str,
    style: str = "Normal",
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    paragraph = doc.add_paragraph(text, style=style)
    if align is not None:
        paragraph.alignment = align


def add_paragraphs(doc: Document, paragraphs: Iterable[str]) -> None:
    for paragraph in paragraphs:
        add_para(doc, paragraph)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], style: str = "Table Grid"):
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = style
    except KeyError:
        pass
    for idx, value in enumerate(headers):
        table.rows[0].cells[idx].text = str(value)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    return table


def add_picture_with_caption(doc: Document, image_path: Path, caption_text: str, width_inches: float = 6.4) -> None:
    if not image_path.exists():
        add_caption_placeholder(doc, f"[ESPACIO RESERVADO: falta la figura {image_path.name}]", caption_text)
        return
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    add_para(doc, caption_text, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_caption_placeholder(doc: Document, reserved_text: str, caption_text: str) -> None:
    add_para(doc, reserved_text)
    add_para(doc, caption_text, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_photo_placeholder(doc: Document, figure_id: str, description: str) -> None:
    out_dir = prepare_generated_figure_dir()
    safe_name = re.sub(r"[^a-zA-Z0-9_]+", "_", figure_id.strip()).strip("_").lower()
    image_path = out_dir / f"placeholder_{safe_name}.png"
    if not image_path.exists():
        fig, ax = plt.subplots(figsize=(8, 4.5))
        ax.set_facecolor("white")
        ax.set_xticks([])
        ax.set_yticks([])
        for spine in ax.spines.values():
            spine.set_edgecolor("#666666")
            spine.set_linewidth(1.5)
            spine.set_linestyle("--")
        ax.text(0.5, 0.72, figure_id, ha="center", va="center", fontsize=18, fontweight="bold")
        ax.text(0.5, 0.48, "Espacio reservado para imagen real", ha="center", va="center", fontsize=14)
        ax.text(0.5, 0.24, description, ha="center", va="center", fontsize=11, wrap=True)
        fig.tight_layout()
        fig.savefig(image_path, dpi=180, bbox_inches="tight")
        plt.close(fig)
    add_picture_with_caption(doc, image_path, f"{figure_id}. {description}", width_inches=6.3)


def safe_read_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def format_seconds(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "N/D"
    return f"{float(value):.1f}"


def format_mebibytes(value: int | float) -> str:
    return f"{float(value) / (1024.0 * 1024.0):.2f}"


def load_run_detail(base: Path) -> dict:
    manifest = safe_read_json(base / "experiment.json")
    points_df = pd.read_csv(base / "points.csv") if (base / "points.csv").exists() else pd.DataFrame()
    samples_df = pd.read_csv(base / "samples.csv") if (base / "samples.csv").exists() else pd.DataFrame()
    dataset_df = pd.read_csv(base / "dataset.csv") if (base / "dataset.csv").exists() else pd.DataFrame()
    metrics_path = base / "models" / "metrics.csv"
    metrics_df = pd.read_csv(metrics_path) if metrics_path.exists() else pd.DataFrame()

    raw_sizes: dict[str, int] = {}
    raw_dir = base / "raw"
    if raw_dir.exists():
        for path in sorted(raw_dir.glob("*.jsonl")):
            raw_sizes[path.stem] = path.stat().st_size

    if not points_df.empty:
        points_df = points_df.copy()
        points_df["point_num"] = points_df["point_id"].str.extract(r"(\d+)").astype(int)
        points_df = points_df.sort_values("point_num").drop(columns=["point_num"])
        points_df["capture_s"] = (
            pd.to_datetime(points_df["capture_complete_at"]) - pd.to_datetime(points_df["started_at"])
        ).dt.total_seconds()
        points_df["close_delay_s"] = (
            pd.to_datetime(points_df["completed_at"]) - pd.to_datetime(points_df["capture_complete_at"])
        ).dt.total_seconds()
    else:
        points_df = pd.DataFrame(
            columns=[
                "point_id",
                "x_m",
                "y_m",
                "z_m",
                "capture_s",
                "close_delay_s",
                "status",
                "samples_per_anchor",
            ]
        )

    samples_per_antenna = (
        samples_df.groupby("anchor_id").size().to_dict()
        if not samples_df.empty and "anchor_id" in samples_df.columns
        else {}
    )
    points_complete = int((points_df["status"] == "complete").sum()) if "status" in points_df.columns else int(len(points_df))
    antenna_count = len(manifest.get("anchors", []))
    samples_target = int(manifest.get("samples_per_anchor", 0))
    expected_samples = points_complete * antenna_count * samples_target
    full_grid_expected_samples = 20 * antenna_count * samples_target
    dataset_exists = (base / "dataset.csv").exists()
    metrics_exist = metrics_path.exists()

    best_p95 = metrics_df.sort_values(["p95", "mae_eucl"]).iloc[0].to_dict() if not metrics_df.empty else None
    best_mae = metrics_df.sort_values(["mae_eucl", "p95"]).iloc[0].to_dict() if not metrics_df.empty else None

    detail = {
        "base": base,
        "session_id": manifest.get("session_id", ""),
        "campaign_id": manifest.get("campaign_id", ""),
        "target_mac": manifest.get("target_mac", ""),
        "samples_per_anchor": samples_target,
        "environment": manifest.get("environment", {}),
        "antennas": manifest.get("anchors", []),
        "points_df": points_df,
        "samples_df": samples_df,
        "dataset_df": dataset_df,
        "metrics_df": metrics_df,
        "raw_sizes": raw_sizes,
        "raw_total_bytes": int(sum(raw_sizes.values())),
        "samples_per_antenna": samples_per_antenna,
        "points_count": int(len(points_df)),
        "points_complete": points_complete,
        "samples_rows": int(len(samples_df)),
        "dataset_rows": int(len(dataset_df)),
        "dataset_cols": int(len(dataset_df.columns)) if not dataset_df.empty else 0,
        "dataset_exists": dataset_exists,
        "metrics_exist": metrics_exist,
        "best_p95": best_p95,
        "best_mae": best_mae,
        "expected_samples": expected_samples,
        "full_grid_expected_samples": full_grid_expected_samples,
    }

    if not points_df.empty:
        detail["capture_mean_s"] = float(points_df["capture_s"].mean())
        detail["capture_median_s"] = float(points_df["capture_s"].median())
        detail["capture_min_s"] = float(points_df["capture_s"].min())
        detail["capture_max_s"] = float(points_df["capture_s"].max())
        detail["close_delay_mean_s"] = float(points_df["close_delay_s"].mean())
        detail["longest_point"] = points_df.loc[points_df["capture_s"].idxmax()].to_dict()
        detail["shortest_point"] = points_df.loc[points_df["capture_s"].idxmin()].to_dict()
    else:
        detail["capture_mean_s"] = None
        detail["capture_median_s"] = None
        detail["capture_min_s"] = None
        detail["capture_max_s"] = None
        detail["close_delay_mean_s"] = None
        detail["longest_point"] = None
        detail["shortest_point"] = None

    return detail


def prepare_generated_figure_dir() -> Path:
    GENERATED_FIGS_DIR.mkdir(parents=True, exist_ok=True)
    return GENERATED_FIGS_DIR


def make_architecture_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(11, 6))
    ax.axis("off")
    boxes = {
        "Emisor\nobjetivo\nESP32": (0.08, 0.55),
        "Antenas\nESP32\nreceptoras": (0.30, 0.55),
        "Servidor\nFastAPI\n+ Uvicorn": (0.54, 0.55),
        "Interfaz\nStreamlit": (0.78, 0.72),
        "Pipeline de\ndataset y\nentrenamiento": (0.78, 0.38),
    }
    for label, (x, y) in boxes.items():
        ax.text(
            x,
            y,
            label,
            ha="center",
            va="center",
            fontsize=12,
            bbox=dict(boxstyle="round,pad=0.6", facecolor="#e9f2fb", edgecolor="#3b73b9", linewidth=2),
        )
    arrows = [
        ((0.14, 0.55), (0.24, 0.55), "Probe Requests"),
        ((0.36, 0.55), (0.48, 0.55), "JSON / HTTP"),
        ((0.60, 0.60), (0.72, 0.69), "Estado y control"),
        ((0.60, 0.49), (0.72, 0.41), "CSV / modelos"),
        ((0.78, 0.58), (0.78, 0.49), "Inferencia"),
    ]
    for start, end, label in arrows:
        ax.annotate("", xy=end, xytext=start, arrowprops=dict(arrowstyle="->", lw=2, color="#444"))
        ax.text((start[0] + end[0]) / 2, (start[1] + end[1]) / 2 + 0.03, label, ha="center", fontsize=10)
    ax.text(
        0.5,
        0.10,
        "Arquitectura general del sistema de localizacion indoor basado en RSSI y antenas ESP32",
        ha="center",
        fontsize=12,
        color="#333",
    )
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_software_stack_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(9, 8))
    ax.axis("off")
    layers = [
        ("Usuario / operador", "#f8d9b6"),
        ("Interfaz de operacion Streamlit", "#d9ead3"),
        ("Servidor FastAPI sobre Uvicorn", "#cfe2f3"),
        ("Procesamiento con pandas y NumPy", "#d9d2e9"),
        ("Modelos scikit-learn + joblib", "#ffe599"),
        ("Firmware ESP32 + Arduino IDE", "#ead1dc"),
    ]
    y = 0.88
    for label, color in layers:
        ax.text(
            0.5,
            y,
            label,
            ha="center",
            va="center",
            fontsize=12,
            bbox=dict(boxstyle="round,pad=0.6", facecolor=color, edgecolor="#555", linewidth=1.5),
        )
        if y > 0.18:
            ax.annotate("", xy=(0.5, y - 0.08), xytext=(0.5, y - 0.01), arrowprops=dict(arrowstyle="->", lw=2))
        y -= 0.14
    ax.text(0.5, 0.05, "Estratificacion del software utilizado en el prototipo", ha="center", fontsize=12)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_flow_figure(path: Path, title: str, steps: list[str], color: str) -> None:
    fig, ax = plt.subplots(figsize=(11, 3.8))
    ax.axis("off")
    for idx, step in enumerate(steps):
        x = 0.08 + idx * (0.84 / max(1, len(steps) - 1))
        ax.text(
            x,
            0.55,
            step,
            ha="center",
            va="center",
            fontsize=11,
            bbox=dict(boxstyle="round,pad=0.5", facecolor=color, edgecolor="#444", linewidth=1.5),
        )
        if idx < len(steps) - 1:
            x2 = 0.08 + (idx + 1) * (0.84 / max(1, len(steps) - 1))
            ax.annotate("", xy=(x2 - 0.06, 0.55), xytext=(x + 0.06, 0.55), arrowprops=dict(arrowstyle="->", lw=2))
    ax.text(0.5, 0.88, title, ha="center", fontsize=14, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_run_comparison_figure(path: Path, run_1a: dict, run_1b: dict) -> None:
    fig, ax = plt.subplots(figsize=(9, 5))
    labels = ["Puntos completos", "Muestras aceptadas", "Tiempo medio por punto [s]"]
    run_1a_values = [run_1a["points_complete"], run_1a["samples_rows"], run_1a["capture_mean_s"] or 0.0]
    run_1b_values = [run_1b["points_complete"], run_1b["samples_rows"], run_1b["capture_mean_s"] or 0.0]
    x = range(len(labels))
    ax.bar([i - 0.18 for i in x], run_1a_values, width=0.36, label="demo_s1 / train_01", color="#d97b66")
    ax.bar([i + 0.18 for i in x], run_1b_values, width=0.36, label="demo_s2 / train_02", color="#5b9bd5")
    ax.set_xticks(list(x))
    ax.set_xticklabels(labels)
    ax.set_title("Comparacion global entre las dos corridas reales con 3 antenas ESP32")
    ax.legend()
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_capture_times_figure(path: Path, run_1b: dict) -> None:
    df = run_1b["points_df"]
    fig, ax = plt.subplots(figsize=(11, 4.8))
    ax.bar(df["point_id"], df["capture_s"], color="#5b9bd5")
    ax.set_title("Tiempo de captura por punto en la corrida principal con 3 antenas ESP32")
    ax.set_xlabel("Punto de referencia")
    ax.set_ylabel("Tiempo de captura [s]")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_metrics_figure(path: Path, run_1b: dict) -> None:
    df = run_1b["metrics_df"].sort_values("p95")
    fig, ax = plt.subplots(figsize=(10, 5))
    ax.bar(df["model"], df["mae_eucl"], label="MAE", color="#7cc576")
    ax.plot(df["model"], df["p95"], marker="o", color="#d9534f", linewidth=2, label="P95")
    ax.set_title("Comparacion de modelos en la corrida principal con 3 antenas ESP32")
    ax.set_ylabel("Error [m]")
    ax.grid(axis="y", alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_raw_sizes_figure(path: Path, run_1b: dict) -> None:
    names = list(run_1b["raw_sizes"].keys())
    values = [run_1b["raw_sizes"][name] / (1024.0 * 1024.0) for name in names]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.bar(names, values, color="#ed7d31")
    ax.set_title("Volumen RAW por antena ESP32 en la corrida principal")
    ax.set_ylabel("Tamano de archivo [MiB]")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def make_points_map_figure(path: Path, run_1b: dict) -> None:
    env = run_1b["environment"]
    points = run_1b["points_df"]
    antennas = run_1b["antennas"]
    fig, ax = plt.subplots(figsize=(9, 5.2))
    ax.set_xlim(0, float(env.get("length_m", 6.0)))
    ax.set_ylim(0, float(env.get("width_m", 4.0)))
    ax.grid(alpha=0.25)
    ax.scatter(points["x_m"], points["y_m"], color="#2ca6a4", s=45, label="Puntos relevados")
    for _, row in points.iterrows():
        ax.text(float(row["x_m"]) + 0.05, float(row["y_m"]) + 0.03, row["point_id"], fontsize=8)
    for antenna in antennas:
        ax.scatter([antenna["x_m"]], [antenna["y_m"]], marker="^", s=110, color="#1f4e79")
        ax.text(float(antenna["x_m"]) + 0.05, float(antenna["y_m"]) + 0.05, antenna["anchor_id"], fontsize=10)
    ax.set_title("Distribucion espacial de puntos y antenas ESP32 en la corrida principal")
    ax.set_xlabel("x [m]")
    ax.set_ylabel("y [m]")
    ax.legend()
    fig.tight_layout()
    fig.savefig(path, dpi=200, bbox_inches="tight")
    plt.close(fig)


def build_generated_figures(run_1a: dict, run_1b: dict) -> dict[str, Path]:
    out_dir = prepare_generated_figure_dir()
    figures = {
        "architecture": out_dir / "fig_architecture_overview.png",
        "software_stack": out_dir / "fig_software_stack.png",
        "firmware_flow": out_dir / "fig_firmware_flow.png",
        "server_flow": out_dir / "fig_server_flow.png",
        "pipeline_flow": out_dir / "fig_pipeline_flow.png",
        "inference_flow": out_dir / "fig_inference_flow.png",
        "run_compare": out_dir / "fig_run_compare.png",
        "run_1b_capture_times": out_dir / "fig_run_1b_capture_times.png",
        "run_1b_metrics": out_dir / "fig_run_1b_metrics.png",
        "run_1b_raw_sizes": out_dir / "fig_run_1b_raw_sizes.png",
        "run_1b_points_map": out_dir / "fig_run_1b_points_map.png",
    }
    make_architecture_figure(figures["architecture"])
    make_software_stack_figure(figures["software_stack"])
    make_flow_figure(
        figures["firmware_flow"],
        "Secuencia funcional de una antena ESP32 receptora",
        ["Modo promiscuo", "Barrido\nde canal", "Buffer local", "Modo estacion", "GET /config", "POST /ingest"],
        "#d9ead3",
    )
    make_flow_figure(
        figures["server_flow"],
        "Flujo del servidor y persistencia experimental",
        ["Recepcion\ndel lote", "RAW JSONL", "Filtro por MAC", "samples.csv", "points.csv", "training_state.json"],
        "#cfe2f3",
    )
    make_flow_figure(
        figures["pipeline_flow"],
        "Pipeline de dataset y entrenamiento",
        ["raw/*.jsonl", "samples.csv", "dataset.csv", "modelos", "metrics.csv", "joblib"],
        "#ffe599",
    )
    make_flow_figure(
        figures["inference_flow"],
        "Flujo de inferencia en linea y mapa de calor",
        ["Paquetes recientes", "Vector de\ncaracteristicas", "Modelo", "Prediccion x,y", "Mapa de calor", "Plano CAD"],
        "#ead1dc",
    )
    make_run_comparison_figure(figures["run_compare"], run_1a, run_1b)
    make_capture_times_figure(figures["run_1b_capture_times"], run_1b)
    make_metrics_figure(figures["run_1b_metrics"], run_1b)
    make_raw_sizes_figure(figures["run_1b_raw_sizes"], run_1b)
    make_points_map_figure(figures["run_1b_points_map"], run_1b)
    return figures


def add_common_intro(doc: Document) -> None:
    add_para(
        doc,
        "El marco teorico ya desarrolla de forma extensa los fundamentos fisicos de la propagacion en interiores, las "
        "metricas RSSI y CSI, la taxonomia de metodos de posicionamiento y los principios del fingerprinting. Por esa "
        "razon, el presente documento concentra su esfuerzo en describir la implementacion real del sistema, su "
        "protocolo experimental y los resultados obtenidos con hardware concreto. La Parte II se enfoca en el diseno y "
        "la implementacion, mientras que la Parte III se dedica a resultados, analisis, mejoras y conclusiones.",
    )
    add_para(
        doc,
        "La redaccion se formula con criterio ingenieril, priorizando decisiones de arquitectura, restricciones de "
        "implementacion, trazabilidad de datos y observaciones de campo. Este criterio resulta necesario porque el "
        "proyecto evoluciono desde una formulacion conceptual hacia una plataforma experimental operativa, donde la "
        "coherencia entre firmware, servidor, interfaz y pipeline de aprendizaje automatico pasa a ser tan relevante "
        "como la teoria de propagacion que sustenta el fenomeno radioelectrico.",
    )


def build_part_ii(doc: Document, run_1a: dict, run_1b: dict, figs: dict[str, Path]) -> None:
    env = run_1b["environment"]
    antennas = run_1b["antennas"]

    add_heading(doc, "PARTE II - DESARROLLO E IMPLEMENTACION", 1)
    add_common_intro(doc)

    add_heading(doc, "11. Alcance tecnico, criterios de diseno y objetivos", 2)
    add_paragraphs(
        doc,
        [
            "El sistema implementado persigue la localizacion bidimensional de un emisor objetivo basado en ESP32 dentro de un recinto domestico. La solucion se fundamenta en la captura de indicadores RSSI asociados a tramas IEEE 802.11, emitidas por el dispositivo objetivo y observadas por un conjunto de antenas ESP32 receptoras. La posicion final se estima mediante modelos supervisados entrenados con datos medidos en el propio ambiente de interes.",
            "A diferencia de otros esquemas de posicionamiento en interiores, la implementacion actual no depende de una infraestructura comercial de puntos de acceso con soporte propietario ni de informacion de tiempo de vuelo. El diseno se orienta a hardware de bajo costo y ampliamente disponible, aceptando la variabilidad inherente del RSSI como parte del problema experimental y compensandola mediante una estrategia de fingerprinting y aprendizaje supervisado.",
            "El alcance inmediato del prototipo comprende un living rectangular de aproximadamente 5,58 m x 3,6 m x 2,5 m, una primera configuracion real con tres antenas ESP32 y una segunda etapa planificada con cuatro antenas ESP32. Esta delimitacion espacial y material no implica una restriccion conceptual del sistema, sino una decision metodologica destinada a controlar la complejidad experimental y a obtener resultados auditables en un ambiente realista.",
            "Desde el punto de vista funcional, la plataforma debe permitir crear una campana, definir la geometria del ambiente, registrar la posicion de cada antena ESP32, cargar manualmente los puntos de referencia, capturar un numero fijo de muestras por punto, construir un dataset supervisado, entrenar varios modelos de regresion y reutilizar el modelo seleccionado en inferencia en linea. Cada etapa debe dejar evidencia persistente para soportar posteriores auditorias y comparaciones metodologicas.",
            "El criterio de validacion adoptado no se limita a una unica metrica de error. Tambien se considera parte del exito del sistema la posibilidad de completar los puntos definidos, mantener estabilidad operativa durante las mediciones, regenerar el dataset a partir de la evidencia primaria y reproducir el entrenamiento. Este enfoque es deliberadamente ingenieril: un sistema experimental valioso no es solo el que entrega un numero de error bajo, sino el que conserva coherencia operativa y trazabilidad completa.",
        ],
    )
    add_para(doc, "Tabla 11.1 - Objetivos tecnicos y criterios de validacion.")
    add_table(
        doc,
        ["Aspecto", "Criterio adoptado"],
        [
            ["Objeto de estimacion", "Coordenadas bidimensionales del emisor objetivo."],
            ["Tecnologia de captura", "RSSI de tramas IEEE 802.11 observadas por antenas ESP32."],
            ["Unidad de entrenamiento", "Punto de referencia etiquetado manualmente."],
            ["Completitud por punto", f"{run_1b['samples_per_anchor']} muestras por antena ESP32 en la corrida principal."],
            ["Persistencia minima", "RAW, muestras aceptadas, puntos cerrados, dataset y metricas."],
            ["Reproducibilidad", "Regeneracion del documento y del modelo desde archivos del repositorio."],
        ],
    )
    add_picture_with_caption(
        doc,
        figs["architecture"],
        "Figura 11.1. Arquitectura general del sistema implementado, desde el emisor objetivo hasta el servidor y el pipeline de entrenamiento.",
        width_inches=6.5,
    )

    add_heading(doc, "12. Plataforma de hardware", 2)
    add_paragraphs(
        doc,
        [
            "La plataforma fisica se basa en placas ESP32, seleccionadas por su disponibilidad, bajo costo, soporte de Wi-Fi en la banda de 2,4 GHz y acceso programatico a funciones de bajo nivel del subsistema radio. Esta combinacion hace posible construir un sistema de medicion y localizacion sin recurrir a equipamiento de laboratorio ni a receptores especializados de alto costo, manteniendo al mismo tiempo un grado de control suficiente para las necesidades del proyecto.",
            "En el sistema se distinguen tres familias de elementos fisicos: el emisor objetivo, las antenas ESP32 receptoras y la infraestructura de soporte. El emisor objetivo cumple la funcion de generar las tramas cuya posicion desea inferirse. Las antenas ESP32 receptoras capturan informacion radioelectrica, la empaquetan y la remiten al servidor. La infraestructura de soporte incluye la computadora que ejecuta el servidor y la interfaz, el punto de acceso Wi-Fi de la red local, la alimentacion y el cableado necesario para el mantenimiento y el monitoreo.",
            "Esta organizacion refleja una separacion funcional clara. El emisor objetivo no estima su propia posicion ni interactua con el usuario. Las antenas ESP32 receptoras no resuelven el problema de localizacion, sino que actuan como nodos de observacion distribuidos. La computadora central asume la coordinacion de campanas, el almacenamiento de datos, la construccion del dataset y la ejecucion de los modelos. Esta particion simplifica la instrumentacion y facilita el diagnostico de fallas en campo.",
        ],
    )
    add_heading(doc, "12.1 Emisor objetivo con ESP32", 3)
    add_paragraphs(
        doc,
        [
            "El emisor objetivo se implementa con una placa ESP32 configurada para transmitir Probe Requests con una direccion MAC fija y conocida. La adopcion de una direccion MAC controlada es esencial para la logica experimental, ya que permite al servidor discriminar entre trafico perteneciente al emisor objetivo y trafico ajeno capturado por las antenas ESP32. Sin esta identificacion explicita, el sistema no podria construir un conjunto de entrenamiento etiquetado con la consistencia necesaria.",
            "Desde el punto de vista de ingenieria, la ESP32 utilizada como emisor objetivo no pretende replicar exactamente el comportamiento estadistico de todos los telefonos inteligentes del mercado. En cambio, se la configura con un perfil de emision razonablemente cercano al de un dispositivo movil comun, incorporando barrido multicanal, potencia moderada y un periodo de transmision controlado. Esta aproximacion permite realizar ensayos repetibles sin perder del todo el realismo operacional.",
            "El emisor objetivo constituye, por lo tanto, una fuente controlada de tramas Wi-Fi. Su papel es doble. Primero, facilita la validacion end to end durante las fases iniciales del proyecto. Segundo, proporciona la referencia necesaria para construir la huella radioelectrica del ambiente durante las campanas de medicion. En ambos casos, su configuracion estable mejora la comparabilidad entre corridas y simplifica el analisis posterior.",
        ],
    )
    add_heading(doc, "12.2 Antenas ESP32 receptoras", 3)
    add_paragraphs(
        doc,
        [
            "Las antenas ESP32 receptoras representan los nodos de observacion del sistema. Cada una alterna entre una fase de captura en modo promiscuo y una fase de reconexion a la red local en modo estacion. Durante la primera fase registra tramas IEEE 802.11, RSSI, canal y tiempo local. Durante la segunda consulta la configuracion vigente en el servidor y transmite los lotes capturados mediante HTTP. Esta alternancia es una concesion impuesta por la plataforma ESP32, pero tambien es una solucion practicamente validada en las pruebas realizadas.",
            "La expresion antena ESP32 se utiliza en este documento para referirse al nodo receptor completo, es decir, a la placa programada, su subsistema radio y la posicion geometrica que ocupa en el ambiente. Desde el punto de vista teorico, el elemento radiante es solo una parte del receptor; sin embargo, para la descripcion experimental interesa la unidad funcional completa, porque cada nodo introduce su propia dinamica de captura, reconexion y sesgo sistematico de medicion.",
            "El numero de antenas ESP32 condiciona directamente la geometria del problema. Con tres receptores se obtiene una primera capacidad de observacion suficiente para una validacion experimental, aunque con asimetrias geometricas relevantes. Con cuatro receptores ubicados en las cuatro esquinas del recinto se espera mejorar la observabilidad espacial y reducir ambiguedades asociadas a zonas con baja discriminacion en alguno de los ejes del plano.",
        ],
    )
    add_heading(doc, "12.3 Servidor local y red Wi-Fi de soporte", 3)
    add_paragraphs(
        doc,
        [
            "El servidor local corre sobre una computadora personal que comparte la misma red Wi-Fi de 2,4 GHz con las antenas ESP32 receptoras. Esta decision simplifica la topologia de despliegue y evita dependencias externas innecesarias. La computadora actua como punto de convergencia de la informacion, aloja la logica de sesion y campana, ejecuta la interfaz de operacion, persiste los datos y corre el pipeline de entrenamiento. En consecuencia, el servidor local es el centro nervioso del sistema implementado.",
            "La red Wi-Fi de soporte cumple una funcion instrumental pero critica. Las antenas ESP32 necesitan reconectarse a ella para consultar configuracion y enviar lotes de captura. Por ello, la estabilidad del punto de acceso, el canal utilizado, la calidad de senal en cada ubicacion y la congestion de la banda influyen directamente en la eficiencia operacional de las campanas. Parte de las dificultades observadas durante las primeras pruebas estuvo asociada precisamente a este acoplamiento entre captura radioelectrica y reconexion de uplink.",
            "Desde una perspectiva de ingenieria de sistemas, esta dependencia de la red local no invalida la arquitectura, pero si la convierte en un subsistema mas a monitorear. En el informe se considera, por tanto, que la infraestructura Wi-Fi no es un mero detalle de implementacion, sino un componente que puede limitar el rendimiento global y que debe ser documentado junto con el resto del montaje experimental.",
        ],
    )
    add_heading(doc, "12.4 Montaje fisico y criterios geometricos", 3)
    add_paragraphs(
        doc,
        [
            "La geometria del despliegue se define en un plano cartesiano bidimensional asociado al living. En la corrida principal, las tres antenas ESP32 se situaron aproximadamente en (0,00; 3,60; 2,00), (2,60; 0,50; 2,00) y (5,58; 3,60; 2,00). Esta configuracion privilegia la cobertura del ambiente completo, aunque no conforma un poligono completamente simetrico. De hecho, la ausencia de una cuarta antena ESP32 en la esquina faltante constituye una de las hipotesis principales para explicar ciertas asimetrias de inferencia observadas cualitativamente.",
            "La altura de instalacion de 2,0 m responde a un compromiso entre visibilidad radioelectrica, practicidad de montaje y separacion respecto de objetos de baja altura. Ubicar las antenas ESP32 demasiado cerca del piso incrementaria la influencia de muebles bajos y reflexiones tempranas, mientras que ubicarlas excesivamente alto complicaria el montaje y podria alejar la sensibilidad respecto de dispositivos portados a la altura de la mano o del torso.",
            "Desde el punto de vista metodologico, el criterio geometrico adoptado para la segunda etapa es mas exigente: cuatro antenas ESP32 en las cuatro esquinas del recinto. Esta configuracion no garantiza por si sola una mejora en todas las metricas, pero si reduce la probabilidad de que la observacion quede dominada por un borde del ambiente. Por esta razon, la prueba con cuatro receptores se plantea como la extension natural de la primera experimentacion.",
        ],
    )
    add_para(doc, "Tabla 12.1 - Inventario y despliegue de hardware en la corrida principal.")
    add_table(
        doc,
        ["Elemento", "Cantidad", "Rol tecnico", "Observaciones"],
        [
            ["ESP32 emisora", "1", "Generacion de Probe Requests del emisor objetivo", f"MAC objetivo {run_1b['target_mac']}"],
            ["Antena ESP32 receptora", str(len(antennas)), "Captura de RSSI y envio de lotes al servidor", "Corrida principal completada con tres nodos."],
            ["Servidor local", "1", "Persistencia, interfaz, entrenamiento e inferencia", "Computadora personal con Python."],
            ["Punto de acceso Wi-Fi", "1", "Conectividad de uplink para las antenas ESP32", "Operacion sobre 2,4 GHz."],
        ],
    )
    add_para(doc, "Tabla 12.2 - Coordenadas reales de las antenas ESP32 en la corrida principal.")
    add_table(
        doc,
        ["Antena ESP32", "x [m]", "y [m]", "z [m]", "Modelo", "Notas"],
        [
            [
                antenna["anchor_id"],
                f"{float(antenna['x_m']):.2f}",
                f"{float(antenna['y_m']):.2f}",
                f"{float(antenna['z_m']):.2f}",
                str(antenna.get("model", "ESP32")),
                str(antenna.get("notes", "")),
            ]
            for antenna in antennas
        ],
    )
    if (EXAMPLE_FIGS_DIR / "fig_13_1_plano_living.png").exists():
        add_picture_with_caption(
            doc,
            EXAMPLE_FIGS_DIR / "fig_13_1_plano_living.png",
            "Figura 12.1. Esquema conceptual del living y del sistema de coordenadas utilizado como apoyo de la descripcion experimental.",
            width_inches=6.3,
        )
    add_picture_with_caption(
        doc,
        figs["run_1b_points_map"],
        "Figura 12.2. Distribucion espacial real de puntos relevados y antenas ESP32 en la corrida principal con tres receptores.",
        width_inches=6.5,
    )
    add_photo_placeholder(doc, "Figura 12.3", "Fotografia panoramica del living utilizada para documentar el montaje fisico de la corrida con tres antenas ESP32.")
    add_photo_placeholder(doc, "Figura 12.4", "Fotografias de detalle de la ubicacion fisica de cada antena ESP32 y del emisor objetivo.")

    add_heading(doc, "13. Plataforma de software", 2)
    add_paragraphs(
        doc,
        [
            "La plataforma de software se organiza como una cadena de componentes especializados que cooperan entre si sin superponer responsabilidades. En un extremo se ubica el firmware cargado en cada ESP32; en el centro se encuentran el servidor FastAPI y la interfaz de operacion construida con Streamlit; y en el plano analitico se despliega el pipeline implementado con Python, pandas, NumPy, scikit-learn, matplotlib y joblib. Esta division no es arbitraria. Responde a una necesidad concreta de mantener separados el subsistema de captura, el subsistema de coordinacion y el subsistema de analisis.",
            "Dado que el marco teorico no describe en detalle estas herramientas, la Parte II incorpora una explicacion breve pero suficiente de su rol. El objetivo no es desarrollar un manual de cada biblioteca, sino justificar por que fueron seleccionadas y que problema tecnico resuelven dentro del proyecto. Esta aclaracion es especialmente importante porque varios resultados experimentales dependen tanto de decisiones de software como de cuestiones puramente radioelectricas.",
            "La eleccion de un stack abierto y ampliamente documentado facilita la reproducibilidad del trabajo. Un tercero que disponga del repositorio, de las dependencias adecuadas y del hardware puede reconstruir los artefactos principales: levantar el servidor, operar la interfaz, regenerar el dataset, reentrenar los modelos y reconstruir el informe tecnico. En el contexto de un proyecto de ingenieria, esta capacidad de reproduccion constituye un criterio de calidad en si mismo.",
        ],
    )
    add_heading(doc, "13.1 Python como lenguaje integrador", 3)
    add_paragraphs(
        doc,
        [
            "Python actua como lenguaje de integracion de la plataforma. Se utiliza para implementar el servidor, la interfaz de operacion, el pipeline de procesamiento, la generacion automatica del informe y diversas tareas auxiliares de validacion. La principal ventaja de esta eleccion reside en la abundancia de bibliotecas cientificas y de desarrollo web disponibles, asi como en la rapidez de iteracion que ofrece para un prototipo experimental en evolucion.",
            "Desde una perspectiva de ingenieria de software, Python permite expresar con claridad estructuras de datos, transformaciones tabulares y flujos de trabajo de entrenamiento sin incurrir en un costo alto de desarrollo. Esto resulta particularmente conveniente en proyectos donde la logica cambia a medida que se avanza en la experimentacion. La flexibilidad del lenguaje facilita adaptar formatos de captura, agregar campos nuevos y regenerar artefactos sin rehacer por completo la plataforma.",
            "La decision de concentrar la logica de alto nivel en un mismo lenguaje tambien reduce la fragmentacion cognitiva del proyecto. El equipo no necesita mantener simultaneamente una pila web en un lenguaje, un pipeline de datos en otro y un sistema de generacion documental en un tercero. Esta uniformidad simplifica mantenimiento, depuracion y transferencia del conocimiento tecnico.",
        ],
    )
    add_heading(doc, "13.2 FastAPI y Uvicorn", 3)
    add_paragraphs(
        doc,
        [
            "FastAPI es el framework web empleado para exponer los endpoints HTTP del servidor. En terminos funcionales, provee la infraestructura necesaria para recibir lotes enviados por las antenas ESP32, devolver configuracion vigente a cada nodo, iniciar y finalizar campanas, abrir y cerrar puntos de referencia y publicar el estado de la sesion. Su utilidad en este proyecto no radica en una complejidad conceptual elevada, sino en la posibilidad de declarar interfaces claras, estructuradas y faciles de mantener.",
            "Uvicorn, por su parte, es el servidor ASGI que ejecuta la aplicacion FastAPI. Conviene distinguir ambos roles. FastAPI define la logica de la aplicacion y la forma en que se responden las solicitudes. Uvicorn es el componente que escucha conexiones entrantes, despacha las solicitudes y mantiene operativa la aplicacion. Esta separacion conceptual resulta importante porque permite comprender el sistema como una combinacion de framework y motor de ejecucion, no como una caja monolitica.",
            "En un escenario experimental como el presente, FastAPI aporta ademas una ventaja metodologica: obliga a formalizar el contrato entre firmware, interfaz y backend. Cada endpoint queda explicitado con sus campos de entrada y salida, lo cual reduce ambiguedades durante la captura. La interfaz de operacion y las antenas ESP32 no intercambian mensajes informales, sino estructuras concretas que pueden validarse, registrarse y depurarse con mayor trazabilidad.",
        ],
    )
    add_heading(doc, "13.3 Streamlit como interfaz de operacion", 3)
    add_paragraphs(
        doc,
        [
            "Streamlit es la herramienta empleada para construir la interfaz de operacion utilizada durante las campanas. Su funcion en el proyecto es ofrecer una vista unica donde se configuran sesiones, se cargan dimensiones del ambiente, se definen puntos de referencia, se monitorea el avance de las capturas, se construyen datasets, se entrenan modelos y se visualiza la inferencia en linea. En consecuencia, Streamlit no se limita a mostrar resultados; participa activamente en la orquestacion experimental.",
            "La razon tecnica para utilizar Streamlit reside en su capacidad para desarrollar interfaces de escritorio basadas en navegador con una cantidad reducida de codigo adicional. Esto permite iterar rapidamente sobre la operacion del sistema sin invertir esfuerzos desproporcionados en front-end tradicional. Dado que la prioridad del proyecto es la validez de las mediciones y no el desarrollo de una aplicacion comercial, esta herramienta resulta adecuada para una plataforma interna de laboratorio y campo.",
            "Tambien es importante destacar que la interfaz no reemplaza al servidor ni ejecuta entrenamiento por si sola. Streamlit actua como capa de interaccion y visualizacion. Cuando el operador presiona un boton para iniciar una campana o construir un dataset, la accion se materializa a traves de llamadas al backend y de funciones del pipeline. Esta distincion conceptual ayuda a evitar confusiones entre interfaz de operacion y logica central del sistema.",
        ],
    )
    add_heading(doc, "13.4 Bibliotecas de procesamiento y aprendizaje automatico", 3)
    add_paragraphs(
        doc,
        [
            "pandas y NumPy conforman la base del procesamiento tabular y numerico. pandas se emplea para leer, ordenar y transformar archivos CSV; calcular estadisticos; agrupar mediciones por antena ESP32 y punto; y construir los dataframes que sirven como entrada al entrenamiento. NumPy proporciona operaciones vectorizadas y una representacion numerica eficiente para las matrices usadas por los algoritmos de aprendizaje automatico.",
            "scikit-learn provee los modelos supervisados evaluados en la corrida principal y las utilidades de entrenamiento y validacion. En el proyecto se emplea para instanciar regresores, separar columnas objetivo y columnas de entrada, ejecutar validacion cruzada por grupos y obtener metricas comparables entre modelos. La biblioteca no aporta conocimiento radioelectrico por si misma; su valor radica en ofrecer implementaciones estables y contrastadas de algoritmos de regresion adecuados para datasets pequenos y medianos.",
            "matplotlib se utiliza para generar graficos tecnicos, tanto dentro de la interfaz como en el proceso de generacion del informe. Su rol es transformar resultados tabulares en figuras interpretable por el operador o por el lector del documento. joblib, finalmente, se utiliza para persistir el modelo seleccionado en disco, de manera que el entrenamiento y la inferencia en linea puedan desacoplarse temporalmente. El modelo no necesita reentrenarse cada vez que la interfaz se abre; basta con cargar el artefacto serializado pertinente.",
        ],
    )
    add_heading(doc, "13.5 Arduino IDE y rol del firmware", 3)
    add_paragraphs(
        doc,
        [
            "Arduino IDE se adopta como entorno practico de compilacion y carga de firmware para las placas ESP32. La herramienta simplifica tareas operativas fundamentales durante la fase experimental: seleccionar puerto serie, compilar sketches, cargar binarios y monitorear trazas de depuracion. Su uso no implica renunciar a conceptos del ecosistema ESP-IDF; por el contrario, el firmware accede a funciones de bajo nivel del subsistema Wi-Fi del ESP32 desde un entorno de trabajo mas accesible.",
            "El firmware constituye la capa software embebida del sistema. Su responsabilidad es netamente distinta de la de Python. Mientras el software de la computadora gestiona campanas y entrenamiento, el firmware gobierna directamente la interfaz radio de cada placa, decide cuando capturar, cuando reconectarse a la red, cuando consultar configuracion y cuando enviar lotes. Esta capa embebida es, por lo tanto, parte integral del sistema de medicion, no un accesorio intercambiable.",
            "La combinacion de Arduino IDE con llamadas especificas al stack Wi-Fi del ESP32 permite un equilibrio util entre rapidez de desarrollo y control tecnico. En un proyecto aplicado, esta relacion costo-beneficio resulta adecuada siempre que se documenten con claridad las limitaciones observadas, entre ellas los reinicios manuales ocasionales, la sensibilidad a la calidad de alimentacion y la necesidad de tolerar fallas de reconexion a la red local.",
        ],
    )
    add_para(doc, "Tabla 13.1 - Componentes principales de software y justificacion tecnica.")
    add_table(
        doc,
        ["Componente", "Tipo", "Funcion en el sistema", "Justificacion de uso"],
        [
            ["Python", "Lenguaje", "Integracion general de backend, interfaz, analitica y generacion documental", "Rapidez de desarrollo y ecosistema cientifico."],
            ["FastAPI", "Framework web", "Definicion de endpoints y logica HTTP del servidor", "Claridad del contrato entre firmware, interfaz y backend."],
            ["Uvicorn", "Servidor ASGI", "Ejecucion de la aplicacion FastAPI", "Despliegue liviano y suficiente para el laboratorio."],
            ["Streamlit", "Interfaz de operacion", "Control de campanas, entrenamiento e inferencia", "Iteracion rapida sobre una interfaz interna."],
            ["pandas / NumPy", "Bibliotecas de datos", "Transformacion tabular y operaciones numericas", "Adecuadas para el pipeline de dataset."],
            ["scikit-learn", "Aprendizaje automatico", "Entrenamiento y evaluacion de regresores", "Implementaciones estables y comparables."],
            ["matplotlib", "Visualizacion", "Graficos tecnicos y mapa de calor", "Control detallado sobre figuras y exportacion."],
            ["joblib", "Persistencia", "Guardado y carga de modelos", "Serializacion simple de artefactos de entrenamiento."],
            ["Arduino IDE", "Entorno de desarrollo", "Compilacion y carga de firmware ESP32", "Operacion practica durante pruebas de campo."],
        ],
    )
    add_picture_with_caption(
        doc,
        figs["software_stack"],
        "Figura 13.1. Estratificacion de software del prototipo, desde el operador hasta el firmware de las placas ESP32.",
        width_inches=5.8,
    )
    add_photo_placeholder(doc, "Figura 13.2", "Captura de la interfaz de operacion Streamlit durante la configuracion de una campana experimental.")
    add_heading(doc, "13.6 Coherencia del stack y razones de no utilizar alternativas mas complejas", 3)
    add_paragraphs(
        doc,
        [
            "Resulta pertinente justificar por que no se adopto, en esta etapa, una arquitectura con base de datos relacional, front-end desacoplado, contenedorizacion completa o servicios distribuidos. La razon principal es metodologica: el proyecto se encuentra en una fase de validacion aplicada donde la prioridad es comprender el fenomeno experimental y asegurar trazabilidad, no maximizar escalabilidad. Introducir capas de complejidad no imprescindibles hubiera elevado el costo de integracion sin aportar necesariamente mejor evidencia sobre la localizacion indoor.",
            "Asimismo, la seleccion de herramientas como FastAPI y Streamlit favorece una cercania estrecha entre quien opera el experimento y quien modifica la logica del sistema. En proyectos de investigacion, esta proximidad suele ser una ventaja, ya que las necesidades cambian a medida que aparecen resultados. Una plataforma excesivamente desacoplada desde el primer momento podria ser elegante desde la arquitectura de software, pero menos eficaz para iterar con rapidez sobre hipotesis experimentales.",
            "En consecuencia, el stack actual debe leerse como una arquitectura deliberadamente suficiente. No es una simplificacion ingenua, sino una eleccion congruente con el estadio del proyecto, la escala del problema, el volumen de datos manejado y el perfil de experimentacion requerido para las etapas descritas en este informe.",
        ],
    )

    add_heading(doc, "14. Arquitectura integral del sistema", 2)
    add_paragraphs(
        doc,
        [
            "La arquitectura integral puede entenderse como un pipeline extremo a extremo que comienza en la emision de tramas Wi-Fi y culmina en una estimacion espacial realizada por un modelo supervisado. Entre ambos extremos intervienen procesos de captura distribuida, reconexion de nodos, persistencia de lotes, filtrado por direccion MAC, consolidacion de muestras, construccion de features y ejecucion de inferencia. Cada etapa introduce sus propias restricciones y, por ende, debe describirse con detalle.",
            "En este sistema, la informacion no fluye solamente en sentido ascendente desde las antenas ESP32 hacia el servidor. Tambien existe un flujo de control descendente: el servidor expone a cada receptor la configuracion activa de la sesion y de la campana. Esto permite modificar MAC objetivo, identificadores de sesion y punto activo sin reprogramar el firmware cada vez. Tal mecanismo aporta flexibilidad operativa y reduce errores humanos durante las mediciones.",
            "La arquitectura implementada se apoyo deliberadamente en archivos persistentes simples y legibles. En lugar de introducir una base de datos relacional o un bus de eventos, se adopto una estructura basada en JSON y CSV por campana. Esta decision responde al tamano moderado del problema y a la conveniencia de inspeccionar datos con herramientas comunes. Aun cuando un sistema de mayor escala podria requerir otra estrategia, la adoptada resulta adecuada para el nivel de complejidad del prototipo actual.",
        ],
    )
    add_heading(doc, "14.1 Flujo de datos extremo a extremo", 3)
    add_paragraphs(
        doc,
        [
            "El flujo de datos comienza cuando el emisor objetivo transmite Probe Requests. Las antenas ESP32, configuradas en modo promiscuo durante la fase de captura, observan tramas de gestion IEEE 802.11, registran su RSSI y conservan informacion minima de direccionamiento y canal. Estas observaciones se acumulan temporalmente en memoria local hasta conformar un lote de envio.",
            "Una vez finalizado el intervalo de captura, cada antena ESP32 abandona el modo promiscuo, vuelve al modo estacion y trata de reconectarse a la red local. Si logra obtener conectividad, consulta al servidor la configuracion al servidor y remite el lote capturado al endpoint de ingesta. El servidor persiste el lote completo como evidencia primaria y, de manera adicional, filtra las mediciones que corresponden a la MAC objetivo y al punto actualmente activo.",
            "Posteriormente, esas muestras aceptadas se transforman en un dataset supervisado mediante agregacion por punto. El entrenamiento opera sobre ese dataset y produce archivos de metricas y un modelo serializado. Durante la inferencia en linea, el sistema reutiliza paquetes recientes almacenados en RAW para componer un nuevo vector de entrada compatible con el formato usado en entrenamiento. El modelo devuelve entonces una estimacion espacial actualizada.",
        ],
    )
    add_heading(doc, "14.2 Endpoints del servidor", 3)
    add_paragraphs(
        doc,
        [
            "El backend expone un conjunto acotado de endpoints cuya semantica se mantuvo estable durante las pruebas principales. Esta estabilidad fue esencial para desacoplar firmware y operacion de interfaz. Mientras el firmware necesita rutas simples y predecibles para consultar configuracion y enviar lotes, la interfaz requiere endpoints orientados a la orquestacion de sesiones y puntos.",
            "La definicion explicita de endpoints tambien constituye un mecanismo de documentacion interna. Permite razonar el sistema en terminos de responsabilidades: algunos endpoints modifican estado, otros publican estado y otros reciben evidencia. Esta separacion conceptual simplifica tanto la depuracion como la generacion del informe tecnico.",
        ],
    )
    add_para(doc, "Tabla 14.1 - Endpoints principales del servidor implementado.")
    add_table(
        doc,
        ["Endpoint", "Metodo", "Funcion"],
        [
            ["/health", "GET", "Verificacion de disponibilidad del servicio."],
            ["/api/session/status", "GET", "Publicacion del estado actual de la sesion y de la campana."],
            ["/api/session/start", "POST", "Creacion e inicializacion de una campana experimental."],
            ["/api/session/finish", "POST", "Cierre administrativo de la campana."],
            ["/api/points/start", "POST", "Apertura del punto de referencia activo."],
            ["/api/points/finish", "POST", "Cierre del punto actual y persistencia de su resumen."],
            ["/api/anchors/{anchor_id}/config", "GET", "Entrega de configuracion vigente a cada antena ESP32."],
            ["/ingest", "POST", "Recepcion de lotes RAW provenientes de las antenas ESP32."],
        ],
    )
    add_heading(doc, "14.3 Estructura de carpetas y persistencia", 3)
    add_paragraphs(
        doc,
        [
            "Cada campana experimental se almacena en una carpeta propia dentro de runs/<session_id>/<campaign_id>. En ella se conservan los archivos experiment.json, training_state.json, raw/*.jsonl, samples.csv, points.csv, dataset.csv y, si corresponde, la subcarpeta models/. Esta estructura persigue un objetivo doble: mantener desacopladas las campanas y facilitar la inspeccion manual de los artefactos.",
            "experiment.json describe la configuracion declarada de la campana: dimensiones del ambiente, MAC objetivo, numero de muestras por antena ESP32 y posiciones de los receptores. training_state.json registra el estado mutable de la sesion y el progreso de cada punto. Los archivos RAW en formato JSONL constituyen la evidencia primaria de captura. samples.csv contiene solo las observaciones aceptadas como pertenecientes al emisor objetivo. points.csv registra el cierre de cada punto. dataset.csv representa la version tabular supervisada apta para entrenamiento.",
            "Esta persistencia por capas permite analizar fallas sin perder contexto. Si un modelo no se entreno, puede verificarse si el problema estuvo en la ausencia de dataset. Si el dataset no pudo construirse, puede inspeccionarse samples.csv. Si las muestras aceptadas no existen, puede auditarse RAW. Esta trazabilidad es especialmente valiosa en proyectos donde la captura depende de varios nodos fisicos y de condiciones radioelectricas cambiantes.",
        ],
    )
    add_picture_with_caption(
        doc,
        figs["server_flow"],
        "Figura 14.1. Flujo de recepcion, filtrado y persistencia en el servidor.",
        width_inches=6.4,
    )
    if (EXAMPLE_FIGS_DIR / "fig_18_1_carpetas.png").exists():
        add_picture_with_caption(
            doc,
            EXAMPLE_FIGS_DIR / "fig_18_1_carpetas.png",
            "Figura 14.2. Esquema conceptual de la organizacion de carpetas y artefactos de una campana experimental.",
            width_inches=6.4,
        )
    add_heading(doc, "14.4 Estados de sesion, campana y punto", 3)
    add_paragraphs(
        doc,
        [
            "La sesion agrupa varias campanas relacionadas y sirve como contenedor logico de un programa experimental mas amplio. La campana es la unidad operativa concreta con una determinada configuracion del ambiente, un conjunto de antenas ESP32 y un objetivo de captura especifico. El punto es la unidad elemental de etiquetado dentro de una campana. Esta jerarquia permite planificar experimentos con multiples escenarios sin perder orden semantico.",
            "El estado del punto es particularmente importante porque define cuando las muestras empiezan a contarse como validas. Mientras no exista un punto activo, los paquetes pueden seguir llegando al servidor y almacenandose en RAW, pero no alimentan samples.csv. Este comportamiento fue fundamental durante las pruebas preliminares, ya que permitio distinguir entre exito de captura radioelectrica y exito de etiquetado experimental.",
            "El cierre explicito del punto, en lugar de un cierre puramente automatico, se adopta como decision metodologica. Obliga al operador a validar que el dispositivo se mantuvo en la posicion prevista y que el conteo de muestras es suficiente. De este modo se evita que una muestra parcialmente capturada pase inadvertida al dataset definitivo.",
        ],
    )
    add_heading(doc, "14.5 Razonamiento ingenieril de la arquitectura", 3)
    add_paragraphs(
        doc,
        [
            "La arquitectura implementada no pretende ser la solucion universal al problema de localizacion en interiores. Su valor reside en haber alcanzado un equilibrio razonable entre simplicidad, trazabilidad y capacidad de extension. Un diseno mas sofisticado podria incorporar sincronizacion mas estricta, streaming continuo o persistencia en base de datos; sin embargo, esas mejoras tambien incrementarian sustancialmente la complejidad del prototipo.",
            "Desde la optica del proyecto, la eleccion actual es adecuada porque permite responder preguntas experimentales concretas: cuantas muestras se logran por punto, cuan estable es la reconexion de las antenas ESP32, que distribucion espacial de puntos se midio realmente, que modelos entregan mejor error y como cambia el comportamiento esperado al pasar de tres a cuatro receptores. La arquitectura, por tanto, esta alineada con los objetivos de la investigacion aplicada.",
        ],
    )
    add_heading(doc, "14.6 Secuencia operativa completa de una campana", 3)
    add_paragraphs(
        doc,
        [
            "Una campana completa puede describirse como una secuencia de once pasos: verificacion del servidor, apertura de la interfaz, declaracion de sesion y campana, definicion del ambiente, carga de coordenadas de las antenas ESP32, verificacion de que cada nodo consulte configuracion, inicio del punto, espera del cupo de muestras, cierre del punto, repeticion hasta completar la grilla y, finalmente, construccion de dataset y entrenamiento. Expresar el flujo de esta manera es util porque convierte una experiencia operativa difusa en un protocolo repetible.",
            "Esta secuencia tambien ayuda a interpretar errores. Si las antenas ESP32 no consultan configuracion, el problema se ubica antes del punto siete. Si hay RAW pero no samples.csv, el problema aparece entre el filtrado por MAC y el punto activo. Si existe samples.csv pero no dataset.csv, la falla se desplaza al pipeline de consolidacion. Esta lectura por etapas permite depurar el sistema de forma estructurada, reduciendo la tentacion de atribuir cualquier inconveniente al modelo de aprendizaje automatico.",
        ],
    )

    add_heading(doc, "15. Firmware", 2)
    add_paragraphs(
        doc,
        [
            "El firmware es el componente que conecta de manera directa la teoria con el entorno fisico. En el emisor objetivo define la frecuencia, el canal y la direccion MAC de las tramas emitidas. En las antenas ESP32 receptoras determina la estrategia de captura, el almacenamiento temporal y la reconexion al servidor. Debido a esta posicion critica dentro del sistema, cualquier limitacion del firmware repercute de forma inmediata sobre la calidad del dataset.",
            "A diferencia del software ejecutado en la computadora, el firmware se encuentra restringido por recursos mas escasos y por un subsistema radio que no puede simultanear cualquier combinacion de modos. Esta condicion explica varias decisiones de diseno, como la captura por lotes, el barrido de canal en una ventana finita y la alternancia entre modo promiscuo y modo estacion. Tales decisiones no son idealmente puras desde el punto de vista teorico, pero son operativamente viables con la plataforma ESP32.",
        ],
    )
    add_heading(doc, "15.1 Firmware del emisor objetivo", 3)
    add_paragraphs(
        doc,
        [
            "El firmware del emisor objetivo transmite Probe Requests con periodicidad configurable y una direccion MAC controlada. El objetivo principal es asegurar que las antenas ESP32 receptoras observen un patron de trafico suficientemente estable para construir la huella radioelectrica del ambiente. Durante las pruebas se ajustaron frecuencia de emision, potencia y estrategia multicanal para aproximar el comportamiento de un dispositivo movil sin sacrificar repetibilidad.",
            "La ventaja metodologica de esta estrategia es evidente: permite conocer con certeza cual es la MAC del dispositivo de interes y evita depender del trafico espontaneo de un telefono o de otro equipo heterogeneo. No obstante, esta simplificacion introduce una limitacion reconocida. El emisor objetivo no reproduce perfectamente la diversidad de chipsets y politicas de ahorro de energia presentes en terminales comerciales. Por ello, la primera etapa debe interpretarse como una caracterizacion controlada y no como una validacion definitiva frente a cualquier dispositivo real.",
        ],
    )
    add_heading(doc, "15.2 Firmware de las antenas ESP32 receptoras", 3)
    add_paragraphs(
        doc,
        [
            "Cada antena ESP32 receptora ejecuta una logica ciclica compuesta por captura, reconexion y envio. Durante la captura, la placa entra en modo promiscuo y realiza barrido de canales. En esta fase se registran tramas de gestion IEEE 802.11 y se extraen RSSI, canal y direcciones MAC relevantes. Las observaciones se almacenan temporalmente en un buffer local hasta completar la ventana de sniffing o alcanzar un umbral de almacenamiento.",
            "Concluida la captura, la placa desactiva el modo promiscuo, vuelve al modo estacion y trata de reconectarse a la red local. Una vez conectada, consulta al servidor la configuracion vigente asociada a su identificador. Esto permite actualizar de forma remota parametros como la sesion, la campana y la MAC objetivo, manteniendo el mismo firmware base y diferenciando nodos unicamente por su identificador de antena ESP32.",
            "El lote se transmite mediante HTTP en formato JSON. El servidor responde con un codigo de estado que permite al nodo determinar si el envio fue exitoso. Si el envio falla o la reconexion no se concreta, el ciclo siguiente intenta recuperarse. Esta estrategia no garantiza ausencia total de perdida de datos, pero si ofrece una operacion suficientemente robusta para el tamano del problema abordado.",
        ],
    )
    add_heading(doc, "15.3 Modo promiscuo, barrido de canal y restricciones", 3)
    add_paragraphs(
        doc,
        [
            "El modo promiscuo permite a la ESP32 observar tramas IEEE 802.11 no destinadas especificamente a ella. Esta capacidad es indispensable para capturar el trafico del emisor objetivo sin asociarse como cliente al punto de acceso. Sin embargo, la observacion efectiva depende del canal en el cual se encuentre el receptor al momento de la transmision. De ahi que el barrido de canal se haya adoptado como mecanismo para ampliar la probabilidad de captura.",
            "El barrido introduce una tension fundamental: cuanto mas tiempo permanece la antena ESP32 en sniffing, mayor es la posibilidad de observar el emisor objetivo; pero tambien mayor es el tiempo durante el cual el nodo no puede reconectarse al servidor. Si, por el contrario, la ventana de captura se reduce excesivamente, mejora la frecuencia de envio pero disminuye la densidad de muestras observadas. La configuracion final surge de un compromiso entre estas dos necesidades.",
            "En la practica, las pruebas mostraron que la estabilidad no depende solo de la teoria del barrido. Tambien influye el router de soporte, la calidad de la alimentacion, la cercania fisica de cada antena ESP32 al punto de acceso y la simultaneidad con la que varios nodos intentan reconectarse. Por ese motivo, el comportamiento del firmware debe analizarse siempre junto con el montaje experimental concreto.",
        ],
    )
    add_heading(doc, "15.4 Reconexion Wi-Fi, envio HTTP y reinicio automatico", 3)
    add_paragraphs(
        doc,
        [
            "Una observacion importante obtenida durante las pruebas fue la tendencia de algunas placas a quedar en estados de reconexion inestable luego de varios ciclos. Para mitigar este comportamiento, el firmware se ajusto de manera que, si la antena ESP32 agota tres intentos consecutivos de reconexion sin recuperar la red, se reinicia automaticamente. Esta medida no resuelve todas las causas de inestabilidad, pero disminuye la necesidad de intervencion manual durante campanas prolongadas.",
            "La reconexion utiliza una logica de reintentos porque el enlace con la red local no siempre es inmediato. El barrido previo, la competencia entre nodos y variaciones de canal pueden demorar la asociacion. Una vez restablecida la conectividad, el lote se remite mediante HTTP al servidor. Esta secuencia captura-reconexion-envio resume una restriccion central del sistema: el mismo hardware radio debe alternar entre medir y reportar.",
            "Desde el punto de vista documental, la incorporacion del reinicio automatico es relevante porque muestra la evolucion del prototipo a partir de problemas observados en campo. No se trata de una caracteristica teorica definida desde el inicio, sino de una mejora de robustez motivada por la evidencia experimental obtenida durante el uso real del sistema.",
        ],
    )
    add_para(doc, "Tabla 15.1 - Responsabilidades del firmware por tipo de nodo.")
    add_table(
        doc,
        ["Nodo", "Responsabilidades principales", "Riesgos operativos observados"],
        [
            ["Emisor objetivo", "Transmitir Probe Requests con MAC fija y parametros de radio controlados", "Patron de emision no identico al de todos los dispositivos comerciales."],
            ["Antena ESP32 receptora", "Capturar tramas, barrer canales, reconectarse, consultar configuracion y enviar lotes", "Bloqueos, fallas de reconexion, variabilidad de captura segun posicion."],
        ],
    )
    add_picture_with_caption(
        doc,
        figs["firmware_flow"],
        "Figura 15.1. Secuencia funcional de una antena ESP32 receptora durante un ciclo de captura y envio.",
        width_inches=6.5,
    )
    add_photo_placeholder(doc, "Figura 15.2", "Captura del monitor serie de una antena ESP32 mostrando mensajes de sniffing, reconexion y envio de lotes.")

    add_heading(doc, "16. Diseno del dataset", 2)
    add_paragraphs(
        doc,
        [
            "El dataset supervisado constituye el nexo formal entre la etapa de medicion y la etapa de aprendizaje automatico. Su unidad de observacion no es el paquete individual, sino el punto de referencia. Para cada punto, el sistema consolida un numero fijo de muestras por antena ESP32 y construye un vector de caracteristicas que combina valores RSSI individuales y estadisticos resumen. Esta eleccion persigue un objetivo metodologico claro: mantener suficiente detalle de la variabilidad intra-punto sin perder una estructura tabular apta para regresion supervisada.",
            "El criterio de fijar un numero de muestras por antena ESP32 surge de la necesidad de controlar la dimensionalidad del problema. Si cada punto pudiera incluir una cantidad arbitraria de paquetes, el dataset resultante seria irregular y dificil de alimentar a modelos tabulares convencionales. Al adoptar un bloque fijo de N muestras por antena, el problema se transforma en un conjunto de filas comparables entre si, aun cuando el costo sea exigir una disciplina de captura durante la campana.",
            "En la corrida principal con tres antenas ESP32 se utilizaron diez muestras por antena y por punto, por lo que cada fila contiene, ademas de las columnas de identificacion y coordenadas, treinta valores RSSI crudos y varios estadisticos adicionales. Esta combinacion conserva informacion mas rica que un simple promedio y permite que el modelo observe tanto el nivel medio de senal como cierta dispersion local.",
        ],
    )
    add_heading(doc, "16.1 Estructura del dataset.csv", 3)
    add_paragraphs(
        doc,
        [
            "El archivo dataset.csv comienza con columnas de identificacion y de verdad de terreno: session_id, campaign_id, point_id, x_m, y_m, z_m y samples_per_anchor. A continuacion se agregan, para cada antena ESP32, N columnas RSSI indexadas secuencialmente, seguidas por un conjunto de estadisticos resumen: media, desvio estandar, minimo, maximo, mediana, promedio de canal y conteo de muestras. El orden de las columnas no es decorativo. Debe mantenerse estable entre entrenamiento e inferencia para que el modelo reciba el vector correcto.",
            "La corrida principal genero un dataset de dieciocho filas y cincuenta y ocho columnas. Las siete primeras corresponden a identificacion y coordenadas. Las cincuenta y una restantes representan caracteristicas derivadas de las tres antenas ESP32. Esta dimensionalidad es suficientemente reducida para trabajar con modelos tabulares convencionales y, al mismo tiempo, suficientemente rica para capturar diferencias espaciales entre puntos del ambiente.",
        ],
    )
    add_para(doc, "Tabla 16.1 - Resumen dimensional del dataset de la corrida principal.")
    add_table(
        doc,
        ["Concepto", "Valor"],
        [
            ["Filas", str(run_1b["dataset_rows"])],
            ["Columnas", str(run_1b["dataset_cols"])],
            ["Puntos completos", str(run_1b["points_complete"])],
            ["Muestras por antena ESP32", str(run_1b["samples_per_anchor"])],
            ["Columnas de identificacion y coordenadas", "7"],
            ["Columnas de caracteristicas", str(run_1b["dataset_cols"] - 7)],
        ],
    )
    dataset_columns = list(run_1b["dataset_df"].columns) if not run_1b["dataset_df"].empty else []
    if dataset_columns:
        for idx in range(0, len(dataset_columns), 14):
            chunk = dataset_columns[idx : idx + 14]
            add_para(doc, f"Tabla 16.{2 + idx // 14} - Columnas del dataset ({idx + 1} a {idx + len(chunk)}).")
            add_table(
                doc,
                ["Indice", "Nombre de columna", "Tipo funcional"],
                [
                    [
                        str(idx + offset + 1),
                        col,
                        (
                            "Identificacion"
                            if col in {"session_id", "campaign_id", "point_id"}
                            else "Coordenada"
                            if col in {"x_m", "y_m", "z_m"}
                            else "Control"
                            if col == "samples_per_anchor"
                            else "Caracteristica de antena ESP32"
                        ),
                    ]
                    for offset, col in enumerate(chunk)
                ],
            )
    add_heading(doc, "16.2 Construccion de caracteristicas", 3)
    add_paragraphs(
        doc,
        [
            "La construccion de caracteristicas se realiza a partir de samples.csv, donde ya se encuentran unicamente las observaciones pertenecientes a la MAC objetivo y al punto activo. Para cada antena ESP32 se ordenan las muestras y se toman las N previstas por configuracion. Luego se generan estadisticos basicos sobre ese subconjunto: media, desvio, minimo, maximo y mediana. Esta estrategia busca conservar tanto la secuencia disponible como una descripcion agregada de la distribucion local de RSSI.",
            "Una consecuencia metodologica importante es que entrenamiento e inferencia deben compartir exactamente la misma representacion. Si en entrenamiento se utilizaron diez RSSI crudos por antena y en inferencia se intentara usar solo promedios, el modelo recibiria un vector de forma incompatible. Por ello, cualquier cambio futuro en el esquema de features exige reconstruir el dataset y reentrenar los modelos.",
            "La estructura elegida no agota las alternativas posibles. Podrian evaluarse representaciones basadas en mediana e intervalo intercuartil, ventanas temporales o estadisticos mas robustos. No obstante, para la primera etapa se privilegio una representacion simple, verificable y alineada con el modo en que se realizo la captura.",
        ],
    )
    add_heading(doc, "16.3 Implicancias metodologicas", 3)
    add_paragraphs(
        doc,
        [
            "La eleccion de un numero fijo de muestras por antena ESP32 tiene implicancias directas sobre el tiempo de captura por punto. Una configuracion mas exigente reduce la variabilidad estadistica de la estimacion local, pero incrementa la duracion total de la campana. Las dos corridas reales con tres antenas ESP32 muestran precisamente este compromiso: la corrida piloto con veinte muestras por antena completo menos puntos y requirio tiempos mucho mas altos; la corrida principal con diez muestras por antena logro una grilla mucho mas extensa.",
            "Tambien debe considerarse que una representacion rica en columnas aumenta el riesgo de sobreajuste cuando el numero de puntos es reducido. En el caso presente, el dataset principal posee dieciocho puntos y cincuenta y ocho columnas. Esta relacion entre filas y variables obliga a utilizar validacion cuidadosa y a interpretar las metricas con prudencia. El objetivo de la primera etapa no es demostrar una solucion definitiva, sino caracterizar un comportamiento experimental y producir una base reproducible para la segunda etapa con cuatro antenas ESP32.",
        ],
    )
    add_picture_with_caption(
        doc,
        figs["pipeline_flow"],
        "Figura 16.1. Pipeline de construccion del dataset y entrenamiento a partir de los archivos persistidos por campana.",
        width_inches=6.4,
    )
    add_heading(doc, "16.4 Semantica de los bloques de caracteristicas", 3)
    add_paragraphs(
        doc,
        [
            "Cada bloque de caracteristicas por antena ESP32 combina dos tipos de informacion. Por un lado, conserva diez RSSI individuales, lo que permite retener cierta variabilidad local de la medicion. Por otro, agrega estadisticos resumen que sintetizan tendencia central, dispersion y extremos. Esta dualidad es metodologicamente relevante porque evita depender de una unica representacion del fenomeno radioelectrico.",
            "En terminos practicos, columnas como media, mediana y maximo informan sobre el nivel caracteristico de senal y sobre eventos de mejor recepcion. El desvio estandar y el rango minimo-maximo sugieren cuan estable fue la captura en el punto. El conteo de muestras confirma que el bloque fue completado segun lo previsto. Aunque estos estadisticos son simples, su interpretacion conjunta ofrece una descripcion bastante rica para un modelo tabular de baja complejidad.",
            "Esta lectura semantica de las features es importante para la siguiente etapa del proyecto. Si la prueba con cuatro antenas ESP32 mostrara mejoras sustanciales, sera posible indagar no solo si mejoro el error final, sino que tipos de variables se volvieron mas informativas: las asociadas al nivel medio de senal, las vinculadas a dispersion o las que reflejan diferencias sistematicas entre receptores.",
        ],
    )

    add_heading(doc, "17. Entrenamiento e inferencia", 2)
    add_paragraphs(
        doc,
        [
            "El entrenamiento supervisado se apoya en el dataset consolidado por punto y persigue la estimacion de coordenadas bidimensionales a partir de las caracteristicas RSSI construidas para cada antena ESP32. La salida del modelo es, por lo tanto, un par ordenado (x, y). No se estima de manera separada la altura z, ya que la primera etapa del proyecto se concentra en posicionamiento sobre el plano del ambiente.",
            "El proceso implementado evalua varios modelos de regresion tabular. Esta estrategia es conveniente porque no existe, a priori, un algoritmo universalmente superior para datasets pequenos, ruidosos y con posible no linealidad. El proyecto adopta entonces un criterio comparativo: se entrenan distintos modelos, se calculan metricas comunes y se selecciona el que ofrece mejor equilibrio entre error medio y percentiles altos del error euclidiano.",
        ],
    )
    add_heading(doc, "17.1 Modelos evaluados y criterio de seleccion", 3)
    add_paragraphs(
        doc,
        [
            "En la corrida principal se evaluaron, entre otros, Random Forest, Extra Trees, K-Nearest Neighbors, Gradient Boosting Regressor y Support Vector Regression con kernel radial. Cada uno de estos modelos representa una hipotesis diferente respecto de la relacion entre las variables RSSI y la posicion espacial. Algunos se apoyan en ensambles de arboles, otros en proximidad entre puntos o en superficies no lineales de regresion.",
            "El criterio de seleccion no se limita al error medio absoluto. Tambien se observan el percentil 50, el percentil 90, el percentil 95 y la raiz del error cuadratico medio. Esta multiplicidad de metricas es importante porque un modelo puede tener buen error medio y, sin embargo, producir fallas grandes en una fraccion reducida pero relevante de los casos. Para una aplicacion de localizacion, la cola alta del error resulta especialmente significativa.",
            "La validacion se implementa con GroupKFold por punto, de modo que la particion de entrenamiento y evaluacion respete la unidad experimental del punto de referencia. Esta decision evita que muestras derivadas del mismo punto aparezcan simultaneamente en entrenamiento y prueba, lo cual generaria una estimacion artificialmente optimista del rendimiento.",
        ],
    )
    add_heading(doc, "17.2 Inferencia en linea", 3)
    add_paragraphs(
        doc,
        [
            "La inferencia en linea reutiliza el modelo ya entrenado para estimar en tiempo real la posicion probable del emisor objetivo y, opcionalmente, de otros dispositivos detectados por las antenas ESP32. Para ello se toman paquetes recientes almacenados en RAW, se agrupan por direccion MAC y se intenta reconstruir un vector de caracteristicas compatible con el formato del dataset de entrenamiento. La posicion estimada no surge, por tanto, de un unico RSSI instantaneo, sino de un bloque reciente de observaciones por antena.",
            "La interfaz permite definir una ventana de deteccion, una memoria temporal por antena, un minimo de paquetes por dispositivo, un minimo de antenas observadoras y umbrales adicionales de calidad. Estos filtros son necesarios porque el trafico del ambiente puede incluir dispositivos ajenos al recinto o equipos observados con muy pocas muestras. El objetivo de la inferencia en linea es ofrecer una estimacion util, no convertir cualquier paquete aislado en un punto espacial aparentemente significativo.",
            "La posibilidad de proyectar dispositivos externos sobre el plano del ambiente constituye una limitacion conocida del enfoque. El modelo siempre entrega una coordenada dentro del dominio aprendido, aun cuando el dispositivo real se encuentre fuera del recinto o se trate de un emisor no contemplado en la etapa de entrenamiento. Por esta razon, la interpretacion de la inferencia en linea debe combinar criterio experimental y filtros operativos.",
        ],
    )
    add_heading(doc, "17.3 Mapa de calor y estabilizacion temporal", 3)
    add_paragraphs(
        doc,
        [
            "El mapa de calor representa la densidad temporal de predicciones acumuladas en la interfaz de operacion. No debe interpretarse como un campo fisico medido directamente, sino como una visualizacion de la persistencia espacial de las estimaciones. Si un dispositivo aparece repetidamente en una zona, el mapa de calor intensifica ese sector. Si se reinicia el historial, la visualizacion vuelve a reflejar solo la informacion reciente.",
            "Esta estrategia tiene utilidad operativa porque suaviza la lectura visual cuando las predicciones individuales fluctuan entre actualizaciones sucesivas. No obstante, tambien puede inducir interpretaciones erroneas si se acumulan dispositivos espurios o si se confunden estimaciones antiguas con presencia actual. Por ello, la interfaz incluye controles para limitar la ventana temporal, resetear el historial y ajustar la exigencia de muestras por antena ESP32.",
        ],
    )
    add_picture_with_caption(
        doc,
        figs["inference_flow"],
        "Figura 17.1. Flujo de inferencia en linea, desde los paquetes recientes hasta la proyeccion sobre el mapa de calor y el plano CAD.",
        width_inches=6.4,
    )
    if (EXAMPLE_FIGS_DIR / "fig_20_heatmap_ejemplo.png").exists():
        add_picture_with_caption(
            doc,
            EXAMPLE_FIGS_DIR / "fig_20_heatmap_ejemplo.png",
            "Figura 17.2. Ejemplo conceptual de mapa de calor para la representacion espacial de predicciones acumuladas.",
            width_inches=6.3,
        )
    add_heading(doc, "17.4 Lectura tecnica de las metricas de entrenamiento", 3)
    add_paragraphs(
        doc,
        [
            "En problemas de localizacion indoor, ninguna metrica aislada resume por completo el comportamiento del sistema. El MAE ofrece una idea del error tipico, pero no informa adecuadamente sobre colas altas de error. El P50 aproxima la experiencia mediana del sistema. Los percentiles 90 y 95 permiten dimensionar escenarios desfavorables. El RMSE penaliza mas los errores grandes y ayuda a identificar modelos con fallas ocasionales severas.",
            "Por esta razon, el informe evita declarar un ganador absoluto en terminos simplistas. El mejor modelo no es necesariamente el que reduce un indicador y empeora drasticamente los demas, sino aquel que mantiene un equilibrio aceptable entre error central, cola alta y estabilidad. Esta forma de leer las metricas es especialmente importante cuando se proyecta la evolucion del sistema hacia la prueba con cuatro antenas ESP32 y multiples campanas.",
        ],
    )

    add_heading(doc, "18. Interfaz de operacion", 2)
    add_paragraphs(
        doc,
        [
            "La interfaz de operacion implementada en Streamlit cumple una funcion central durante las campanas. No se limita a presentar resultados finales, sino que articula el flujo de trabajo completo del operador: configuracion inicial de la campana, control del estado de las antenas ESP32, carga manual de puntos, cierre de puntos, construccion de dataset, entrenamiento y supervision de la inferencia en linea. Su existencia reduce la necesidad de interactuar manualmente con archivos o endpoints durante el trabajo de campo.",
            "La disposicion de la interfaz responde a la secuencia real del experimento. Primero se declara la campana y la geometria del ambiente. Luego se observa el estado de cada antena ESP32. A continuacion se opera punto por punto. Una vez completada la captura, se construye el dataset y se entrena el modelo. Finalmente se utiliza la seccion de inferencia para validar el comportamiento del sistema en tiempo real. Esta organizacion coincide con la logica de operacion efectivamente empleada durante las corridas realizadas.",
        ],
    )
    add_heading(doc, "18.1 Configuracion de la campana", 3)
    add_paragraphs(
        doc,
        [
            "La primera seccion de la interfaz permite declarar session_id, campaign_id, MAC objetivo, cantidad de antenas, cantidad de campanas planificadas, dimensiones del ambiente y altura por defecto de las antenas ESP32. Esta informacion se transmite al servidor y queda persistida en experiment.json. La operacion explicita de esta pantalla aporta dos beneficios: obliga a documentar el contexto experimental y reduce el riesgo de inconsistencias entre corridas.",
            "Tambien se cargan o editan manualmente las coordenadas de las antenas ESP32. Este punto es critico porque la geometria forma parte del problema de aprendizaje. Si la ubicacion declarada no coincide con la ubicacion real, el dataset queda contaminado desde su origen. Por esta razon, el informe insiste en considerar la interfaz como parte activa de la calidad metrologica del sistema.",
        ],
    )
    add_heading(doc, "18.2 Captura punto por punto", 3)
    add_paragraphs(
        doc,
        [
            "La captura punto por punto se opera desde una seccion especifica donde el usuario ingresa point_id, coordenadas x, y, z y habilita el inicio del punto. A partir de ese momento, el servidor contabiliza muestras validas de la MAC objetivo recibidas desde cada antena ESP32. La interfaz informa el progreso y permite cerrar el punto una vez alcanzado el numero requerido de observaciones.",
            "Este flujo tiene una consecuencia metodologica importante: el etiquetado espacial no se realiza automaticamente a posteriori, sino durante la campana. El operador conoce en cada instante cual es el punto activo y puede decidir si la captura es satisfactoria o si debe prolongarse. Esta supervisacion humana agrega trabajo operativo, pero mejora la trazabilidad y disminuye la probabilidad de mezclar posiciones distintas dentro de una misma fila del dataset.",
        ],
    )
    add_heading(doc, "18.3 Construccion de dataset, entrenamiento e inferencia", 3)
    add_paragraphs(
        doc,
        [
            "Una vez finalizada la campana, la misma interfaz expone botones para construir el dataset de campana, entrenar el modelo de campana, construir el dataset maestro de la sesion y entrenar el modelo maestro. Esta integracion reduce la fragmentacion del flujo de trabajo y evita que el operador deba cambiar constantemente entre terminales, scripts y exploradores de archivos.",
            "La seccion de inferencia en linea muestra, a su vez, indicadores sinteticos como cantidad de dispositivos detectados y coordenadas estimadas del emisor objetivo. Debajo se despliega una tabla de predicciones por direccion MAC y un plano con el mapa de calor correspondiente. Adicionalmente se ofrece una vista sobre plano CAD, pensada para cargar una imagen 2D exportada del ambiente real y superponer sobre ella las predicciones.",
        ],
    )
    add_heading(doc, "18.4 Recomendaciones de uso operativo", 3)
    add_paragraphs(
        doc,
        [
            "La experiencia de campo sugiere operar la interfaz con una disciplina clara: verificar siempre el estado de las antenas ESP32 antes de abrir un punto, confirmar la MAC objetivo, revisar que el contador de muestras avance de manera pareja y no cerrar el punto hasta que el servidor muestre el cupo esperado. Estas acciones, aunque simples, reducen significativamente la probabilidad de contaminar el dataset con mediciones incompletas o mal etiquetadas.",
            "Tambien se recomienda registrar en paralelo observaciones de campo no estructuradas, tales como demoras inusuales, reinicios de placas, cambios del ambiente o dificultades de conexion. Si bien estos eventos no siempre quedan reflejados en los archivos CSV, pueden ser decisivos para interpretar despues una desviacion del modelo o un incremento puntual del tiempo de captura en determinados puntos.",
        ],
    )
    add_photo_placeholder(doc, "Figura 18.1", "Captura de la interfaz de operacion durante la definicion de una campana y la carga de antenas ESP32.")
    add_photo_placeholder(doc, "Figura 18.2", "Captura de la interfaz durante la medicion punto por punto, con el contador de muestras visibles por antena ESP32.")
    add_photo_placeholder(doc, "Figura 18.3", "Captura de la interfaz en la etapa de entrenamiento y comparacion de modelos.")
    add_photo_placeholder(doc, "Figura 18.4", "Captura de la inferencia en linea sobre el plano abstracto del ambiente.")
    add_photo_placeholder(doc, "Figura 18.5", "Captura de la inferencia en linea superpuesta a un plano CAD del ambiente.")

    add_heading(doc, "19. Robustez, trazabilidad y consideraciones operativas", 2)
    add_paragraphs(
        doc,
        [
            "Las pruebas realizadas evidenciaron que la robustez del sistema depende no solo del algoritmo de entrenamiento, sino de la estabilidad integral del conjunto hardware-software. La tendencia de algunas antenas ESP32 a requerir reinicio manual tras varios ciclos, la sensibilidad a la reconexion Wi-Fi y la variabilidad del volumen capturado por nodo demostraron que la plataforma experimental debe analizarse como un sistema distribuido sujeto a fallas parciales.",
            "Frente a esta realidad, se incorporaron mecanismos de trazabilidad y recuperacion. El almacenamiento RAW en JSONL permite inspeccionar que observo cada antena ESP32, incluso si la muestra luego no fue aceptada para entrenamiento. training_state.json conserva el estado operativo, points.csv registra el cierre de los puntos y el reinicio automatico del firmware reduce el tiempo de inactividad ante fallas repetidas de reconexion. Estas medidas no eliminan toda incertidumbre, pero convierten al sistema en una plataforma auditable.",
            "Desde el punto de vista de control de calidad, la posibilidad de comparar el numero esperado de muestras con el numero efectivamente aceptado, revisar tiempos de captura por punto y verificar la existencia de dataset y metricas representa una ventaja considerable. El proyecto no queda reducido a una visualizacion puntual en pantalla, sino que deja una huella documental completa de cada campana.",
            "La experiencia adquirida durante las corridas con tres antenas ESP32 sugiere que la segunda etapa con cuatro receptores debe ejecutarse con protocolos de chequeo previos mas rigurosos, incluyendo verificaciones de conectividad, de alimentacion y de sincronizacion operativa entre nodos. Lejos de ser un detalle administrativo, esta recomendacion constituye una consecuencia directa de la evidencia obtenida en campo.",
        ],
    )
    add_para(doc, "Tabla 19.1 - Problemas operativos observados y acciones de mitigacion.")
    add_table(
        doc,
        ["Problema observado", "Manifestacion", "Mitigacion incorporada", "Estado"],
        [
            ["Reconexiones fallidas", "Antenas ESP32 que dejan de enviar lotes luego de varios ciclos", "Reintentos y reinicio automatico despues de tres fallas", "Implementado"],
            ["Capturas desbalanceadas entre nodos", "Una antena ESP32 completa el cupo antes que otras", "Ajuste del firmware y revision de ubicacion fisica", "Parcial"],
            ["Ruido de dispositivos externos", "Predicciones de dispositivos fuera del ambiente", "Filtros de RSSI, ventana temporal y minimo por antena", "Implementado"],
            ["Tiempo alto de captura", "Puntos con demora excesiva al exigir muchas muestras", "Reduccion de muestras por antena de 20 a 10 en la corrida principal", "Implementado"],
            ["Asimetria geometrica", "Mayor error cualitativo en uno de los ejes", "Planificacion de segunda etapa con cuatro antenas ESP32", "Pendiente"],
        ],
    )
    add_heading(doc, "19.1 Lecciones aprendidas de robustez", 3)
    add_paragraphs(
        doc,
        [
            "La primera leccion aprendida es que la estabilidad de las antenas ESP32 no puede asumirse como una condicion dada. Debe instrumentarse, observarse y, cuando sea necesario, corregirse por firmware o por mejoras de alimentacion. La segunda es que los tiempos de captura son una variable de diseno tan importante como el error del modelo. Un sistema potencialmente preciso pero impracticable en tiempo de medicion no resulta metodologicamente conveniente.",
            "La tercera leccion es que la persistencia detallada de datos no es un lujo, sino una necesidad. Sin RAW, sin samples.csv y sin points.csv seria muy dificil reconstruir que ocurrio realmente en cada campana. El valor de la trazabilidad se hizo evidente precisamente cuando aparecieron fallas o comportamientos inesperados, ya que la informacion persistida permitio distinguir problemas de captura, de filtrado o de entrenamiento.",
        ],
    )
    add_para(
        doc,
        "Con lo expuesto se cierra la Parte II. En ella se establecen las bases de hardware, software, arquitectura y metodologia de operacion que permiten interpretar correctamente los resultados experimentales de la Parte III. Sin esta descripcion detallada, las metricas del modelo carecerian del contexto necesario para ser evaluadas en terminos ingenieriles.",
    )


def add_point_tables(doc: Document, title_prefix: str, points_df: pd.DataFrame, chunk_size: int = 9) -> None:
    table_idx = 1
    for start in range(0, len(points_df), chunk_size):
        chunk = points_df.iloc[start : start + chunk_size]
        add_para(doc, f"{title_prefix} {table_idx}.")
        add_table(
            doc,
            ["Punto", "x [m]", "y [m]", "z [m]", "Captura [s]", "Cierre [s]", "Estado"],
            [
                [
                    row["point_id"],
                    f"{float(row['x_m']):.2f}",
                    f"{float(row['y_m']):.2f}",
                    f"{float(row['z_m']):.2f}",
                    format_seconds(row.get("capture_s")),
                    format_seconds(row.get("close_delay_s")),
                    row.get("status", "N/D"),
                ]
                for _, row in chunk.iterrows()
            ],
        )
        table_idx += 1


def add_raw_table(doc: Document, title: str, raw_sizes: dict[str, int]) -> None:
    add_para(doc, title)
    add_table(
        doc,
        ["Antena ESP32", "Tamano [bytes]", "Tamano [MiB]"],
        [
            [antenna_id, str(size), format_mebibytes(size)]
            for antenna_id, size in raw_sizes.items()
        ],
    )


def add_sample_count_table(doc: Document, title: str, counts: dict[str, int]) -> None:
    add_para(doc, title)
    add_table(
        doc,
        ["Antena ESP32", "Muestras aceptadas"],
        [[antenna_id, str(count)] for antenna_id, count in counts.items()],
    )


def build_part_iii(doc: Document, run_1a: dict, run_1b: dict, figs: dict[str, Path]) -> None:
    doc.add_page_break()
    add_heading(doc, "PARTE III - RESULTADOS, ANALISIS, MEJORAS Y CONCLUSIONES", 1)
    add_paragraphs(
        doc,
        [
            "La Parte III se concentra en la evidencia empirica obtenida a partir de las corridas reales realizadas hasta el momento y en la planificacion detallada de la siguiente etapa experimental. Se incluyen tanto la corrida piloto parcial como la corrida principal completada con tres antenas ESP32, debido a que ambas aportan informacion distinta y complementaria. La primera pone de manifiesto limitaciones operativas y costos temporales de un esquema mas exigente en cantidad de muestras. La segunda ofrece ya un dataset completo, metricas comparables entre modelos y observaciones de inferencia en linea.",
            "Debe destacarse que, al momento de generacion de este documento, no existen aun corridas reales completadas con cuatro antenas ESP32 en el repositorio. Por ese motivo, la segunda etapa experimental se documenta como diseno metodologico, con tablas, espacios y criterios de registro preparados para completarse una vez ejecutadas las nuevas campanas. Esta decision evita inventar resultados y mantiene el caracter trazable del informe.",
        ],
    )

    add_heading(doc, "20. Programa experimental general", 2)
    add_paragraphs(
        doc,
        [
            "El programa experimental se estructura en dos etapas principales. La primera etapa corresponde a la validacion con tres antenas ESP32. Dentro de ella se distinguen una corrida piloto parcial, de naturaleza exploratoria, y una corrida principal suficientemente extensa como para permitir entrenamiento y evaluacion de modelos. La segunda etapa, planificada pero aun no ejecutada al momento de esta version del informe, incorpora cuatro antenas ESP32 y tres campanas diferenciadas para estudiar de manera mas rigurosa la influencia de la geometria y de modificaciones del ambiente.",
            "La decision de organizar las pruebas en etapas responde a un criterio metodologico prudente. Antes de invertir esfuerzo en una configuracion de cuatro receptores y multiples campanas, era necesario verificar que el flujo extremo a extremo funcionara: captura, etiquetado, consolidacion, entrenamiento e inferencia. Las corridas con tres antenas ESP32 cumplieron justamente esa funcion y permitieron identificar tanto fortalezas como puntos criticos del sistema.",
            "Desde el punto de vista de la interpretacion de resultados, la primera etapa no debe juzgarse solo por las metricas numericas alcanzadas. Tambien interesa la informacion que brinda sobre tiempos de captura, volumen de datos RAW, estabilidad de las antenas ESP32 y sensibilidad a la geometria de despliegue. Esta informacion operativa sera crucial para optimizar la prueba con cuatro receptores.",
        ],
    )
    add_para(doc, "Tabla 20.1 - Estructura general del programa experimental.")
    add_table(
        doc,
        ["Etapa", "Identificador", "Configuracion", "Estado", "Objetivo principal"],
        [
            ["Prueba 1A", "demo_s1 / train_01", "3 antenas ESP32, 20 muestras por antena y punto", "Realizada", "Validar flujo end to end y detectar limites operativos iniciales."],
            ["Prueba 1B", "demo_s2 / train_02", "3 antenas ESP32, 10 muestras por antena y punto", "Realizada", "Completar grilla extensa y entrenar modelos comparables."],
            ["Prueba 2", "P2_C01_base_4A", "4 antenas ESP32, campana base", "Planificada", "Cuantificar el impacto de la geometria con cuatro esquinas."],
            ["Prueba 2", "P2_C02_layout_4A", "4 antenas ESP32, ambiente modificado", "Planificada", "Evaluar sensibilidad a cambios de mobiliario."],
            ["Prueba 2", "P2_C03_robustez_4A", "4 antenas ESP32, repeticion controlada", "Planificada", "Analizar repetibilidad y estabilidad temporal."],
        ],
    )
    add_picture_with_caption(
        doc,
        figs["run_compare"],
        "Figura 20.1. Comparacion global entre la corrida piloto parcial y la corrida principal con tres antenas ESP32.",
        width_inches=6.3,
    )

    add_heading(doc, "21. Resultados de la Prueba 1A: demo_s1 / train_01", 2)
    add_paragraphs(
        doc,
        [
            "La corrida demo_s1 / train_01 corresponde a la primera validacion end to end efectivamente operativa del sistema con tres antenas ESP32. Su importancia no radica en la cantidad de puntos completados, sino en haber demostrado que el firmware, el servidor, la interfaz de operacion y la logica de conteo por MAC objetivo podian funcionar de manera integrada. Esta corrida debe entenderse como una instancia piloto, orientada a desriesgar el sistema completo antes de una campana mas extensa.",
            "En esta prueba se exigieron veinte muestras por antena ESP32 y por punto. Tal configuracion representaba una apuesta por maximizar informacion local, pero implicaba tambien un costo temporal elevado. El resultado empirico confirmo esta intuicion: solo se completaron dos puntos y el tiempo de captura por punto fue muy superior al observado luego en la corrida principal con diez muestras por antena.",
            "Aun sin llegar a generar un dataset final ni metricas de modelos, la corrida piloto produjo evidencia valiosa. Se verifico la persistencia RAW, se confirmo el conteo correcto de muestras por antena ESP32, se observaron demoras significativas asociadas a la exigencia de captura y se documentaron las primeras dificultades de estabilidad de algunas placas. Por lo tanto, su ausencia de modelo entrenado debe leerse como resultado experimental, no como una mera omision de proceso.",
        ],
    )
    add_para(doc, "Tabla 21.1 - Configuracion real de la corrida piloto parcial con tres antenas ESP32.")
    add_table(
        doc,
        ["Parametro", "Valor"],
        [
            ["Sesion", run_1a["session_id"]],
            ["Campana", run_1a["campaign_id"]],
            ["MAC objetivo", run_1a["target_mac"]],
            ["Largo del ambiente [m]", f"{float(run_1a['environment']['length_m']):.2f}"],
            ["Ancho del ambiente [m]", f"{float(run_1a['environment']['width_m']):.2f}"],
            ["Alto del ambiente [m]", f"{float(run_1a['environment']['height_m']):.2f}"],
            ["Antenas ESP32", str(len(run_1a["antennas"]))],
            ["Muestras por antena y punto", str(run_1a["samples_per_anchor"])],
            ["Puntos completos", str(run_1a["points_complete"])],
            ["Muestras aceptadas", str(run_1a["samples_rows"])],
        ],
    )
    add_para(doc, "Tabla 21.2 - Coordenadas de las antenas ESP32 utilizadas en la corrida piloto.")
    add_table(
        doc,
        ["Antena ESP32", "x [m]", "y [m]", "z [m]"],
        [
            [a["anchor_id"], f"{float(a['x_m']):.2f}", f"{float(a['y_m']):.2f}", f"{float(a['z_m']):.2f}"]
            for a in run_1a["antennas"]
        ],
    )
    add_point_tables(doc, "Tabla 21.3 - Puntos medidos en la corrida piloto (bloque", run_1a["points_df"], chunk_size=6)
    add_sample_count_table(doc, "Tabla 21.4 - Muestras aceptadas por antena ESP32 en la corrida piloto.", run_1a["samples_per_antenna"])
    add_raw_table(doc, "Tabla 21.5 - Volumen de evidencia RAW en la corrida piloto.", run_1a["raw_sizes"])
    add_heading(doc, "21.1 Observaciones puntuales de la corrida piloto", 3)
    for _, row in run_1a["points_df"].iterrows():
        add_para(
            doc,
            f"El punto {row['point_id']} se ubico en ({float(row['x_m']):.2f}, {float(row['y_m']):.2f}, {float(row['z_m']):.2f}) m y requirio {format_seconds(row['capture_s'])} s para completar la captura. El tiempo adicional entre la captura completa y el cierre administrativo fue de {format_seconds(row['close_delay_s'])} s. Considerando que la corrida piloto exigia {run_1a['samples_per_anchor']} muestras por antena ESP32, este punto confirma que la metodologia era funcional pero costosa en tiempo. Su registro resulta importante porque aporta una medicion real de la penalizacion temporal asociada a una configuracion de captura mas intensiva.",
        )
    add_paragraphs(
        doc,
        [
            f"La corrida piloto completo {run_1a['points_complete']} puntos y acepto {run_1a['samples_rows']} muestras, exactamente {run_1a['expected_samples']} observaciones, lo que indica que para los puntos efectivamente cerrados se alcanzo el cupo previsto de {run_1a['samples_per_anchor']} muestras por antena ESP32. Este dato es importante porque demuestra que el flujo de captura y filtrado por MAC funciono correctamente a nivel funcional.",
            f"Sin embargo, el tiempo medio de captura fue de {format_seconds(run_1a['capture_mean_s'])} s por punto, con un minimo de {format_seconds(run_1a['capture_min_s'])} s y un maximo de {format_seconds(run_1a['capture_max_s'])} s. Estas cifras son demasiado elevadas para una grilla extensa de veinte puntos, ya que hubieran implicado una campana de varias horas aun sin contar pausas, reposicionamientos o incidencias operativas.",
            f"El punto mas costoso fue {run_1a['longest_point']['point_id']} con {format_seconds(run_1a['longest_point']['capture_s'])} s de captura, mientras que el punto mas rapido fue {run_1a['shortest_point']['point_id']} con {format_seconds(run_1a['shortest_point']['capture_s'])} s. Esta dispersion temporal ya anticipaba que el sistema necesitaba ajustes tanto en la cantidad de muestras requerida como en la robustez del firmware.",
            "La ausencia de dataset.csv y de metrics.csv en esta corrida debe registrarse expresamente. No se trata de archivos perdidos ni de un descuido posterior, sino de una consecuencia metodologica del caracter piloto de la prueba. La campana se interrumpio una vez comprobado el funcionamiento general del sistema y detectado que el costo temporal de veinte muestras por antena era poco compatible con una grilla amplia dentro del living.",
            "En sintesis, la Prueba 1A cumplio una funcion de validacion funcional y de exploracion de limites. Aporto evidencia de captura real, puso en evidencia la carga temporal de una configuracion muy exigente y justifico el rediseño operativo que dio origen a la corrida principal con diez muestras por antena ESP32.",
        ],
    )

    add_heading(doc, "22. Resultados de la Prueba 1B: demo_s2 / train_02", 2)
    add_paragraphs(
        doc,
        [
            "La corrida demo_s2 / train_02 constituye la primera campana extensa y completa del sistema con tres antenas ESP32. A diferencia de la corrida piloto, esta version logro cubrir una grilla mucho mas amplia del living y produjo un dataset supervisado apto para entrenamiento y comparacion de modelos. En ese sentido, representa la primera base cuantitativa solida del proyecto.",
            "La principal decision metodologica que diferencia a esta corrida de la piloto es la reduccion de veinte a diez muestras por antena y por punto. Esta modificacion no fue arbitraria. Se adopto a partir de la evidencia reunida en la Prueba 1A y tuvo como objetivo disminuir el tiempo de captura sin perder la estructura de features necesaria para el entrenamiento. El resultado empirico muestra que la medida fue efectiva: se lograron dieciocho puntos completos y se generaron archivos de entrenamiento y metricas comparables.",
            "La Prueba 1B debe interpretarse como una corrida principal de caracterizacion, no como una validacion definitiva del sistema. Aporta resultados cuantitativos relevantes, revela limitaciones de geometria y estabilidad, y habilita una primera discusion seria sobre el potencial y los limites de la arquitectura con tres antenas ESP32.",
        ],
    )
    add_para(doc, "Tabla 22.1 - Configuracion real de la corrida principal con tres antenas ESP32.")
    add_table(
        doc,
        ["Parametro", "Valor"],
        [
            ["Sesion", run_1b["session_id"]],
            ["Campana", run_1b["campaign_id"]],
            ["MAC objetivo", run_1b["target_mac"]],
            ["Largo del ambiente [m]", f"{float(run_1b['environment']['length_m']):.2f}"],
            ["Ancho del ambiente [m]", f"{float(run_1b['environment']['width_m']):.2f}"],
            ["Alto del ambiente [m]", f"{float(run_1b['environment']['height_m']):.2f}"],
            ["Antenas ESP32", str(len(run_1b["antennas"]))],
            ["Muestras por antena y punto", str(run_1b["samples_per_anchor"])],
            ["Puntos completos", str(run_1b["points_complete"])],
            ["Muestras aceptadas", str(run_1b["samples_rows"])],
            ["Dataset", f"{run_1b['dataset_rows']} filas x {run_1b['dataset_cols']} columnas"],
        ],
    )
    add_para(doc, "Tabla 22.2 - Coordenadas de las antenas ESP32 utilizadas en la corrida principal.")
    add_table(
        doc,
        ["Antena ESP32", "x [m]", "y [m]", "z [m]"],
        [
            [a["anchor_id"], f"{float(a['x_m']):.2f}", f"{float(a['y_m']):.2f}", f"{float(a['z_m']):.2f}"]
            for a in run_1b["antennas"]
        ],
    )
    add_point_tables(doc, "Tabla 22.3 - Puntos medidos en la corrida principal (bloque", run_1b["points_df"], chunk_size=9)
    add_para(doc, "Tabla 22.5 - Resumen temporal de la corrida principal.")
    add_table(
        doc,
        ["Indicador", "Valor"],
        [
            ["Tiempo medio de captura por punto [s]", format_seconds(run_1b["capture_mean_s"])],
            ["Mediana de captura por punto [s]", format_seconds(run_1b["capture_median_s"])],
            ["Tiempo minimo de captura [s]", format_seconds(run_1b["capture_min_s"])],
            ["Tiempo maximo de captura [s]", format_seconds(run_1b["capture_max_s"])],
            ["Tiempo medio entre captura completa y cierre [s]", format_seconds(run_1b["close_delay_mean_s"])],
            ["Punto mas rapido", f"{run_1b['shortest_point']['point_id']} ({format_seconds(run_1b['shortest_point']['capture_s'])} s)"],
            ["Punto mas lento", f"{run_1b['longest_point']['point_id']} ({format_seconds(run_1b['longest_point']['capture_s'])} s)"],
        ],
    )
    add_sample_count_table(doc, "Tabla 22.6 - Muestras aceptadas por antena ESP32 en la corrida principal.", run_1b["samples_per_antenna"])
    add_raw_table(doc, "Tabla 22.7 - Volumen de evidencia RAW por antena ESP32 en la corrida principal.", run_1b["raw_sizes"])
    add_para(doc, "Tabla 22.8 - Rendimiento de modelos entrenados con el dataset de la corrida principal.")
    add_table(
        doc,
        ["Modelo", "MAE [m]", "P50 [m]", "P90 [m]", "P95 [m]", "RMSE [m]", "Evaluacion"],
        [
            [
                row["model"],
                f"{float(row['mae_eucl']):.4f}",
                f"{float(row['p50']):.4f}",
                f"{float(row['p90']):.4f}",
                f"{float(row['p95']):.4f}",
                f"{float(row['rmse_eucl']):.4f}",
                row["eval_mode"],
            ]
            for _, row in run_1b["metrics_df"].sort_values("mae_eucl").iterrows()
        ],
    )
    add_picture_with_caption(
        doc,
        figs["run_1b_capture_times"],
        "Figura 22.1. Tiempo de captura por punto en la corrida principal con tres antenas ESP32.",
        width_inches=6.5,
    )
    add_picture_with_caption(
        doc,
        figs["run_1b_metrics"],
        "Figura 22.2. Comparacion de modelos entrenados con el dataset de la corrida principal.",
        width_inches=6.2,
    )
    add_picture_with_caption(
        doc,
        figs["run_1b_raw_sizes"],
        "Figura 22.3. Volumen RAW por antena ESP32 en la corrida principal.",
        width_inches=5.8,
    )
    if (EXAMPLE_FIGS_DIR / "fig_20_1_cdf_ejemplo.png").exists():
        add_picture_with_caption(
            doc,
            EXAMPLE_FIGS_DIR / "fig_20_1_cdf_ejemplo.png",
            "Figura 22.4. Figura conceptual de referencia para discutir distribuciones acumuladas de error.",
            width_inches=6.0,
        )
    if (EXAMPLE_FIGS_DIR / "fig_20_scatter_ejemplo.png").exists():
        add_picture_with_caption(
            doc,
            EXAMPLE_FIGS_DIR / "fig_20_scatter_ejemplo.png",
            "Figura 22.5. Figura conceptual de referencia para discutir dispersion espacial de predicciones.",
            width_inches=6.1,
        )
    add_paragraphs(
        doc,
        [
            f"La corrida principal completo {run_1b['points_complete']} puntos, lo que representa una cobertura sustancialmente mayor que la corrida piloto. Se aceptaron {run_1b['samples_rows']} muestras, exactamente {run_1b['expected_samples']} observaciones, por lo que el cupo previsto de {run_1b['samples_per_anchor']} muestras por antena ESP32 y por punto se cumplio para toda la grilla efectivamente relevada. Este dato confirma la consistencia interna de la campana y la correcta operacion del pipeline de etiquetado.",
            f"El tiempo medio de captura fue de {format_seconds(run_1b['capture_mean_s'])} s por punto, con una mediana de {format_seconds(run_1b['capture_median_s'])} s. Aunque estas cifras siguen siendo elevadas para una operacion plenamente automatizada, representan una reduccion muy marcada respecto de la corrida piloto. La mejora se explica principalmente por la menor exigencia de muestras por antena ESP32, pero tambien por los ajustes introducidos en el flujo operativo y en el firmware.",
            f"El punto mas lento fue {run_1b['longest_point']['point_id']} con {format_seconds(run_1b['longest_point']['capture_s'])} s, mientras que el mas rapido fue {run_1b['shortest_point']['point_id']} con {format_seconds(run_1b['shortest_point']['capture_s'])} s. La amplitud entre ambos extremos indica que la captura no depende unicamente del numero de muestras requerido. Tambien influyen la posicion relativa del emisor, la geometria respecto de las antenas ESP32, la congestion radioelectrica y el comportamiento de reconexion de los nodos.",
            f"En cuanto al dataset, la campana produjo {run_1b['dataset_rows']} filas y {run_1b['dataset_cols']} columnas. Esta dimension confirma que la estructura elegida pudo generarse correctamente y que el pipeline completo de consolidacion funciono tal como fue disenado. En otras palabras, la Prueba 1B no solo produjo archivos de captura, sino un conjunto de entrenamiento utilizable y reproducible.",
            f"El mejor modelo por error medio absoluto fue {run_1b['best_mae']['model']} con un MAE de {float(run_1b['best_mae']['mae_eucl']):.4f} m. El mejor modelo por percentil 95 tambien fue {run_1b['best_p95']['model']} con un P95 de {float(run_1b['best_p95']['p95']):.4f} m. Este resultado sugiere que, en la configuracion actual, los ensambles de arboles ofrecen una relacion favorable entre capacidad de ajuste y robustez frente a un dataset pequeno.",
            "Aun asi, las metricas deben interpretarse con cautela. Un MAE del orden de 1,26 m a 1,35 m puede considerarse prometedor para una primera campana con tres antenas ESP32 de bajo costo, pero todavia deja margen significativo de mejora para aplicaciones que requieran mayor precision. La evidencia cualitativa observada durante la inferencia en linea, particularmente cierta desviacion en uno de los ejes del plano, refuerza esta lectura prudente.",
            "Otro dato relevante es el volumen RAW por antena ESP32. Los archivos muestran que la cantidad de informacion capturada y persistida no fue uniforme entre nodos. Esta asimetria es coherente con la experiencia operativa recogida en campo: algunas placas tendieron a comportarse de forma mas estable o a observar mayor volumen de trafico. La interpretacion de este hallazgo es doble. Por un lado, demuestra que la plataforma produce trazas suficientes para auditoria. Por otro, confirma que la homogeneidad entre receptores no puede darse por garantizada.",
        ],
    )
    add_heading(doc, "22.1 Lectura detallada por punto de referencia", 3)
    for _, row in run_1b["points_df"].iterrows():
        add_para(
            doc,
            f"Para el punto {row['point_id']}, ubicado en ({float(row['x_m']):.2f}, {float(row['y_m']):.2f}, {float(row['z_m']):.2f}) m, el tiempo de captura fue de {format_seconds(row['capture_s'])} s y el tiempo de cierre posterior fue de {format_seconds(row['close_delay_s'])} s. Este registro puntual permite observar que la productividad de la campana no fue uniforme sobre todo el plano. La variacion entre puntos sugiere que la dificultad de completar el cupo de {run_1b['samples_per_anchor']} muestras por antena ESP32 depende tanto de la geometria relativa respecto de las antenas como del comportamiento temporal de los nodos durante la reconexion y el envio de lotes.",
        )
    add_heading(doc, "22.2 Comentario por modelo evaluado", 3)
    for _, row in run_1b["metrics_df"].sort_values("mae_eucl").iterrows():
        add_para(
            doc,
            f"El modelo {row['model']} obtuvo un MAE de {float(row['mae_eucl']):.4f} m, un P50 de {float(row['p50']):.4f} m, un P90 de {float(row['p90']):.4f} m, un P95 de {float(row['p95']):.4f} m y un RMSE de {float(row['rmse_eucl']):.4f} m. Este perfil sugiere un compromiso especifico entre error tipico y errores de cola. Su inclusion en el estudio es relevante porque muestra que, aun con el mismo dataset, la eleccion del modelo modifica de forma apreciable la lectura final del problema. Comparar estas cifras permite comprender por que los ensambles de arboles emergen como candidatos particularmente favorables para la primera etapa del proyecto.",
        )
    add_heading(doc, "22.3 Comentario por antena ESP32", 3)
    for antenna_id, size in run_1b["raw_sizes"].items():
        count = run_1b["samples_per_antenna"].get(antenna_id, 0)
        add_para(
            doc,
            f"La antena ESP32 {antenna_id} genero {format_mebibytes(size)} MiB de evidencia RAW y contribuyo con {count} muestras aceptadas al dataset. La diferencia de volumen respecto de otras antenas revela que el comportamiento de captura no fue identico entre nodos. Este hallazgo debe ser considerado en la planificacion de la prueba con cuatro antenas ESP32, ya que la homogeneidad de receptores no puede asumirse sin medicion.",
        )

    add_heading(doc, "23. Sintesis comparativa de resultados reales disponibles", 2)
    add_paragraphs(
        doc,
        [
            "La comparacion entre demo_s1 / train_01 y demo_s2 / train_02 permite extraer una conclusion metodologica muy clara: exigir demasiadas muestras por antena ESP32 y por punto puede volver impracticable una grilla extensa, aun cuando el sistema funcione correctamente a nivel end to end. La reduccion de veinte a diez muestras por antena resulto determinante para ampliar la cobertura espacial y habilitar el entrenamiento de modelos.",
            "La corrida piloto y la corrida principal no deben verse como experiencias desconectadas. La segunda es consecuencia directa de los hallazgos de la primera. Sin la evidencia de tiempos excesivos y de costos operativos altos en la Prueba 1A, la decision de redisenar el protocolo de captura hubiera carecido de sustento empírico. Esta relacion entre corridas ilustra un proceso genuino de iteracion ingenieril apoyado en medicion real.",
        ],
    )
    add_para(doc, "Tabla 23.1 - Comparacion sintetica entre las dos corridas reales con tres antenas ESP32.")
    add_table(
        doc,
        ["Indicador", "Prueba 1A", "Prueba 1B"],
        [
            ["Muestras por antena y punto", str(run_1a["samples_per_anchor"]), str(run_1b["samples_per_anchor"])],
            ["Puntos completos", str(run_1a["points_complete"]), str(run_1b["points_complete"])],
            ["Muestras aceptadas", str(run_1a["samples_rows"]), str(run_1b["samples_rows"])],
            ["Tiempo medio por punto [s]", format_seconds(run_1a["capture_mean_s"]), format_seconds(run_1b["capture_mean_s"])],
            ["Dataset.csv", "No generado", f"{run_1b['dataset_rows']} x {run_1b['dataset_cols']}"],
            ["Metricas de modelos", "No disponibles", f"{len(run_1b['metrics_df'])} modelos comparados"],
        ],
    )
    add_paragraphs(
        doc,
        [
            "La evidencia comparativa tambien respalda la necesidad de avanzar hacia una geometria mas equilibrada con cuatro antenas ESP32. Las tres antenas utilizadas en la primera etapa permitieron validar el sistema y producir resultados medibles, pero dejaron visible una asimetria espacial que podria estar contribuyendo a errores mayores sobre uno de los ejes del plano. La ampliacion a cuatro esquinas apunta precisamente a reducir esta limitacion.",
            "En resumen, la primera etapa ya aporta resultados cuantitativos y cualitativos suficientes para justificar la segunda. No se trata de repetir la misma prueba por mera acumulacion de datos, sino de introducir una modificacion estructural significativa en la geometria de observacion y evaluar su impacto de manera controlada a traves de tres campanas planificadas.",
        ],
    )

    add_heading(doc, "24. Prueba 2 con 4 antenas ESP32", 2)
    add_paragraphs(
        doc,
        [
            "La segunda etapa experimental se disena con cuatro antenas ESP32 ubicadas en las cuatro esquinas del recinto. El objetivo tecnico central es mejorar la observabilidad espacial del problema, disminuir la asimetria geométrica y evaluar si la precision del sistema mejora de manera mensurable respecto de la configuracion de tres receptores. Esta etapa se considera una evolucion natural del prototipo, no un cambio radical de paradigma.",
            "Se planifican tres campanas diferenciadas. La primera mantendra un ambiente base y servira como referencia con cuatro receptores. La segunda introducira modificaciones controladas en el mobiliario para evaluar sensibilidad a cambios de propagacion. La tercera se orientara a repetibilidad y robustez, repitiendo el protocolo bajo condiciones similares para estimar variabilidad inter-campana. Todas las tablas necesarias para registrar resultados se dejan preparadas en esta version del informe.",
            "Es importante subrayar que las siguientes tablas y espacios no contienen valores medidos aun. Se presentan con fines de documentacion metodologica y para facilitar el registro sistematico una vez ejecutadas las campanas reales. De este modo, el informe queda listo para evolucionar sin romper su estructura.",
        ],
    )
    add_para(doc, "Tabla 24.1 - Geometria prevista para la prueba con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Antena ESP32", "x [m]", "y [m]", "z [m]", "Estado"],
        [
            ["A1", "A completar", "A completar", "2.00", "Planificada"],
            ["A2", "A completar", "A completar", "2.00", "Planificada"],
            ["A3", "A completar", "A completar", "2.00", "Planificada"],
            ["A4", "A completar", "A completar", "2.00", "Planificada"],
        ],
    )
    add_para(doc, "Tabla 24.2 - Campanas planificadas para la configuracion con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Campana", "Objetivo", "Resultados a registrar"],
        [[camp_id, desc, "Tiempo por punto, muestras aceptadas, dataset, metricas, observaciones"] for camp_id, desc in TRIAL_2_CAMPAIGNS],
    )
    add_para(doc, "Tabla 24.3 - Plantilla de resultados para la campana P2_C01_base_4A.")
    add_table(
        doc,
        ["Indicador", "Valor medido", "Observaciones"],
        [
            ["Puntos completos", "A completar", ""],
            ["Muestras aceptadas", "A completar", ""],
            ["Tiempo medio por punto [s]", "A completar", ""],
            ["Dataset.csv", "A completar", ""],
            ["Mejor MAE [m]", "A completar", ""],
            ["Mejor P95 [m]", "A completar", ""],
        ],
    )
    add_para(doc, "Tabla 24.4 - Plantilla de resultados para la campana P2_C02_layout_4A.")
    add_table(
        doc,
        ["Indicador", "Valor medido", "Observaciones"],
        [
            ["Puntos completos", "A completar", ""],
            ["Muestras aceptadas", "A completar", ""],
            ["Tiempo medio por punto [s]", "A completar", ""],
            ["Dataset.csv", "A completar", ""],
            ["Mejor MAE [m]", "A completar", ""],
            ["Mejor P95 [m]", "A completar", ""],
        ],
    )
    add_para(doc, "Tabla 24.5 - Plantilla de resultados para la campana P2_C03_robustez_4A.")
    add_table(
        doc,
        ["Indicador", "Valor medido", "Observaciones"],
        [
            ["Puntos completos", "A completar", ""],
            ["Muestras aceptadas", "A completar", ""],
            ["Tiempo medio por punto [s]", "A completar", ""],
            ["Dataset.csv", "A completar", ""],
            ["Mejor MAE [m]", "A completar", ""],
            ["Mejor P95 [m]", "A completar", ""],
        ],
    )
    add_photo_placeholder(doc, "Figura 24.1", "Fotografia general del montaje con cuatro antenas ESP32 en las cuatro esquinas del ambiente.")
    add_photo_placeholder(doc, "Figura 24.2", "Fotografias de cada una de las tres campanas con cuatro antenas ESP32.")
    add_photo_placeholder(doc, "Figura 24.3", "Capturas de pantalla de la interfaz durante la ejecucion de la prueba con cuatro antenas ESP32.")
    add_heading(doc, "24.1 Campana P2_C01_base_4A", 3)
    add_paragraphs(
        doc,
        [
            "La campana base con cuatro antenas ESP32 debe reproducir, en la medida de lo posible, las condiciones geometricas generales de la Prueba 1B, agregando un receptor en la esquina faltante. Su finalidad es construir una nueva referencia experimental donde el unico cambio estructural significativo sea la topologia de observacion. Esto permitira atribuir diferencias de rendimiento principalmente a la nueva geometria y no a variaciones metodologicas no controladas.",
            "Se recomienda mantener la misma grilla de puntos de referencia y la misma politica de diez muestras por antena ESP32 y por punto, salvo que las primeras observaciones de campo indiquen la conveniencia de un ajuste menor. El uso de una grilla equivalente sera esencial para comparar resultados de forma interpretable con la primera etapa.",
        ],
    )
    add_heading(doc, "24.2 Campana P2_C02_layout_4A", 3)
    add_paragraphs(
        doc,
        [
            "La segunda campana con cuatro antenas ESP32 debe introducir modificaciones controladas del mobiliario o de la disposicion del ambiente. El objetivo no es degradar artificialmente el sistema, sino medir su sensibilidad ante cambios plausibles de propagacion. Esta etapa es relevante porque un sistema de localizacion indoor basado en fingerprinting no solo debe funcionar en un instante ideal, sino sostener comportamiento razonable frente a cambios moderados del entorno.",
            "En esta campana conviene documentar fotograficamente el estado del ambiente antes y despues de la modificacion, registrar con detalle que objetos se desplazaron y mantener sin cambios la posicion de las antenas ESP32. Solo asi sera posible atribuir diferencias de rendimiento a la variacion del ambiente y no a cambios simultaneos de geometria del despliegue.",
        ],
    )
    add_heading(doc, "24.3 Campana P2_C03_robustez_4A", 3)
    add_paragraphs(
        doc,
        [
            "La tercera campana con cuatro antenas ESP32 debe dedicarse a repetibilidad y robustez. La idea es repetir el protocolo con minima alteracion del contexto para estimar cuanta variabilidad surge simplemente de factores temporales, de pequeñas diferencias de operacion o del comportamiento estocastico del trafico radioelectrico. Esta informacion es particularmente importante si se pretende defender que el sistema no solo funciona una vez, sino que posee un comportamiento replicable.",
            "Para esta campana se sugiere extremar el registro de incidencias: reinicios de placas, demoras anormales, paquetes perdidos, cortes de red, tiempos por punto y cualquier otra observacion que permita contextualizar resultados. La robustez no se mide solo por el error final del modelo, sino tambien por la estabilidad operacional necesaria para obtener datos comparables.",
        ],
    )

    add_heading(doc, "25. Analisis global", 2)
    add_paragraphs(
        doc,
        [
            "El analisis global de la evidencia disponible permite sostener que el sistema ya supero la etapa de prueba puramente conceptual. Existen corridas reales, persistencia completa de datos, un dataset entrenable, modelos comparados y observaciones de comportamiento en linea. Esto no implica que la solucion este madura para un despliegue definitivo, pero si que la plataforma se encuentra en un estadio de validacion experimental genuina.",
            "La principal fortaleza observada es la coherencia end to end. Las antenas ESP32 capturan, el servidor persiste, la interfaz permite operar la campana y el pipeline produce artefactos reproducibles. Esta continuidad funcional es un logro relevante en un problema que combina hardware embebido, radiofrecuencia, backend, interfaz y aprendizaje automatico. Muchas veces los proyectos de localizacion fallan no por el modelo, sino por la imposibilidad de sostener una cadena experimental completa; en el presente caso, esa cadena ya existe.",
            "La principal debilidad de la primera etapa reside en la geometria con tres antenas ESP32 y en la estabilidad de largo plazo de algunos nodos. La evidencia cualitativa de inferencia y el razonamiento geométrico sugieren que la observabilidad no es uniforme sobre todo el plano. Del mismo modo, las incidencias de reconexion y la necesidad de robustecer el firmware indican que la confiabilidad del sistema depende todavia de ajustes adicionales.",
            "En terminos de precision, la corrida principal ofrece valores prometedores para un prototipo de bajo costo, pero todavia insuficientes para considerar cerrado el problema. Un MAE alrededor de 1,26 m y percentiles altos por encima de 2 m son defendibles como resultado preliminar, aunque dejan margen amplio para optimizacion. En este sentido, la futura prueba con cuatro antenas ESP32 no es simplemente deseable: es metodologicamente necesaria para verificar si la limitacion principal es geometrica o si persiste aun con una topologia mas favorable.",
            "La inferencia en linea, por su parte, debe interpretarse como una capacidad demostrada pero en proceso de consolidacion. El sistema puede producir estimaciones actuales a partir de paquetes recientes, pero la presencia de dispositivos ajenos al ambiente, la naturaleza oportunista del trafico observado y la no simultaneidad perfecta entre receptores introducen ruido que requiere filtros y criterio de operacion. Este hallazgo no invalida la funcionalidad, pero define claramente el alcance real de la primera version.",
        ],
    )
    add_heading(doc, "25.1 Analisis espacial de la geometria con tres antenas ESP32", 3)
    add_paragraphs(
        doc,
        [
            "La geometria utilizada en la primera etapa no forma un rectangulo completo de observacion, sino una disposicion con dos antenas ESP32 en un borde superior y una tercera mas centrada hacia el borde inferior. Esta distribucion aporta cobertura, pero no simetria plena. Desde una perspectiva geométrica, es razonable esperar que ciertas regiones del plano queden mejor condicionadas que otras para la inferencia.",
            "La experiencia cualitativa observada durante las pruebas coincide con esta interpretacion. Si bien el sistema fue capaz de ubicar aproximadamente el emisor objetivo en el ambiente, tambien se observaron desviaciones mas marcadas sobre uno de los ejes. Esta situacion refuerza la hipotesis de que una cuarta antena ESP32 en la esquina faltante podria aportar informacion adicional significativa y reducir ambigüedades espaciales.",
        ],
    )
    add_heading(doc, "25.2 Analisis temporal del proceso de captura", 3)
    add_paragraphs(
        doc,
        [
            "Los tiempos por punto muestran que la captura no es un proceso determinista de duracion fija. Aun manteniendo el mismo cupo de muestras por antena ESP32, diferentes puntos requirieron tiempos muy distintos. Esta variabilidad temporal es, en si misma, un resultado experimental. Indica que el sistema responde tanto a factores geométricos como a condiciones transitorias de reconexion, canal y trafico radioelectrico.",
            "En consecuencia, cualquier estimacion del tiempo total de una campana futura debe contemplar no solo un promedio, sino tambien dispersion y casos desfavorables. Esta conclusion es relevante para la planificacion de la prueba con cuatro antenas ESP32, ya que sumar un receptor puede mejorar observabilidad, pero tambien agregar otra fuente potencial de demora o inestabilidad si el firmware no se comporta de forma suficientemente robusta.",
        ],
    )
    add_heading(doc, "25.3 Lectura critica de la inferencia en linea", 3)
    add_paragraphs(
        doc,
        [
            "La inferencia en linea representa uno de los resultados mas ilustrativos del prototipo, pero tambien uno de los mas faciles de sobreinterpretar. Ver un punto o un mapa de calor en pantalla no equivale automaticamente a demostrar precision real. La estimacion en vivo depende del modelo entrenado, de la ventana temporal seleccionada, del historial acumulado y del trafico efectivamente observado en ese instante por las antenas ESP32.",
            "Por ese motivo, la inferencia en linea debe entenderse como una herramienta de supervision y validacion cualitativa, complementaria de las metricas obtenidas offline. Su valor es grande, porque permite visualizar inmediatamente tendencias y detectar comportamientos anomalos. Sin embargo, la conclusion principal sobre el rendimiento del sistema debe seguir apoyandose en campanas bien etiquetadas y en evaluaciones reproducibles sobre datasets persistidos.",
        ],
    )

    add_heading(doc, "26. Posibles mejoras", 2)
    add_paragraphs(
        doc,
        [
            "Las mejoras posibles del proyecto pueden organizarse en cuatro grupos: geometria y hardware, firmware y comunicaciones, procesamiento de datos y modelos, e instrumentacion experimental. Esta clasificacion permite priorizar acciones sin perder de vista que el sistema es integral. Mejorar solo el modelo sin fortalecer la captura, por ejemplo, tendria un impacto limitado.",
        ],
    )
    add_para(doc, "Tabla 26.1 - Propuestas de mejora priorizadas para la evolucion del sistema.")
    add_table(
        doc,
        ["Mejora", "Categoria", "Justificacion tecnica", "Impacto esperado"],
        [
            ["Agregar una cuarta antena ESP32 en la esquina faltante", "Geometria", "Mejorar observabilidad espacial y simetria del despliegue", "Alto"],
            ["Mantener altura y orientacion mas consistentes entre receptores", "Geometria", "Reducir sesgos sistematicos entre nodos", "Medio"],
            ["Usar fuentes de alimentacion mas estables o mejor filtradas", "Hardware", "Disminuir reinicios y estados erraticos", "Medio"],
            ["Optimizar la estrategia de barrido de canal", "Firmware", "Balancear captura y reconexion con menor perdida", "Alto"],
            ["Persistir parcialmente el buffer local ante fallas de red", "Firmware", "Reducir perdida de datos en reconexiones fallidas", "Medio"],
            ["Introducir telemetria mas detallada de errores de las placas", "Firmware", "Aumentar capacidad diagnostica", "Medio"],
            ["Comparar features basadas en mediana e IQR", "Datos", "Reducir sensibilidad a outliers de RSSI", "Alto"],
            ["Evaluar normalizacion por antena ESP32", "Datos", "Mitigar diferencias sistematicas entre receptores", "Medio"],
            ["Entrenar un modelo especifico para inferencia multidispositivo", "Modelos", "Evitar extrapolacion indebida de dispositivos externos", "Alto"],
            ["Aumentar la cantidad de campanas y unificar dataset maestro", "Modelos", "Mejorar generalizacion frente a cambios del ambiente", "Alto"],
            ["Agregar estimacion de incertidumbre por prediccion", "Modelos", "Distinguir predicciones confiables de las dudosas", "Medio"],
            ["Automatizar exportacion de resumen tecnico por campana", "Software", "Reducir tareas manuales de documentacion", "Medio"],
            ["Integrar capturas fotograficas y de interfaz al pipeline documental", "Instrumentacion", "Fortalecer evidencia visual del experimento", "Medio"],
            ["Fijar procedimientos operativos estandar para inicio de campana", "Metodologia", "Reducir errores humanos y variabilidad entre dias", "Alto"],
        ],
    )
    add_paragraphs(
        doc,
        [
            "Entre todas las mejoras posibles, la mas prioritaria en el corto plazo es completar la prueba con cuatro antenas ESP32 y tres campanas formalmente documentadas. Esta accion es la que posee mayor potencial para responder la pregunta central que dejan abierta las corridas actuales: hasta que punto la geometria de tres receptores limita la precision final.",
            "En segundo termino, conviene profundizar la robustez del firmware. La automatizacion del reinicio tras fallas repetidas fue un paso importante, pero todavia seria conveniente registrar mejor las causas de desconexion, el tiempo efectivo en cada fase del ciclo y el comportamiento diferencial entre nodos. Esa informacion permitiria reducir el componente ad hoc de la depuracion.",
            "Por ultimo, la evolucion del pipeline de modelado deberia orientarse a dos metas complementarias. La primera es mejorar el modelo del emisor objetivo a partir de mas campanas y de una geometria mas rica. La segunda es diferenciar con mas claridad la inferencia del objetivo conocido de la inferencia multidispositivo, ya que ambas tareas no comparten exactamente las mismas condiciones de informacion disponible.",
        ],
    )
    add_heading(doc, "26.1 Mejoras especificas para la prueba con cuatro antenas ESP32", 3)
    add_paragraphs(
        doc,
        [
            "La segunda etapa deberia iniciar con un protocolo de pre-chequeo fijo: verificacion de alimentacion de cada antena ESP32, confirmacion del identificador de nodo cargado, prueba corta de conectividad al servidor, validacion del router de soporte y comprobacion de que cada placa consulta correctamente su configuracion. Formalizar este pre-chequeo permitira que los problemas de la campana se analicen sobre datos de mejor calidad.",
            "Tambien conviene preparar de antemano un plan de nombres y convenciones para las tres campanas con cuatro antenas ESP32, asegurando que las carpetas y tablas se generen con trazabilidad uniforme. La claridad documental previa reducira carga operativa en el momento de medir y facilitara la comparacion posterior con la etapa de tres receptores.",
        ],
    )

    add_heading(doc, "27. Conclusiones", 2)
    add_paragraphs(
        doc,
        [
            "El proyecto logro materializar una plataforma experimental operativa para localizacion indoor basada en RSSI y antenas ESP32. Esta afirmacion se sustenta en hechos verificables: el sistema captura trafico, lo persiste con trazabilidad, permite operar campanas punto por punto, construye datasets supervisados, entrena modelos y ejecuta inferencia en linea. Por lo tanto, la propuesta supera ya el plano de una idea conceptual y se convierte en un prototipo funcional de investigacion aplicada.",
            "La primera etapa con tres antenas ESP32 mostro con claridad tanto el potencial como las limitaciones de la arquitectura. Por un lado, demostro que es posible obtener un dataset consistente y alcanzar errores medios del orden del metro y medio con hardware de bajo costo. Por otro, revelo que la geometria actual y la estabilidad operativa de los nodos todavia condicionan el rendimiento final. Esta tension entre viabilidad y limite es precisamente el tipo de resultado que se espera en una investigacion ingenieril honesta.",
            "La corrida piloto parcial tuvo valor porque hizo visibles los costos temporales y los problemas operativos de una configuracion muy exigente en cantidad de muestras. La corrida principal, en cambio, aporto la primera evidencia cuantitativa completa del sistema. En conjunto, ambas corridas conforman una base suficientemente solida para justificar la segunda etapa con cuatro antenas ESP32, la cual emerge como continuidad natural y no como un agregado arbitrario.",
            "Desde una perspectiva metodologica, una de las principales contribuciones del trabajo es la integracion coherente entre hardware embebido, backend, interfaz de operacion, persistencia documental y modelado supervisado. Esta integracion es relevante porque muchas dificultades practicas de la localizacion indoor aparecen en las transiciones entre componentes, no solo dentro del algoritmo de estimacion. El proyecto demostro que dichas transiciones pueden gestionarse de forma reproducible con herramientas accesibles.",
            "La prueba con cuatro antenas ESP32 y tres campanas formales constituye, por lo tanto, el paso experimental mas importante a continuacion. Sera ella la que permita determinar con mayor fundamento si la precision alcanzada en la primera etapa estaba limitada principalmente por la geometria, por el volumen de entrenamiento o por restricciones mas profundas del enfoque basado en RSSI. Sea cual sea el resultado, el trabajo realizado hasta aqui ya ofrece una base metodologica y tecnica suficientemente madura para sostener esa siguiente fase.",
        ],
    )


def build_references(doc: Document) -> None:
    doc.add_page_break()
    add_heading(doc, "28. Referencias tecnicas complementarias", 2)
    refs = [
        "Python Software Foundation. Python Documentation.",
        "FastAPI. Official Documentation.",
        "Uvicorn. ASGI Server Documentation.",
        "Streamlit. Official Documentation.",
        "pandas Documentation.",
        "NumPy Documentation.",
        "scikit-learn Documentation.",
        "matplotlib Documentation.",
        "joblib Documentation.",
        "Arduino IDE Documentation.",
        "Espressif Systems. ESP32 and ESP-IDF Wi-Fi Programming Documentation.",
    ]
    for ref in refs:
        add_para(doc, ref)


def validate_output(path: Path) -> dict[str, int]:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs]
    full_text = "\n".join(paragraphs).lower()
    if re.search(r"\bancla\b", full_text):
        raise RuntimeError("El documento generado contiene la palabra prohibida 'ancla'.")
    if re.search(r"\bpoc\b", full_text):
        raise RuntimeError("El documento generado contiene la abreviatura prohibida 'PoC'.")
    required_terms = [
        "PARTE II - DESARROLLO E IMPLEMENTACION",
        "PARTE III - RESULTADOS, ANALISIS, MEJORAS Y CONCLUSIONES",
        "demo_s1 / train_01",
        "demo_s2 / train_02",
        "Prueba 2 con 4 antenas ESP32",
    ]
    for term in required_terms:
        if term.lower() not in full_text:
            raise RuntimeError(f"Falta el termino obligatorio: {term}")
    return {
        "paragraphs": len(doc.paragraphs),
        "tables": len(doc.tables),
        "images": len(doc.inline_shapes),
    }


def build_document() -> tuple[Path, dict[str, int]]:
    run_1a = load_run_detail(RUN_1A)
    run_1b = load_run_detail(RUN_1B)
    figs = build_generated_figures(run_1a, run_1b)

    doc = Document(str(TEMPLATE_DOCX))
    clear_document(doc)
    build_part_ii(doc, run_1a, run_1b, figs)
    build_part_iii(doc, run_1a, run_1b, figs)
    build_references(doc)
    doc.save(str(OUTPUT_DOCX))
    stats = validate_output(OUTPUT_DOCX)
    return OUTPUT_DOCX, stats


def main() -> None:
    output, stats = build_document()
    print(f"Documento generado: {output}")
    print(f"Parrafos: {stats['paragraphs']}")
    print(f"Tablas: {stats['tables']}")
    print(f"Imagenes: {stats['images']}")


if __name__ == "__main__":
    main()
