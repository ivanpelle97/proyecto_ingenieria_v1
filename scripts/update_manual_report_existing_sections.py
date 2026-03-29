from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = Path(r"C:\Users\hi_iv\Downloads\PROYECTO_FINAL_CORREGIDO_MANUAL.docx")
OUTPUT_DOC = ROOT / "PROYECTO_FINAL_CORREGIDO_MANUAL_v3_integrado.docx"
RUN_3A = ROOT / "runs" / "demo_s2" / "train_02"
RUN_4A = ROOT / "runs" / "demo_s4" / "train_04"
SCREENSHOT_DIR = Path(r"C:\Users\hi_iv\OneDrive\Imágenes\Capturas de pantalla")
GENERATED_DIR = ROOT / "doc_assets" / "generated"


REFERENCE_TEXTS = [
    "Python Software Foundation. Python Documentation.",
    "FastAPI. Official Documentation.",
    "Uvicorn. ASGI Server Documentation.",
    "Streamlit. Official Documentation.",
    "pandas Documentation.",
    "NumPy Documentation.",
    "scikit-learn Documentation.",
    "matplotlib Documentation.",
    "joblib Documentation.",
    "Arduino IDE Documentation.",
    "Espressif Systems. ESP32 and ESP-IDF Wi-Fi Programming Documentation.",
]


def ensure_dirs() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_table_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True


def set_cell_text(cell, text: object, bold: bool = False) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)
            run.bold = bold


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, width: float = 6.3) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def create_placeholder_image(title: str, subtitle: str, output_path: Path) -> Path:
    image = Image.new("RGB", (1600, 900), (242, 242, 242))
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    draw.rectangle((40, 40, 1560, 860), outline=(120, 120, 120), width=4)
    draw.multiline_text((130, 250), title, fill=(20, 20, 20), font=font, spacing=12)
    draw.multiline_text((130, 430), subtitle, fill=(70, 70, 70), font=font, spacing=10)
    image.save(output_path)
    return output_path


def create_contact_sheet(paths: list[Path], output_path: Path) -> Path:
    tile_w, tile_h = 420, 240
    cols = 2
    rows = (len(paths) + cols - 1) // cols
    canvas = Image.new("RGB", (cols * tile_w + 80, rows * (tile_h + 50) + 80), (255, 255, 255))
    draw = ImageDraw.Draw(canvas)
    font = ImageFont.load_default()
    for index, path in enumerate(paths):
        row = index // cols
        col = index % cols
        x = 40 + col * tile_w
        y = 40 + row * (tile_h + 50)
        with Image.open(path) as image:
            image = image.convert("RGB")
            image.thumbnail((tile_w - 20, tile_h - 20))
            tile = Image.new("RGB", (tile_w - 20, tile_h - 20), (248, 249, 250))
            ox = (tile.width - image.width) // 2
            oy = (tile.height - image.height) // 2
            tile.paste(image, (ox, oy))
            canvas.paste(tile, (x + 10, y + 10))
        draw.rectangle((x, y, x + tile_w - 1, y + tile_h - 1), outline=(185, 185, 185), width=2)
        draw.text((x, y + tile_h + 8), path.stem.replace("Captura de pantalla ", ""), fill=(20, 20, 20), font=font)
    canvas.save(output_path)
    return output_path


def load_run_summary(run_dir: Path) -> dict:
    experiment = json.loads((run_dir / "experiment.json").read_text(encoding="utf-8"))
    points = pd.read_csv(run_dir / "points.csv")
    samples = pd.read_csv(run_dir / "samples.csv")
    dataset = pd.read_csv(run_dir / "dataset.csv")
    metrics = pd.read_csv(run_dir / "models" / "metrics.csv")

    points["started_at"] = pd.to_datetime(points["started_at"])
    points["capture_complete_at"] = pd.to_datetime(points["capture_complete_at"])
    points["completed_at"] = pd.to_datetime(points["completed_at"])
    points["capture_duration_s"] = (points["capture_complete_at"] - points["started_at"]).dt.total_seconds()
    points["close_duration_s"] = (points["completed_at"] - points["capture_complete_at"]).dt.total_seconds()
    points["total_duration_s"] = (points["completed_at"] - points["started_at"]).dt.total_seconds()

    rssi_stats = (
        samples.groupby("anchor_id")["rssi"]
        .agg(["count", "min", "max", "mean", "std"])
        .reset_index()
        .rename(columns={"count": "n"})
    )

    return {
        "experiment": experiment,
        "points": points,
        "samples": samples,
        "dataset": dataset,
        "metrics": metrics,
        "rssi_stats": rssi_stats,
        "raw_sizes": {path.stem: path.stat().st_size for path in sorted((run_dir / "raw").glob("*.jsonl"))},
        "best_mae": metrics.sort_values("mae_eucl").iloc[0].to_dict(),
        "best_p95": metrics.sort_values("p95").iloc[0].to_dict(),
    }


def find_screenshots() -> list[Path]:
    return sorted([path for path in SCREENSHOT_DIR.glob("*.png") if "2026-03-23" in path.name])


def selected_screenshots(paths: list[Path]) -> list[Path]:
    selected_names = [
        "Captura de pantalla 2026-03-23 163502.png",
        "Captura de pantalla 2026-03-23 163526.png",
        "Captura de pantalla 2026-03-23 180641.png",
        "Captura de pantalla 2026-03-23 181743.png",
        "Captura de pantalla 2026-03-23 184258.png",
    ]
    selected = []
    by_name = {path.name: path for path in paths}
    for name in selected_names:
        if name in by_name:
            selected.append(by_name[name])
    return selected


def remove_body_from_heading_to_end(doc: Document, heading_prefix: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_para = next(p for p in doc.paragraphs if p.text.strip().startswith(heading_prefix))
    start_idx = children.index(start_para._element)
    for child in children[start_idx:-1]:
        body.remove(child)


def add_section_24(doc: Document, run4: dict, screenshots: list[Path]) -> None:
    exp = run4["experiment"]
    points = run4["points"]
    metrics = run4["metrics"].sort_values("mae_eucl")

    add_heading(doc, "24. Prueba 2 con 4 antenas ESP32", level=1)
    add_body_paragraph(
        doc,
        "La segunda etapa experimental ya no constituye una planificación futura, sino una prueba efectivamente ejecutada. "
        "La campaña demo_s4 / train_04 se realizó con cuatro antenas ESP32 ubicadas en las cuatro esquinas del recinto, con el objetivo "
        "de reducir la asimetría geométrica observada en la etapa previa y de verificar si una topología de observación más cerrada "
        "mejoraba la calidad del posicionamiento indoor basado en RSSI."
    )
    add_body_paragraph(
        doc,
        "La campaña se desarrolló con la misma grilla de dieciocho puntos utilizada en la etapa principal de tres antenas ESP32, "
        "con diez muestras por antena ESP32 y por punto y con el mismo emisor objetivo identificado por la MAC 02:11:22:33:44:55. "
        "La principal diferencia experimental fue, por lo tanto, la inclusión del cuarto receptor y la evaluación posterior del comportamiento "
        "de la inferencia en línea bajo esta nueva geometría."
    )

    add_heading(doc, "24.1 Campaña ejecutada con 4 antenas ESP32: configuración y resultados medidos", level=2)
    add_body_paragraph(
        doc,
        "La campaña demo_s4 / train_04 logró completarse en los dieciocho puntos de referencia previstos. En todos los puntos se alcanzó el "
        "cupo de diez muestras por cada una de las cuatro antenas ESP32, lo que produjo un total de setecientas veinte muestras aceptadas y un "
        "dataset final de dieciocho filas por setenta y cinco columnas. La evidencia de captura muestra, por lo tanto, una completitud "
        "experimental plena para esta segunda prueba."
    )

    add_table_title(doc, "Tabla 24.1 - Geometría real empleada en la prueba con cuatro antenas ESP32.")
    geometry_rows = [
        [anchor["anchor_id"], f"{anchor['x_m']:.2f}", f"{anchor['y_m']:.2f}", f"{anchor['z_m']:.2f}", anchor.get("model", "ESP32")]
        for anchor in exp["anchors"]
    ]
    add_table(doc, ["Antena ESP32", "x [m]", "y [m]", "z [m]", "Modelo"], geometry_rows)

    add_table_title(doc, "Tabla 24.2 - Resumen global de la campaña demo_s4 / train_04.")
    add_table(
        doc,
        ["Variable", "Valor medido"],
        [
            ["Sesión", exp["session_id"]],
            ["Campaña", exp["campaign_id"]],
            ["MAC del emisor objetivo", exp["target_mac"]],
            ["Antenas ESP32", len(exp["anchors"])],
            ["Puntos completados", len(points)],
            ["Muestras aceptadas", len(run4["samples"])],
            ["Muestras por antena ESP32 y por punto", exp["samples_per_anchor"]],
            ["Dataset final", f"{run4['dataset'].shape[0]} filas x {run4['dataset'].shape[1]} columnas"],
            ["Mejor MAE", f"{run4['best_mae']['model']} = {run4['best_mae']['mae_eucl']:.4f} m"],
            ["Mejor P95", f"{run4['best_p95']['model']} = {run4['best_p95']['p95']:.4f} m"],
        ],
    )

    add_table_title(doc, "Tabla 24.3 - Resultados por punto de referencia en la campaña con cuatro antenas ESP32.")
    point_rows = []
    for row in points.itertuples(index=False):
        point_rows.append(
            [
                row.point_id,
                f"{row.x_m:.1f}",
                f"{row.y_m:.1f}",
                f"{row.z_m:.1f}",
                row.started_at.strftime("%H:%M:%S"),
                f"{row.capture_duration_s:.2f}",
                f"{row.close_duration_s:.2f}",
                f"{row.total_duration_s:.2f}",
                int(row.A1_count),
                int(row.A2_count),
                int(row.A3_count),
                int(row.A4_count),
            ]
        )
    add_table(
        doc,
        ["Punto", "x [m]", "y [m]", "z [m]", "Inicio", "Captura [s]", "Cierre [s]", "Total [s]", "A1", "A2", "A3", "A4"],
        point_rows,
    )

    add_table_title(doc, "Tabla 24.4 - Volumen de muestras y tamaño de archivos RAW por antena ESP32.")
    sample_counts = run4["samples"]["anchor_id"].value_counts().to_dict()
    volume_rows = []
    for anchor_id in sorted(sample_counts):
        raw_size = run4["raw_sizes"].get(anchor_id, 0)
        volume_rows.append([anchor_id, sample_counts[anchor_id], raw_size, f"{raw_size / 1024:.1f} KiB"])
    add_table(doc, ["Antena ESP32", "Muestras aceptadas", "RAW [bytes]", "RAW [KiB]"], volume_rows)

    add_table_title(doc, "Tabla 24.5 - Métricas de modelos obtenidas con la campaña demo_s4 / train_04.")
    metric_rows = []
    for row in metrics.itertuples(index=False):
        metric_rows.append(
            [
                row.model,
                f"{row.mae_eucl:.4f}",
                f"{row.p50:.4f}",
                f"{row.p90:.4f}",
                f"{row.p95:.4f}",
                f"{row.rmse_eucl:.4f}",
            ]
        )
    add_table(doc, ["Modelo", "MAE [m]", "P50 [m]", "P90 [m]", "P95 [m]", "RMSE [m]"], metric_rows)

    add_body_paragraph(
        doc,
        "El mejor modelo por error absoluto medio fue extratrees_500, con un MAE de "
        f"{run4['best_mae']['mae_eucl']:.4f} m. El mejor percentil 95 correspondió a {run4['best_p95']['model']}, con un valor de "
        f"{run4['best_p95']['p95']:.4f} m. En términos de completitud de captura, la campaña fue consistente; sin embargo, desde el punto de vista "
        "predictivo, estos valores no mejoraron a la mejor corrida con tres antenas ESP32."
    )

    add_heading(doc, "24.2 Registro visual de la campaña y de la inferencia en línea", level=2)
    add_body_paragraph(
        doc,
        "Las capturas de pantalla obtenidas durante la ejecución documentan tres momentos de interés: la configuración de la campaña, el avance "
        "de la toma de puntos y el comportamiento de la inferencia en línea una vez entrenado el modelo. Estas imágenes son útiles porque muestran "
        "cómo se manifestaron visualmente los resultados que luego quedaron persistidos en los archivos CSV y JSONL."
    )

    contact_sheet = create_contact_sheet(screenshots, GENERATED_DIR / "contacto_prueba2_integrada.png")
    add_picture_with_caption(
        doc,
        contact_sheet,
        "Figura 24.1. Conjunto de capturas de pantalla correspondientes a la campaña del 23 de marzo de 2026 con cuatro antenas ESP32.",
    )

    selected = selected_screenshots(screenshots)
    captions = {
        "Captura de pantalla 2026-03-23 163502.png": "Figura 24.2. Configuración inicial de la campaña demo_s4 / train_04 con definición de la geometría de cuatro antenas ESP32.",
        "Captura de pantalla 2026-03-23 163526.png": "Figura 24.3. Estado del plano durante la fase de captura, con el punto P03 como punto actual de entrenamiento.",
        "Captura de pantalla 2026-03-23 180641.png": "Figura 24.4. Plano completo con los dieciocho puntos registrados al finalizar la campaña de entrenamiento.",
        "Captura de pantalla 2026-03-23 181743.png": "Figura 24.5. Ejemplo de inferencia en línea con múltiples dispositivos detectados y mapa de calor activo.",
        "Captura de pantalla 2026-03-23 184258.png": "Figura 24.6. Estado final de inferencia en línea donde se aprecia la estimación del emisor objetivo respecto del plano entrenado.",
    }
    for image_path in selected:
        add_picture_with_caption(doc, image_path, captions[image_path.name])

    placeholder_mount = create_placeholder_image(
        "Insertar fotografía del montaje físico de la prueba con cuatro antenas ESP32",
        "Se recomienda incluir la vista del recinto completo, la ubicación de cada antena ESP32,\nel router de soporte y la orientación respecto del plano cargado en la interfaz.",
        GENERATED_DIR / "placeholder_24_montaje_4_antenas.png",
    )
    placeholder_tag = create_placeholder_image(
        "Insertar fotografía de la posición del TAG durante la inferencia en línea",
        "La imagen debe mostrar la ubicación real del emisor objetivo,\nla prueba concreta realizada y la captura de interfaz asociada.",
        GENERATED_DIR / "placeholder_24_tag_inferencia.png",
    )
    placeholder_train = create_placeholder_image(
        "Insertar fotografía de la medición de puntos de entrenamiento",
        "La imagen debe mostrar el punto rotulado, la colocación del TAG y la correspondencia\ncon la tabla de puntos utilizada durante la campaña.",
        GENERATED_DIR / "placeholder_24_punto_entrenamiento.png",
    )
    add_picture_with_caption(doc, placeholder_mount, "Figura 24.7. Espacio reservado para la fotografía del montaje físico de la prueba con cuatro antenas ESP32.")
    add_picture_with_caption(doc, placeholder_tag, "Figura 24.8. Espacio reservado para la fotografía de la posición real del TAG durante la inferencia en línea.")
    add_picture_with_caption(doc, placeholder_train, "Figura 24.9. Espacio reservado para la fotografía de la medición de un punto de entrenamiento.")

    add_heading(doc, "24.3 Reformulación experimental posterior y alcance de la tercera prueba", level=2)
    add_body_paragraph(
        doc,
        "La lectura conjunta de métricas, tiempos y comportamiento visual de la inferencia condujo a una reformulación metodológica. "
        "La segunda prueba con cuatro antenas ESP32 mostró que la cobertura espacial mejoró y que el sistema podía recolectar datos completos; "
        "sin embargo, la inferencia en línea continuó condicionada por ventanas de detección amplias y por la espera práctica hasta disponer de "
        "muestras de las cuatro antenas ESP32. Ese mecanismo introdujo demoras y redujo la sensación de continuidad operacional."
    )
    add_body_paragraph(
        doc,
        "A partir de esta observación se definió una tercera prueba centrada específicamente en inferencia en línea continua. Su objetivo ya no "
        "es volver a levantar un dataset completo punto a punto, sino reutilizar una campaña entrenada y evaluar un flujo constante desde las cuatro "
        "antenas ESP32 mediante firmware de captura continua, canal fijo del punto de acceso, emisión del emisor objetivo sobre el mismo canal y "
        "ventanas temporales más estrechas para consolidación de la estimación."
    )
    add_table_title(doc, "Tabla 24.6 - Lineamientos operativos de la tercera prueba propuesta a partir de los hallazgos de la prueba con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Componente", "Condición observada en la prueba 2", "Ajuste previsto para la prueba 3", "Objetivo"],
        [
            ["Antenas ESP32", "Reporte con demoras variables", "Captura continua con conexión persistente", "Asegurar flujo estable hacia el servidor"],
            ["Emisor objetivo", "Operación sin sincronización explícita con el canal del AP", "Emisión en el mismo canal fijo del AP", "Aumentar coincidencia temporal de observación"],
            ["Interfaz de operación", "Campaña de entrenamiento e inferencia posterior", "Modo de solo inferencia sobre campaña ya entrenada", "Separar evaluación en línea del costo de captura"],
            ["Ventanas temporales", "Amplias", "Reducidas y parametrizadas", "Evitar arrastre temporal excesivo"],
            ["Criterio de disponibilidad", "Dependencia práctica de cuatro antenas ESP32", "Uso de memoria por antena y mínimos configurables", "Reducir tiempo hasta una predicción útil"],
        ],
    )


def add_section_25(doc: Document, run3: dict, run4: dict) -> None:
    add_heading(doc, "25. Análisis global", level=1)
    add_body_paragraph(
        doc,
        "La evidencia acumulada hasta este punto permite afirmar que el proyecto ya fue sometido a dos configuraciones geométricas reales: una "
        "campaña principal con tres antenas ESP32 y una campaña completa con cuatro antenas ESP32. Esto fortalece el carácter experimental del trabajo, "
        "porque la discusión ya no se apoya en una única geometría ni en una sola corrida satisfactoria."
    )
    add_body_paragraph(
        doc,
        "La prueba con cuatro antenas ESP32 confirmó que el sistema puede sostener una adquisición completa y balanceada de datos en una grilla de "
        "dieciocho puntos. Sin embargo, también mostró que una geometría más rica no garantiza por sí sola una mejora del modelo. El análisis debe "
        "desplazarse, entonces, desde la mera cobertura espacial hacia la interacción entre geometría, ventanas temporales, sincronía práctica entre "
        "receptores y modo de construcción de la inferencia en línea."
    )

    add_heading(doc, "25.1 Comparación directa entre la campaña principal con tres antenas ESP32 y la prueba con cuatro antenas ESP32", level=2)
    add_table_title(doc, "Tabla 25.1 - Comparación cuantitativa entre la mejor corrida con tres antenas ESP32 y la campaña ejecutada con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Configuración", "Puntos", "Muestras", "Dataset", "Mejor modelo MAE", "MAE [m]", "Mejor modelo P95", "P95 [m]"],
        [
            [
                "Tres antenas ESP32 (demo_s2 / train_02)",
                len(run3["points"]),
                len(run3["samples"]),
                f"{run3['dataset'].shape[0]} x {run3['dataset'].shape[1]}",
                run3["best_mae"]["model"],
                f"{run3['best_mae']['mae_eucl']:.4f}",
                run3["best_p95"]["model"],
                f"{run3['best_p95']['p95']:.4f}",
            ],
            [
                "Cuatro antenas ESP32 (demo_s4 / train_04)",
                len(run4["points"]),
                len(run4["samples"]),
                f"{run4['dataset'].shape[0]} x {run4['dataset'].shape[1]}",
                run4["best_mae"]["model"],
                f"{run4['best_mae']['mae_eucl']:.4f}",
                run4["best_p95"]["model"],
                f"{run4['best_p95']['p95']:.4f}",
            ],
        ],
    )
    add_body_paragraph(
        doc,
        "La comparación deja una conclusión metodológica importante: la campaña de cuatro antenas ESP32 consiguió más información por punto, pero no "
        "produjo un mejor error final. En la práctica, esto significa que el desempeño del sistema está siendo afectado por factores adicionales a la "
        "geometría, entre ellos el modo de consolidación temporal de las muestras y la forma en que la inferencia en línea consume la información disponible."
    )

    add_heading(doc, "25.2 Lectura crítica de la inferencia en línea y de las ventanas temporales", level=2)
    add_body_paragraph(
        doc,
        "Las capturas de la interfaz mostraron un fenómeno consistente: el mapa de calor tendía a concentrar dispositivos en una región central del plano "
        "y la estimación del emisor objetivo no siempre reaccionaba con la inmediatez esperada. Este comportamiento es compatible con el uso de ventanas "
        "temporales amplias y con la necesidad práctica de esperar a que la información de las cuatro antenas ESP32 complete un patrón suficientemente estable."
    )
    add_body_paragraph(
        doc,
        "El hallazgo no invalida el enfoque. Lo que indica es que la fase en línea requiere una estrategia específica, distinta de la etapa de captura "
        "supervisada. La tercera prueba propuesta apunta precisamente a esa necesidad: reducir la inercia temporal y observar si una política de memoria "
        "por antena ESP32 más acotada mejora la respuesta sin perder estabilidad."
    )

    add_heading(doc, "25.3 Estabilidad operativa de las placas y continuidad de reporte", level=2)
    add_body_paragraph(
        doc,
        "La segunda prueba también reforzó la necesidad de robustecer el funcionamiento embebido de las placas. Aunque la campaña pudo completarse, "
        "se observaron episodios de reconexión lenta y situaciones en las que algunas antenas ESP32 demoraban en recuperar el enlace con el punto de acceso. "
        "Esta dimensión operacional es crucial porque, en una arquitectura distribuida, la calidad de la inferencia depende de la disponibilidad efectiva de "
        "cada receptor y no solamente del algoritmo de aprendizaje automático."
    )
    add_body_paragraph(
        doc,
        "En este contexto, la decisión de pasar a una tercera prueba con firmware continuo, canal fijo y reinicio automático tras fallas repetidas "
        "aparece plenamente justificada. No se trata de una mejora incremental menor, sino de una respuesta concreta a un cuello de botella evidenciado "
        "durante la ejecución real de la prueba con cuatro antenas ESP32."
    )


def add_section_26(doc: Document) -> None:
    add_heading(doc, "26. Posibles mejoras", level=1)
    add_body_paragraph(
        doc,
        "Las mejoras más relevantes del proyecto dejaron de ser puramente hipotéticas, porque ya están conectadas con problemas efectivamente medidos. "
        "La primera línea de trabajo consiste en separar con mayor claridad las exigencias del entrenamiento supervisado y las exigencias de la inferencia "
        "en línea. La segunda se orienta a consolidar la estabilidad operativa de las antenas ESP32. La tercera apunta a enriquecer la documentación visual "
        "y experimental para sostener la defensa metodológica del proyecto."
    )

    add_table_title(doc, "Tabla 26.1 - Acciones de mejora priorizadas a partir de la prueba con cuatro antenas ESP32.")
    add_table(
        doc,
        ["Prioridad", "Mejora", "Justificación técnica", "Impacto esperado"],
        [
            ["Alta", "Firmware continuo para antenas ESP32", "Reducir reconexiones y sostener flujo constante", "Mayor continuidad durante inferencia en línea"],
            ["Alta", "Emisor objetivo en canal fijo del AP", "Aumentar coincidencia temporal entre emisión y escucha", "Menor espera hasta una predicción utilizable"],
            ["Alta", "Ventanas temporales reducidas", "Evitar arrastre excesivo de historial", "Mayor respuesta temporal del punto estimado"],
            ["Media", "Criterios mínimos por antena ESP32 configurables", "Disminuir dependencia de simultaneidad perfecta", "Robustez frente a asincronías parciales"],
            ["Media", "Mayor registro fotográfico de campo", "Mejor trazabilidad espacial y metodológica", "Fortalecimiento documental del informe final"],
            ["Media", "Nuevas campañas con muebles y escenarios controlados", "Ampliar diversidad del entorno", "Mejor generalización del modelo"],
        ],
    )

    add_heading(doc, "26.1 Mejoras específicas para la tercera prueba de inferencia continua", level=2)
    add_body_paragraph(
        doc,
        "La tercera prueba debe diseñarse como una verificación operativa enfocada en continuidad temporal. Para ello conviene reutilizar la campaña ya "
        "entrenada y activar el sistema en modo de solo inferencia, evitando levantar un nuevo dataset en esta fase. Esta decisión permitirá evaluar el "
        "desempeño en línea sin mezclarlo con los tiempos de captura de puntos supervisados."
    )
    add_table_title(doc, "Tabla 26.2 - Plantilla preparada para registrar resultados de la tercera prueba.")
    add_table(
        doc,
        ["Escenario", "Ventana [s]", "Memoria por antena [s]", "Mínimo por antena", "Antenas requeridas", "TAG ubicado en", "Comportamiento observado", "Conclusión"],
        [
            ["T3_E01", "Completar", "Completar", "Completar", "Completar", "Completar", "Completar luego de la prueba", "Completar luego de la prueba"],
            ["T3_E02", "Completar", "Completar", "Completar", "Completar", "Completar", "Completar luego de la prueba", "Completar luego de la prueba"],
            ["T3_E03", "Completar", "Completar", "Completar", "Completar", "Completar", "Completar luego de la prueba", "Completar luego de la prueba"],
        ],
    )


def add_section_27(doc: Document, run3: dict, run4: dict) -> None:
    add_heading(doc, "27. Conclusiones", level=1)
    add_body_paragraph(
        doc,
        "El proyecto alcanzó un grado de madurez experimental superior al que presentaba al cierre de la primera versión del informe. Hoy existe evidencia "
        "real no solo de una campaña principal con tres antenas ESP32, sino también de una campaña completa con cuatro antenas ESP32, con persistencia de datos, "
        "modelos entrenados, resultados comparables y capturas visuales del proceso de inferencia en línea."
    )
    add_body_paragraph(
        doc,
        "La principal conclusión técnica es que el aumento de receptores no resolvió automáticamente la precisión final del sistema. La campaña con cuatro "
        "antenas ESP32 aportó completitud de medición y mejor cierre geométrico del recinto, pero sus métricas no superaron a las de la mejor campaña con "
        "tres antenas ESP32. Este resultado es valioso porque desplaza la discusión desde una explicación simplista basada solo en geometría hacia una lectura "
        "más profunda, donde la dinámica temporal de adquisición y la forma de construir la inferencia en línea resultan determinantes."
    )
    add_body_paragraph(
        doc,
        "En consecuencia, la tercera prueba propuesta no constituye una expansión arbitraria del trabajo, sino la respuesta natural al hallazgo más importante "
        "de la segunda prueba. Si las ventanas de detección amplias y la espera práctica de muestras desde las cuatro antenas ESP32 están penalizando la operación "
        "en línea, entonces corresponde rediseñar esa fase con firmware continuo, canal fijo compartido y criterios temporales más estrictos."
    )
    add_body_paragraph(
        doc,
        "Desde el punto de vista ingenieril, el valor del proyecto reside precisamente en esta secuencia de aprendizaje: una primera validación end to end, una "
        "segunda prueba con cuatro antenas ESP32 y una reformulación posterior basada en evidencia medida. Esta trayectoria muestra una investigación aplicada que "
        "no se limita a presentar un resultado favorable, sino que documenta también los límites reales del sistema y transforma esos límites en decisiones concretas "
        "de rediseño experimental."
    )


def add_references(doc: Document) -> None:
    add_heading(doc, "28. Referencias técnicas complementarias", level=1)
    for reference in REFERENCE_TEXTS:
        paragraph = doc.add_paragraph(reference)
        for run in paragraph.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(11)


def main() -> None:
    ensure_dirs()
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"No se encontró el documento fuente: {SOURCE_DOC}")

    run3 = load_run_summary(RUN_3A)
    run4 = load_run_summary(RUN_4A)
    screenshots = find_screenshots()
    if not screenshots:
        raise FileNotFoundError(f"No se encontraron capturas del 23 de marzo de 2026 en {SCREENSHOT_DIR}")

    doc = Document(str(SOURCE_DOC))
    remove_body_from_heading_to_end(doc, "24. Prueba 2 con 4 antenas ESP32")
    add_section_24(doc, run4, screenshots)
    add_section_25(doc, run3, run4)
    add_section_26(doc)
    add_section_27(doc, run3, run4)
    add_references(doc)
    doc.save(str(OUTPUT_DOC))
    print(f"Documento generado en: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
