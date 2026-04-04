from __future__ import annotations

import json
from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
SOURCE_DOC = ROOT / "PROYECTO_FINAL_CORREGIDO_MANUAL_v3_integrado.docx"
OUTPUT_DOC = ROOT / "PROYECTO_FINAL_CORREGIDO_MANUAL_v4_entrega_final.docx"
THEORETICAL_DOC = ROOT / "PROYECTO DE INGENIERIA v3_MARCO_TEORICO.docx"

RUN_S1 = ROOT / "runs" / "demo_s1" / "train_01"
RUN_S2 = ROOT / "runs" / "demo_s2" / "train_02"
RUN_S3 = ROOT / "runs" / "demo_s3" / "train_03"
RUN_S4 = ROOT / "runs" / "demo_s4" / "train_04"
RUN_S5 = ROOT / "runs" / "demo_s5" / "train_05"

SCREENSHOT_DIR = Path("C:/Users/hi_iv/OneDrive/Im\u00e1genes/Capturas de pantalla")
GENERATED_DIR = ROOT / "doc_assets" / "generated"
PENDING_PHOTO_DIR = ROOT / "doc_assets" / "pending_phone_photos"

REFERENCE_TEXTS = [
    "Documento interno de referencia teorica: PROYECTO DE INGENIERIA v3_MARCO_TEORICO.docx.",
    "Python Software Foundation. Python Documentation.",
    "FastAPI. Official Documentation.",
    "Uvicorn. Official Documentation.",
    "Streamlit. Official Documentation.",
    "pandas Documentation.",
    "NumPy Documentation.",
    "scikit-learn Documentation.",
    "matplotlib Documentation.",
    "joblib Documentation.",
    "Espressif Systems. ESP32 and ESP-IDF Wi-Fi Programming Documentation.",
    "Arduino IDE Documentation.",
]

OLD_ANTENNA_CODE = """static void sta_set_config(uint8_t ap_channel_hint, bool use_channel_hint, bool fast_scan) {
  wifi_config_t cfg = {};
  strncpy((char*)cfg.sta.ssid, WIFI_SSID, sizeof(cfg.sta.ssid) - 1);
  strncpy((char*)cfg.sta.password, WIFI_PASS, sizeof(cfg.sta.password) - 1);
  cfg.sta.scan_method = fast_scan ? WIFI_FAST_SCAN : WIFI_ALL_CHANNEL_SCAN;
  cfg.sta.sort_method = WIFI_CONNECT_AP_BY_SIGNAL;
  cfg.sta.channel = use_channel_hint ? ap_channel_hint : 0;
  esp_wifi_set_config(WIFI_IF_STA, &cfg);
}

static bool enter_sta_and_wait_ip(uint8_t ap_channel_hint, bool use_channel_hint, bool fast_scan) {
  exit_sniffer_mode();
  esp_wifi_set_mode(WIFI_MODE_STA);
  sta_set_config(ap_channel_hint, use_channel_hint, fast_scan);
  esp_wifi_connect();
  EventBits_t bits = xEventGroupWaitBits(evg, GOT_IP_BIT, pdFALSE, pdTRUE, pdMS_TO_TICKS(STA_TIMEOUT_MS));
  return (bits & GOT_IP_BIT) != 0;
}"""

NEW_TRAINING_CODE = """static bool should_capture_now() {
  return config_enabled && capture_active && target_mac_valid;
}

static void apply_sta_config() {
  wifi_config_t cfg = {};
  strncpy((char *)cfg.sta.ssid, WIFI_SSID, sizeof(cfg.sta.ssid) - 1);
  strncpy((char *)cfg.sta.password, WIFI_PASS, sizeof(cfg.sta.password) - 1);
  cfg.sta.scan_method = WIFI_ALL_CHANNEL_SCAN;
  cfg.sta.channel = USE_BSSID_LOCK ? WIFI_AP_CHANNEL : 0;
  cfg.sta.bssid_set = USE_BSSID_LOCK ? 1 : 0;
  if (USE_BSSID_LOCK) {
    memcpy(cfg.sta.bssid, WIFI_AP_BSSID, sizeof(cfg.sta.bssid));
  }
  esp_wifi_set_config(WIFI_IF_STA, &cfg);
}

static void sniffer_cb(void *buff, wifi_promiscuous_pkt_type_t type) {
  if (type != WIFI_PKT_MGMT) return;
  if (!should_capture_now()) return;
  ...
}"""

NEW_INFERENCE_CODE = """static uint32_t anchor_connect_slot_delay_ms() {
  if (strcmp(ANCHOR_ID, "A1") == 0) return 0;
  if (strcmp(ANCHOR_ID, "A2") == 0) return 700;
  if (strcmp(ANCHOR_ID, "A3") == 0) return 1400;
  if (strcmp(ANCHOR_ID, "A4") == 0) return 2100;
  return 0;
}

static void ensure_wifi_connected_or_restart() {
  for (int attempt = 1; attempt <= WIFI_RETRIES; attempt++) {
    uint32_t slot_delay = (attempt == 1) ? anchor_connect_slot_delay_ms() : 0;
    uint32_t jitter = retry_jitter_ms();
    vTaskDelay(pdMS_TO_TICKS(slot_delay + jitter));
    if (connect_sta_once()) return;
    hard_reset_wifi_radio();
  }
  restart_after_wifi_failure();
}

static void sniffer_cb(void *buff, wifi_promiscuous_pkt_type_t type) {
  if (type != WIFI_PKT_MGMT) return;
  if (!config_enabled) return;
  if (target_mac_valid && (capture_active || live_mode_target_only())) {
    ...
  }
}"""

OLD_TAG_CODE = """static const bool BURST_ALL_CHANNELS = true;
static const uint8_t TAG_CHANNELS[] = {1, 2, 3, 4, 5, 6, 7, 8, 9, 10, 11, 12, 13};
static const uint32_t TX_CYCLE_INTERVAL_MS = 500;
static const uint32_t INTER_CHANNEL_GAP_MS = 18;
static const int8_t TAG_TX_POWER_QDBM = 50;

void loop() {
  if (BURST_ALL_CHANNELS) {
    for (size_t idx = 0; idx < num_channels; idx++) {
      send_probe_on_channel(TAG_CHANNELS[idx]);
      delay(INTER_CHANNEL_GAP_MS);
    }
  } else {
    send_probe_on_channel(TAG_CHANNEL);
  }
}"""

NEW_TAG_CODE = """static const uint8_t TAG_CHANNEL = WIFI_AP_CHANNEL;
static const uint32_t TX_INTERVAL_MS = 90;
static const uint8_t BURST_PROBES_PER_CYCLE = 3;
static const uint32_t BURST_GAP_MS = 12;
static const int8_t TAG_TX_POWER_QDBM = 76;

void loop() {
  if (now_ms - last_tx_ms < TX_INTERVAL_MS) return;
  last_tx_ms = now_ms;
  for (uint8_t i = 0; i < BURST_PROBES_PER_CYCLE; i++) {
    send_probe();
    if (i + 1 < BURST_PROBES_PER_CYCLE) {
      delay(BURST_GAP_MS);
    }
  }
}"""

LIVE_PADDING_CODE = """def _build_feature_row_from_packets(..., pad_missing: bool = True):
    ...
    if len(rssis) < samples_per_anchor:
        if not pad_missing:
            return pd.DataFrame(), counts
        rssis = ([-100.0] * (samples_per_anchor - len(rssis))) + rssis

strict_full_window = st.checkbox(
    "Exigir ventana completa por antena (sin relleno artificial)",
    value=True,
)
...
pad_missing=not bool(strict_full_window)"""


def ensure_dirs() -> None:
    GENERATED_DIR.mkdir(parents=True, exist_ok=True)
    PENDING_PHOTO_DIR.mkdir(parents=True, exist_ok=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph(text, style=f"Heading {level}")
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.bold = True


def add_body_paragraph(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)


def add_table_title(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True


def add_page_break(doc: Document) -> None:
    doc.add_page_break()


def set_cell_text(cell, text: object, bold: bool = False, font_name: str = "Times New Roman", size: int = 10) -> None:
    cell.text = str(text)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(size)
            run.bold = bold


def shade_cell(cell, fill: str = "F2F2F2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, headers: list[str], rows: Iterable[Iterable[object]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph("")


def add_caption(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.italic = True


def add_picture_with_caption(doc: Document, image_path: Path, caption: str, width: float = 6.2) -> None:
    doc.add_picture(str(image_path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_caption(doc, caption)


def add_code_block(doc: Document, title: str, code_text: str, explanation: str) -> None:
    title_par = doc.add_paragraph(title)
    for run in title_par.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(10)
        run.bold = True

    table = doc.add_table(rows=1, cols=1)
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    cell = table.rows[0].cells[0]
    shade_cell(cell, "F7F7F7")
    cell.text = code_text
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        for run in paragraph.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8.5)
    add_body_paragraph(doc, explanation)


def create_placeholder_image(title: str, subtitle: str, output_path: Path) -> Path:
    image = Image.new("RGB", (1600, 900), (242, 242, 242))
    draw = ImageDraw.Draw(image)
    font = ImageFont.load_default()
    draw.rectangle((40, 40, 1560, 860), outline=(120, 120, 120), width=4)
    draw.multiline_text((120, 260), title, fill=(20, 20, 20), font=font, spacing=12)
    draw.multiline_text((120, 430), subtitle, fill=(70, 70, 70), font=font, spacing=10)
    image.save(output_path)
    return output_path


def create_contact_sheet(paths: list[Path], output_path: Path, cols: int = 2) -> Path:
    if not paths:
        raise ValueError("No hay imagenes para crear la hoja de contacto")
    tile_w, tile_h = 460, 280
    rows = (len(paths) + cols - 1) // cols
    canvas = Image.new("RGB", (cols * tile_w + 80, rows * (tile_h + 60) + 80), (255, 255, 255))
    draw = ImageDraw.Draw(canvas)
    font = ImageFont.load_default()
    for index, path in enumerate(paths):
        row = index // cols
        col = index % cols
        x = 40 + col * tile_w
        y = 40 + row * (tile_h + 60)
        with Image.open(path) as image:
            image = image.convert("RGB")
            image.thumbnail((tile_w - 20, tile_h - 20))
            tile = Image.new("RGB", (tile_w - 20, tile_h - 20), (248, 249, 250))
            ox = (tile.width - image.width) // 2
            oy = (tile.height - image.height) // 2
            tile.paste(image, (ox, oy))
            canvas.paste(tile, (x + 10, y + 10))
        draw.rectangle((x, y, x + tile_w - 1, y + tile_h - 1), outline=(185, 185, 185), width=2)
        draw.text((x, y + tile_h + 10), path.stem.replace("Captura de pantalla ", ""), fill=(20, 20, 20), font=font)
    canvas.save(output_path)
    return output_path


def plot_capture_durations(points: pd.DataFrame, title: str, output_path: Path) -> Path:
    if points.empty:
        return output_path
    frame = points.copy()
    frame["started_at"] = pd.to_datetime(frame["started_at"])
    frame["capture_complete_at"] = pd.to_datetime(frame["capture_complete_at"])
    frame["completed_at"] = pd.to_datetime(frame["completed_at"])
    frame["capture_duration_s"] = (frame["capture_complete_at"] - frame["started_at"]).dt.total_seconds()

    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.bar(frame["point_id"], frame["capture_duration_s"], color="#0a9396")
    ax.set_title(title)
    ax.set_ylabel("Tiempo de captura [s]")
    ax.set_xlabel("Punto")
    ax.tick_params(axis="x", rotation=45)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(output_path, dpi=200)
    plt.close(fig)
    return output_path


def plot_metrics(metrics: pd.DataFrame, title: str, output_path: Path) -> Path:
    if metrics.empty:
        return output_path
    frame = metrics.sort_values("mae_eucl").copy()
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.bar(frame["model"], frame["mae_eucl"], color="#005f73", label="MAE")
    ax.plot(frame["model"], frame["p95"], color="#bb3e03", marker="o", linewidth=2, label="P95")
    ax.set_title(title)
    ax.set_ylabel("Error [m]")
    ax.set_xlabel("Modelo")
    ax.grid(axis="y", alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(output_path, dpi=200)
    plt.close(fig)
    return output_path


def plot_raw_sizes(raw_sizes: dict[str, int], title: str, output_path: Path) -> Path:
    if not raw_sizes:
        return output_path
    labels = list(raw_sizes.keys())
    values = [raw_sizes[label] / 1024.0 for label in labels]
    fig, ax = plt.subplots(figsize=(8, 4.5))
    ax.bar(labels, values, color="#ee9b00")
    ax.set_title(title)
    ax.set_ylabel("Tamano RAW [KiB]")
    ax.set_xlabel("Antena ESP32")
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(output_path, dpi=200)
    plt.close(fig)
    return output_path


def plot_campaign_compare(rows: list[dict[str, object]], output_path: Path) -> Path:
    frame = pd.DataFrame(rows)
    fig, ax = plt.subplots(figsize=(10, 4.8))
    ax.bar(frame["label"], frame["mae"], color="#0a9396", label="MAE")
    ax.plot(frame["label"], frame["p95"], color="#9b2226", marker="o", linewidth=2, label="P95")
    ax.set_title("Comparacion de campanas con dataset y modelo")
    ax.set_ylabel("Error [m]")
    ax.set_xlabel("Campana")
    ax.tick_params(axis="x", rotation=20)
    ax.grid(axis="y", alpha=0.25)
    ax.legend()
    fig.tight_layout()
    fig.savefig(output_path, dpi=200)
    plt.close(fig)
    return output_path


def load_run_summary(run_dir: Path) -> dict:
    summary: dict[str, object] = {"run_dir": run_dir}
    if not run_dir.exists():
        return summary

    experiment = json.loads((run_dir / "experiment.json").read_text(encoding="utf-8"))
    summary["experiment"] = experiment

    state_path = run_dir / "training_state.json"
    summary["state"] = json.loads(state_path.read_text(encoding="utf-8")) if state_path.exists() else {}

    points_path = run_dir / "points.csv"
    if points_path.exists():
        points = pd.read_csv(points_path)
        if {"started_at", "capture_complete_at", "completed_at"}.issubset(points.columns):
            points["started_at"] = pd.to_datetime(points["started_at"])
            points["capture_complete_at"] = pd.to_datetime(points["capture_complete_at"])
            points["completed_at"] = pd.to_datetime(points["completed_at"])
            points["capture_duration_s"] = (points["capture_complete_at"] - points["started_at"]).dt.total_seconds()
            points["close_duration_s"] = (points["completed_at"] - points["capture_complete_at"]).dt.total_seconds()
            points["total_duration_s"] = (points["completed_at"] - points["started_at"]).dt.total_seconds()
        summary["points"] = points
    else:
        summary["points"] = pd.DataFrame()

    samples_path = run_dir / "samples.csv"
    if samples_path.exists():
        samples = pd.read_csv(samples_path)
        summary["samples"] = samples
        summary["sample_counts"] = samples.groupby("anchor_id").size().to_dict()
        summary["rssi_stats"] = (
            samples.groupby("anchor_id")["rssi"]
            .agg(["count", "min", "max", "mean", "std"])
            .reset_index()
            .rename(columns={"count": "n"})
        )
    else:
        summary["samples"] = pd.DataFrame()
        summary["sample_counts"] = {}
        summary["rssi_stats"] = pd.DataFrame()

    dataset_path = run_dir / "dataset.csv"
    summary["dataset"] = pd.read_csv(dataset_path) if dataset_path.exists() else pd.DataFrame()

    metrics_path = run_dir / "models" / "metrics.csv"
    if metrics_path.exists():
        metrics = pd.read_csv(metrics_path)
        summary["metrics"] = metrics
        summary["best_mae"] = metrics.sort_values("mae_eucl").iloc[0].to_dict()
        summary["best_p95"] = metrics.sort_values("p95").iloc[0].to_dict()
    else:
        summary["metrics"] = pd.DataFrame()
        summary["best_mae"] = {}
        summary["best_p95"] = {}

    raw_dir = run_dir / "raw"
    summary["raw_sizes"] = {path.stem: path.stat().st_size for path in sorted(raw_dir.glob("*.jsonl"))} if raw_dir.exists() else {}
    return summary


def remove_body_from_heading_to_end(doc: Document, heading_prefix: str) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_para = next(p for p in doc.paragraphs if p.text.strip().startswith(heading_prefix))
    start_idx = children.index(start_para._element)
    for child in children[start_idx:-1]:
        body.remove(child)


def selected_screenshots(date_prefix: str, chosen_names: list[str]) -> list[Path]:
    available = {path.name: path for path in SCREENSHOT_DIR.glob(f"Captura de pantalla {date_prefix}*.png")}
    return [available[name] for name in chosen_names if name in available]


def recent_phone_photos() -> list[Path]:
    candidates: list[Path] = []
    roots = [
        PENDING_PHOTO_DIR,
        Path("C:/Users/hi_iv/OneDrive/Im\u00e1genes/Galer\u00eda de Samsung/DCIM/Camera"),
        Path("C:/Users/hi_iv/OneDrive/Im\u00e1genes/\u00c1lbum de c\u00e1mara"),
    ]
    for root in roots:
        if not root.exists():
            continue
        for path in root.rglob("*"):
            if path.is_file() and path.suffix.lower() in {".jpg", ".jpeg", ".png", ".heic"}:
                if path.stat().st_mtime >= pd.Timestamp("2026-03-20").timestamp():
                    candidates.append(path)
    return sorted(candidates, key=lambda path: path.stat().st_mtime, reverse=True)


def build_figures(run_s2: dict, run_s4: dict, run_s5: dict) -> dict[str, Path]:
    return {
        "s5_capture": plot_capture_durations(
            run_s5["points"],
            "Tiempos de captura por punto - demo_s5 / train_05",
            GENERATED_DIR / "fig_run_5_capture_times.png",
        ),
        "s5_metrics": plot_metrics(
            run_s5["metrics"],
            "Metricas de modelos - demo_s5 / train_05",
            GENERATED_DIR / "fig_run_5_metrics.png",
        ),
        "s5_raw": plot_raw_sizes(
            run_s5["raw_sizes"],
            "Tamano de archivos RAW por antena - demo_s5 / train_05",
            GENERATED_DIR / "fig_run_5_raw_sizes.png",
        ),
        "campaign_compare": plot_campaign_compare(
            [
                {
                    "label": "demo_s2",
                    "mae": float(run_s2["best_mae"]["mae_eucl"]),
                    "p95": float(run_s2["best_p95"]["p95"]),
                },
                {
                    "label": "demo_s4",
                    "mae": float(run_s4["best_mae"]["mae_eucl"]),
                    "p95": float(run_s4["best_p95"]["p95"]),
                },
                {
                    "label": "demo_s5",
                    "mae": float(run_s5["best_mae"]["mae_eucl"]),
                    "p95": float(run_s5["best_p95"]["p95"]),
                },
            ],
            GENERATED_DIR / "fig_campaign_compare_final.png",
        ),
    }


def fmt_float(value: object, digits: int = 2) -> str:
    if value is None:
        return "-"
    try:
        if pd.isna(value):
            return "-"
    except TypeError:
        pass
    return f"{float(value):.{digits}f}"


def fmt_seconds(value: object) -> str:
    if value is None:
        return "-"
    try:
        if pd.isna(value):
            return "-"
    except TypeError:
        pass
    total = float(value)
    if total >= 60.0:
        minutes = int(total // 60.0)
        seconds = total - minutes * 60.0
        return f"{minutes} min {seconds:.1f} s"
    return f"{total:.1f} s"


def add_table_from_dataframe(doc: Document, title: str, frame: pd.DataFrame) -> None:
    if frame.empty:
        return
    add_table_title(doc, title)
    headers = [str(column) for column in frame.columns]
    rows = frame.astype(object).values.tolist()
    add_table(doc, headers, rows)


def build_visual_assets() -> dict[str, Path]:
    assets: dict[str, Path] = {}

    screenshots_s4 = selected_screenshots(
        "2026-03-23",
        [
            "Captura de pantalla 2026-03-23 163526.png",
            "Captura de pantalla 2026-03-23 180641.png",
            "Captura de pantalla 2026-03-23 181743.png",
            "Captura de pantalla 2026-03-23 184258.png",
        ],
    )
    if screenshots_s4:
        assets["sheet_s4"] = create_contact_sheet(
            screenshots_s4,
            GENERATED_DIR / "contacto_campana4_2026_03_23.png",
            cols=2,
        )
        assets["s4_single_1"] = screenshots_s4[0]
        assets["s4_single_2"] = screenshots_s4[-1]

    screenshots_s5 = selected_screenshots(
        "2026-03-29",
        [
            "Captura de pantalla 2026-03-29 181318.png",
            "Captura de pantalla 2026-03-29 182253.png",
            "Captura de pantalla 2026-03-29 182425.png",
            "Captura de pantalla 2026-03-29 184015.png",
        ],
    )
    if screenshots_s5:
        assets["sheet_s5"] = create_contact_sheet(
            screenshots_s5,
            GENERATED_DIR / "contacto_campana5_2026_03_29.png",
            cols=2,
        )
        assets["s5_single_1"] = screenshots_s5[0]
        assets["s5_single_2"] = screenshots_s5[-1]

    screenshots_final = selected_screenshots(
        "2026-04-03",
        [
            "Captura de pantalla 2026-04-03 093940.png",
            "Captura de pantalla 2026-04-03 094228.png",
            "Captura de pantalla 2026-04-03 094417.png",
            "Captura de pantalla 2026-04-03 095556.png",
        ],
    )
    if screenshots_final:
        assets["sheet_final"] = create_contact_sheet(
            screenshots_final,
            GENERATED_DIR / "contacto_inferencia_final_2026_04_03.png",
            cols=2,
        )
        assets["final_single_1"] = screenshots_final[0]
        assets["final_single_2"] = screenshots_final[-1]

    router_screens = selected_screenshots(
        "2026-03-26",
        [
            "Captura de pantalla 2026-03-26 214636.png",
            "Captura de pantalla 2026-03-26 214651.png",
            "Captura de pantalla 2026-03-26 214716.png",
            "Captura de pantalla 2026-03-26 214728.png",
        ],
    )
    if router_screens:
        assets["sheet_router"] = create_contact_sheet(
            router_screens,
            GENERATED_DIR / "contacto_router_2026_03_26.png",
            cols=2,
        )

    board_screen = SCREENSHOT_DIR / "Captura de pantalla 2026-04-03 085858.png"
    if board_screen.exists():
        assets["board_selection"] = board_screen

    phone_photos = recent_phone_photos()
    if phone_photos:
        assets["sheet_phone"] = create_contact_sheet(
            phone_photos[:4],
            GENERATED_DIR / "contacto_fotos_campo_recientes.png",
            cols=2,
        )

    assets["placeholder_montaje"] = create_placeholder_image(
        "Espacio reservado para fotografia del montaje fisico de la campana final",
        "Insertar aqui una fotografia tomada con telefono movil donde se observe el recinto, la ubicacion de las cuatro antenas ESP32, el servidor y el recorrido experimental utilizado durante demo_s5.",
        GENERATED_DIR / "placeholder_final_montaje_campana5.png",
    )
    assets["placeholder_tag"] = create_placeholder_image(
        "Espacio reservado para fotografia del TAG durante la inferencia en linea",
        "Insertar aqui una fotografia donde se observe la posicion real del emisor objetivo durante la validacion final posterior a la recuperacion de la placa, para contrastar la ubicacion estimada en la interfaz.",
        GENERATED_DIR / "placeholder_final_tag_inferencia.png",
    )
    assets["placeholder_point"] = create_placeholder_image(
        "Espacio reservado para fotografia del punto de entrenamiento",
        "Insertar aqui una fotografia del TAG ubicado sobre un punto marcado del plano, con identificacion visible del punto de la grilla y de la distribucion de antenas ESP32.",
        GENERATED_DIR / "placeholder_final_punto_entrenamiento.png",
    )
    return assets


def add_section_24(doc: Document, run_s1: dict, run_s2: dict, run_s3: dict, run_s4: dict, run_s5: dict, figures: dict[str, Path], assets: dict[str, Path]) -> None:
    inventory_rows = [
        ["demo_s1 / train_01", "Piloto inicial con tres antenas ESP32", len(run_s1["experiment"]["anchors"]), run_s1["experiment"]["samples_per_anchor"], len(run_s1["points"]), len(run_s1["samples"]), "-", "-", "-"],
        ["demo_s2 / train_02", "Campana principal con tres antenas ESP32", len(run_s2["experiment"]["anchors"]), run_s2["experiment"]["samples_per_anchor"], len(run_s2["points"]), len(run_s2["samples"]), run_s2["best_mae"]["model"], fmt_float(run_s2["best_mae"]["mae_eucl"], 3), fmt_float(run_s2["best_p95"]["p95"], 3)],
        ["demo_s3 / train_03", "Corrida de transicion sin dataset utilizable", len(run_s3["experiment"]["anchors"]), run_s3["experiment"]["samples_per_anchor"], 0, 0, "-", "-", "-"],
        ["demo_s4 / train_04", "Primera campana completa con cuatro antenas ESP32", len(run_s4["experiment"]["anchors"]), run_s4["experiment"]["samples_per_anchor"], len(run_s4["points"]), len(run_s4["samples"]), run_s4["best_mae"]["model"], fmt_float(run_s4["best_mae"]["mae_eucl"], 3), fmt_float(run_s4["best_p95"]["p95"], 3)],
        ["demo_s5 / train_05", "Campana final con firmware de canal fijo y cincuenta muestras por antena", len(run_s5["experiment"]["anchors"]), run_s5["experiment"]["samples_per_anchor"], len(run_s5["points"]), len(run_s5["samples"]), run_s5["best_mae"]["model"], fmt_float(run_s5["best_mae"]["mae_eucl"], 3), fmt_float(run_s5["best_p95"]["p95"], 3)],
    ]

    add_heading(doc, "24. Cierre experimental consolidado y campañas ejecutadas", level=1)
    add_body_paragraph(doc, "La etapa final del proyecto se estructuro como un proceso iterativo de maduracion experimental. A diferencia de una campana unica y lineal, el desarrollo practico requirio una secuencia de corridas sucesivas en las que se ajustaron la geometria de las antenas ESP32, la estabilidad de la asociacion Wi-Fi, la logica de captura, la densidad de muestras por punto y la consistencia entre el firmware del emisor objetivo y la inferencia en linea. En consecuencia, el cierre experimental debe leerse como una consolidacion de cinco campanas, cada una con un objetivo distinto y con un aporte especifico a la validacion del sistema.")
    add_body_paragraph(doc, "La Tabla 24.1 resume el inventario completo de campanas ejecutadas y explicita el rol metodologico de cada una. Esta sintesis permite identificar con claridad cuales corridas tuvieron caracter exploratorio, cuales generaron un dataset utilizable para entrenamiento y cuales sirvieron para estabilizar el funcionamiento del sistema de localizacion en interiores.")
    add_table_title(doc, "Tabla 24.1. Inventario consolidado de campanas experimentales ejecutadas")
    add_table(doc, ["Campana", "Rol experimental", "Antenas", "Muestras por antena", "Puntos", "Muestras aceptadas", "Mejor modelo", "MAE [m]", "P95 [m]"], inventory_rows)

    add_heading(doc, "24.1 Campana piloto inicial con tres antenas ESP32: demo_s1 / train_01", level=2)
    add_body_paragraph(doc, "La corrida demo_s1 / train_01 constituyo la primera validacion operativa de la instrumentacion. Se trabajo con tres antenas ESP32, veinte muestras por antena y solo dos puntos completados. El objetivo de esta etapa no fue entrenar un modelo definitivo, sino verificar que el circuito experimental pudiera abrir un punto, recibir paquetes desde varias antenas, consolidar muestras por identificador de punto y almacenar los archivos RAW correspondientes.")
    add_table_title(doc, "Tabla 24.2. Resumen cuantitativo de la campana demo_s1 / train_01")
    add_table(doc, ["Parametro", "Valor"], [["Antenas ESP32 utilizadas", len(run_s1["experiment"]["anchors"])], ["Muestras por antena", run_s1["experiment"]["samples_per_anchor"]], ["Puntos completos", len(run_s1["points"])], ["Muestras aceptadas", len(run_s1["samples"])], ["Distribucion por antena", ", ".join(f"{key}: {value}" for key, value in run_s1["sample_counts"].items())], ["Dataset consolidado", "No generado"], ["Metricas de modelo", "No disponibles"]])
    add_body_paragraph(doc, "El valor de esta campana residio en poner en evidencia la viabilidad del pipeline extremo a extremo y, al mismo tiempo, la insuficiencia del volumen muestral para sostener un entrenamiento reproducible. La ausencia de dataset.csv y metrics.csv no debe interpretarse como una omision, sino como un hallazgo metodologico: la cantidad de puntos relevados aun no representaba de manera razonable el recinto.")

    add_heading(doc, "24.2 Campana principal con tres antenas ESP32: demo_s2 / train_02", level=2)
    add_body_paragraph(doc, "La campana demo_s2 / train_02 fue la primera corrida completa con resultados de entrenamiento plenamente utilizables. Se relevaron dieciocho puntos de la grilla, con diez muestras por antena y tres antenas ESP32 en funcionamiento estable. El dataset consolidado resultante tuvo dieciocho filas y cincuenta y ocho columnas, lo que permitio comparar cinco modelos de regresion distintos mediante validacion GroupKFold.")
    add_table_title(doc, "Tabla 24.3. Metricas completas de la campana demo_s2 / train_02")
    add_table(doc, ["Modelo", "MAE [m]", "RMSE [m]", "P95 [m]"], [[row["model"], fmt_float(row["mae_eucl"], 3), fmt_float(row["rmse_eucl"], 3), fmt_float(row["p95"], 3)] for _, row in run_s2["metrics"].sort_values("mae_eucl").iterrows()])
    add_body_paragraph(doc, "En esta campana el mejor modelo fue Extra Trees con quinientos estimadores, con un MAE de 1,263 m y un percentil 95 de 2,375 m. Este resultado establecio la primera linea de base cuantitativa del proyecto. La utilidad practica de la campana demo_s2 consistio en demostrar que el enfoque de fingerprinting RSSI podia producir estimaciones consistentes, aunque todavia con una geometria de observacion incompleta al trabajar solo con tres antenas ESP32.")

    add_heading(doc, "24.3 Corrida de transicion sin dataset utilizable: demo_s3 / train_03", level=2)
    add_body_paragraph(doc, "La corrida demo_s3 / train_03 opero como una etapa de transicion entre la plataforma de tres antenas y la posterior migracion a cuatro antenas ESP32. Si bien el experimento fue formalmente inicializado con diez muestras por antena, no se completaron puntos utiles y no se genero un dataset consolidado. Desde el punto de vista ingenieril, esta corrida sirvio para detectar problemas de continuidad operativa y para justificar la reformulacion del firmware y de la metodologia de captura antes de avanzar a la siguiente campana.")

    add_page_break(doc)
    add_heading(doc, "24.4 Primera campana completa con cuatro antenas ESP32: demo_s4 / train_04", level=2)
    add_body_paragraph(doc, "La campana demo_s4 / train_04 incorporo una cuarta antena ESP32 y, con ello, una geometria mas cerrada del recinto. Esta decision respondio directamente a una limitacion observada en las corridas previas: la falta de simetria espacial tendia a degradar la discriminacion direccional del RSSI. Se mantuvieron diez muestras por antena y se completaron dieciocho puntos, con setecientas veinte muestras aceptadas y un dataset de dieciocho filas por setenta y cinco columnas.")
    add_table_title(doc, "Tabla 24.4. Coordenadas de las cuatro antenas ESP32 utilizadas en demo_s4 / train_04")
    add_table(doc, ["Antena", "x [m]", "y [m]", "z [m]"], [[anchor["anchor_id"], fmt_float(anchor["x_m"]), fmt_float(anchor["y_m"]), fmt_float(anchor["z_m"])] for anchor in run_s4["experiment"]["anchors"]])
    add_table_title(doc, "Tabla 24.5. Metricas completas de la campana demo_s4 / train_04")
    add_table(doc, ["Modelo", "MAE [m]", "RMSE [m]", "P95 [m]"], [[row["model"], fmt_float(row["mae_eucl"], 3), fmt_float(row["rmse_eucl"], 3), fmt_float(row["p95"], 3)] for _, row in run_s4["metrics"].sort_values("mae_eucl").iterrows()])
    add_body_paragraph(doc, "El mejor modelo de esta campana volvio a ser Extra Trees con quinientos estimadores, pero el error medio absoluto ascendio a 1,512 m y el P95 a 3,185 m. Este resultado mostro que incorporar una cuarta antena no garantiza por si mismo una mejora inmediata del modelo. La campana fue metodologicamente valiosa porque expuso que la calidad del aprendizaje depende, ademas de la geometria, de la estabilidad del firmware, de la coherencia del emisor objetivo y de la igualdad efectiva de condiciones de captura entre antenas.")
    if "sheet_s4" in assets:
        add_picture_with_caption(doc, assets["sheet_s4"], "Figura 24.1. Capturas de la interfaz correspondientes a la primera campana completa con cuatro antenas ESP32 (demo_s4 / train_04). Se observan el plano, la localizacion estimada y la evolucion del mapa de calor durante las pruebas iniciales de inferencia en linea.", width=6.3)
    if "s4_single_2" in assets:
        add_picture_with_caption(doc, assets["s4_single_2"], "Figura 24.2. Ejemplo de visualizacion puntual de demo_s4. La imagen ilustra el comportamiento espacial del mapa de calor durante la primera fase de validacion con cuatro antenas ESP32.", width=5.9)

    add_heading(doc, "24.5 Campana final con firmware de canal fijo y cincuenta muestras por antena: demo_s5 / train_05", level=2)
    add_body_paragraph(doc, "La campana demo_s5 / train_05 representa la instancia experimental final y la base de los resultados de cierre del proyecto. En ella se relevaron veintiun puntos, con cuatro antenas ESP32, cincuenta muestras por antena y un firmware de entrenamiento reformulado para operar siempre sobre el canal fijo del punto de acceso. Esta campana incorporo, ademas, la version revisada del emisor objetivo, de manera que la adquisicion de datos y la futura inferencia compartieran la misma logica de transmision sobre el canal uno de la banda de 2,4 GHz.")
    add_table_title(doc, "Tabla 24.6. Coordenadas de las cuatro antenas ESP32 utilizadas en demo_s5 / train_05")
    add_table(doc, ["Antena", "x [m]", "y [m]", "z [m]"], [[anchor["anchor_id"], fmt_float(anchor["x_m"]), fmt_float(anchor["y_m"]), fmt_float(anchor["z_m"])] for anchor in run_s5["experiment"]["anchors"]])
    add_table_title(doc, "Tabla 24.7. Resumen cuantitativo de demo_s5 / train_05")
    add_table(doc, ["Parametro", "Valor"], [["Puntos completos", len(run_s5["points"])], ["Muestras por antena", run_s5["experiment"]["samples_per_anchor"]], ["Muestras aceptadas", len(run_s5["samples"])], ["Distribucion por antena", ", ".join(f"{key}: {value}" for key, value in run_s5["sample_counts"].items())], ["Dimensiones de dataset.csv", f"{run_s5['dataset'].shape[0]} filas x {run_s5['dataset'].shape[1]} columnas"], ["Mejor modelo por MAE", run_s5["best_mae"]["model"]], ["MAE minimo", fmt_float(run_s5["best_mae"]["mae_eucl"], 3)], ["P95 minimo", fmt_float(run_s5["best_p95"]["p95"], 3)]])
    add_table_title(doc, "Tabla 24.8. Metricas completas de la campana demo_s5 / train_05")
    add_table(doc, ["Modelo", "MAE [m]", "RMSE [m]", "P50 [m]", "P90 [m]", "P95 [m]"], [[row["model"], fmt_float(row["mae_eucl"], 3), fmt_float(row["rmse_eucl"], 3), fmt_float(row["p50"], 3), fmt_float(row["p90"], 3), fmt_float(row["p95"], 3)] for _, row in run_s5["metrics"].sort_values("mae_eucl").iterrows()])
    add_table_title(doc, "Tabla 24.9. Tiempos reales de captura por punto en demo_s5 / train_05")
    add_table(doc, ["Punto", "x [m]", "y [m]", "Captura", "Duracion total"], [[row["point_id"], fmt_float(row["x_m"]), fmt_float(row["y_m"]), fmt_seconds(row["capture_duration_s"]), fmt_seconds(row["total_duration_s"])] for _, row in run_s5["points"].iterrows()])
    add_body_paragraph(doc, "La campana final produjo un dataset significativamente mas denso que los anteriores, con 4200 muestras aceptadas y 235 columnas de atributos. El mejor resultado correspondio al modelo Random Forest con trescientos estimadores, con un MAE de 1,670 m, RMSE de 1,858 m y P95 de 2,756 m. Aunque este MAE no mejoro al obtenido en demo_s2, la corrida final fue la mas representativa desde el punto de vista operativo, porque incorporo cuatro antenas, mayor volumen muestral, estabilidad de canal y una trazabilidad experimental mucho mas completa.")
    add_body_paragraph(doc, "Las duraciones de captura por punto muestran que el proceso no fue homogeneo. Puntos como P05, P06, P07, P10, P12 y P17 registraron tiempos extraordinariamente altos, asociados a problemas reales de continuidad del emisor objetivo, recuperacion de placas y ajuste fino de la conectividad. Esta dispersion temporal no invalida la campana; por el contrario, documenta de forma transparente las restricciones de una implementacion real y permite justificar las correcciones incorporadas en la fase final.")
    if "s5_capture" in figures:
        add_picture_with_caption(doc, figures["s5_capture"], "Figura 24.3. Tiempos de captura por punto en la campana demo_s5 / train_05. La dispersion temporal evidencia que los puntos no tuvieron igual dificultad operativa y que la estabilidad del sistema todavia dependio de la continuidad de transmision y recepcion.", width=6.1)
    if "s5_metrics" in figures:
        add_picture_with_caption(doc, figures["s5_metrics"], "Figura 24.4. Comparacion de modelos para demo_s5 / train_05. Se grafican el error medio absoluto y el percentil 95 obtenidos sobre el dataset consolidado final.", width=6.1)
    if "s5_raw" in figures:
        add_picture_with_caption(doc, figures["s5_raw"], "Figura 24.5. Tamano de archivos RAW por antena ESP32 en demo_s5 / train_05. La uniformidad relativa del volumen por antena es coherente con la exigencia de cincuenta muestras por punto y con el cierre correcto de los veintiun puntos relevados.", width=5.8)
    if "sheet_s5" in assets:
        add_picture_with_caption(doc, assets["sheet_s5"], "Figura 24.6. Capturas de la interfaz correspondientes a demo_s5 / train_05. Se incluyen ejemplos de la localizacion estimada y del mapa de calor obtenido durante la fase final de entrenamiento y verificacion.", width=6.3)

    add_page_break(doc)
    add_heading(doc, "24.6 Validacion final de inferencia en linea posterior a la recuperacion de la placa", level=2)
    add_body_paragraph(doc, "La fase de cierre practico no se limito al entrenamiento. Una vez completada la campana demo_s5, se ejecuto una validacion final de inferencia en linea con el firmware continuo de las antenas ESP32 y con la placa recuperada mediante reset del pin EN. Esta instancia fue especialmente relevante porque permitio verificar, ya sobre la plataforma estabilizada, que las cuatro antenas continuaran consultando configuracion, capturando paquetes utiles y reportando lotes al servidor con continuidad suficiente para sostener la inferencia del emisor objetivo.")
    add_body_paragraph(doc, "En el estado final registrado por el sistema, las cuatro antenas ESP32 reportaron actividad valida en la campana demo_s5 / train_05. Los ultimos lotes observados en training_state.json correspondieron a cuarenta paquetes para A1, treinta y ocho para A2, treinta y tres para A3 y cuarenta y dos para A4, con tiempos de ultimo reporte concentrados en la manana del 3 de abril de 2026. Este comportamiento constituye la evidencia mas clara de que, luego de los ajustes de firmware y de hardware, el sistema alcanzo una operacion simultanea de las cuatro antenas sobre el canal fijo del punto de acceso.")
    add_table_title(doc, "Tabla 24.10. Ultimo estado registrado de las antenas ESP32 durante la validacion final de inferencia")
    add_table(doc, ["Antena", "Ultima IP", "Ultimo lote", "Ultimo reporte"], [[anchor_id, data.get("last_ip", "-"), data.get("last_batch_packets", "-"), data.get("last_batch_at", "-")] for anchor_id, data in run_s5["state"].get("anchor_statuses", {}).items()])
    if "sheet_final" in assets:
        add_picture_with_caption(doc, assets["sheet_final"], "Figura 24.7. Registro visual de la inferencia final en linea posterior a la recuperacion de la placa. Las capturas del 3 de abril de 2026 muestran la interfaz operando sobre la campana demo_s5 y reflejan el comportamiento del mapa de calor una vez estabilizadas las cuatro antenas ESP32.", width=6.3)
    if "final_single_2" in assets:
        add_picture_with_caption(doc, assets["final_single_2"], "Figura 24.8. Vista individual de la inferencia final. La figura ilustra el uso de la campana demo_s5 como referencia para la estimacion de posicion del emisor objetivo en tiempo real.", width=5.9)

    add_heading(doc, "24.7 Registro visual complementario y espacios reservados para fotografias de campo", level=2)
    add_body_paragraph(doc, "Con el fin de dejar el documento listo para entrega formal, se incorporaron capturas efectivamente generadas durante las campanas y se reservaron espacios explicitamente rotulados para las fotografias tomadas con telefono movil. Al momento de reconstruir el presente documento, las capturas de pantalla estaban disponibles en el arbol de trabajo, mientras que las fotografias de campo recientes no se encontraban aun ubicadas de manera automatica en una carpeta accesible. Por esa razon, se dejaron espacios formales de insercion para que puedan ser reemplazados sin alterar la estructura del informe.")
    if "sheet_phone" in assets:
        add_picture_with_caption(doc, assets["sheet_phone"], "Figura 24.9. Contacto de fotografias de campo correspondientes a la campana final y a la validacion practica del emisor objetivo.", width=6.3)
    else:
        add_picture_with_caption(doc, assets["placeholder_montaje"], "Figura 24.9. Espacio reservado para fotografia general del montaje fisico de la campana final con cuatro antenas ESP32.", width=6.1)
        add_picture_with_caption(doc, assets["placeholder_tag"], "Figura 24.10. Espacio reservado para fotografia del emisor objetivo durante la inferencia en linea.", width=6.1)
        add_picture_with_caption(doc, assets["placeholder_point"], "Figura 24.11. Espacio reservado para fotografia del TAG ubicado sobre un punto de entrenamiento de la grilla experimental.", width=6.1)


def add_section_25(doc: Document, assets: dict[str, Path]) -> None:
    add_heading(doc, "25. Explicacion tecnica del codigo implementado", level=1)
    add_body_paragraph(doc, "La evolucion del proyecto no estuvo determinada unicamente por la cantidad de antenas ESP32 o por el algoritmo de aprendizaje automatico. Una parte sustantiva del resultado final dependio del rediseño progresivo del firmware y del software de soporte. A continuacion se describen los bloques mas relevantes, no en forma de inventario superficial, sino como una secuencia de decisiones de ingenieria que modificaron directamente la calidad de la captura, la continuidad operativa y la representatividad del dataset.")

    add_heading(doc, "25.1 Firmware original de antenas ESP32 para captura por puntos", level=2)
    add_body_paragraph(doc, "El firmware original de las antenas ESP32 utilizaba una estrategia de alternancia entre modo promiscuo y asociacion STA. Durante la captura, la placa escuchaba tramas de gestion en modo promiscuo; luego abandonaba temporalmente la captura, se asociaba al punto de acceso, consultaba la configuracion vigente y enviaba un lote HTTP al servidor. El siguiente fragmento resume esa logica.")
    add_code_block(doc, "Fragmento 25.1. Cambio de configuracion en el firmware original de antenas", OLD_ANTENNA_CODE, "Desde el punto de vista funcional, este bloque implementa la idea de utilizar el canal del punto de acceso como pista de configuracion, sin fijarlo de manera permanente. El enfoque fue util para las primeras pruebas, pero presento una debilidad evidente: la placa debia salir del modo de captura para recuperar IP y transmitir el lote, lo que introducia huecos temporales y hacia mas sensible el sistema a variaciones de conectividad y de roaming del punto de acceso.")
    add_body_paragraph(doc, "La consecuencia metodologica de esta primera version fue doble. Por un lado, permitio construir las primeras campanas con tres antenas ESP32. Por otro, revelo la necesidad de desacoplar, en la medida de lo posible, la captura del trafico del procedimiento de reconexion a la red Wi-Fi. Esta observacion se volvio central cuando el proyecto migro hacia cuatro antenas y hacia la inferencia en tiempo real.")

    add_heading(doc, "25.2 Firmware nuevo de entrenamiento con canal fijo y asociacion estable", level=2)
    add_body_paragraph(doc, "La siguiente evolucion consistio en un firmware especifico para entrenamiento, orientado a capturar solamente durante un punto abierto de la campana y a mantener la asociacion a la red sobre el canal fijo del punto de acceso. Para ello se incorporo bloqueo por BSSID, seleccion de canal fijo y una funcion explicita que habilita la captura solo si la campaña se encuentra en modo capture_active.")
    add_code_block(doc, "Fragmento 25.2. Firmware de entrenamiento con canal fijo y BSSID bloqueado", NEW_TRAINING_CODE, "En este diseño, should_capture_now() actua como compuerta de seguridad experimental. La antena ESP32 no acumula trafico irrelevante cuando no hay un punto activo y, al mismo tiempo, la configuracion de STA se aplica sobre un canal y un BSSID conocidos. El resultado practico es una mejora en la trazabilidad de la campaña y una disminucion de las capturas espurias fuera del punto de entrenamiento.")
    add_body_paragraph(doc, "Este firmware fue el que hizo posible la campana demo_s5 / train_05. Su principal aporte fue alinear la logica de captura con la hipotesis experimental definitiva: todas las mediciones debian obtenerse en el mismo canal en el que luego se ejecutaria la inferencia en linea. Asi, el dataset final ya no quedo contaminado por una dinamica distinta de la que se esperaba en la operacion real.")

    add_heading(doc, "25.3 Firmware nuevo de inferencia continua", level=2)
    add_body_paragraph(doc, "Para la inferencia en linea se adopto una estrategia diferente. A diferencia del firmware de entrenamiento, el objetivo ya no era capturar solo cuando un punto estuviera abierto, sino permanecer asociado al punto de acceso de forma permanente y sostener una escucha continua de la MAC objetivo. En esta version se incorporaron reintentos escalonados por antena, jitter pseudoaleatorio, retroceso adicional frente a AUTH_EXPIRE y filtrado TARGET_ONLY para evitar que el buffer se llenara con trafico de otros dispositivos.")
    add_code_block(doc, "Fragmento 25.3. Firmware continuo para inferencia en linea", NEW_INFERENCE_CODE, "La funcion anchor_connect_slot_delay_ms() asigna un retardo de arranque distinto a cada antena ESP32, con el fin de reducir la probabilidad de que todas intenten autenticarse simultaneamente contra el mismo punto de acceso. Este detalle, aparentemente menor, fue importante para atenuar fallas intermitentes de autenticacion. Ademas, el callback de captura permite seguir observando al emisor objetivo incluso cuando capture_active es falso, lo cual constituye la base operativa de la seccion de inferencia en linea de la interfaz.")
    add_body_paragraph(doc, "La adopcion de este firmware diferenciado para inferencia responde a una premisa de ingenieria clara: entrenamiento e inferencia son fases con restricciones operativas distintas y, por lo tanto, no conviene forzarlas a compartir exactamente la misma politica de captura de red. El firmware continuo se optimizo para persistencia, estabilidad de asociacion y bajo trafico irrelevante; el de entrenamiento, para trazabilidad y control del muestreo por punto.")

    add_page_break(doc)
    add_heading(doc, "25.4 Firmware original del emisor objetivo", level=2)
    add_body_paragraph(doc, "El primer firmware del emisor objetivo buscaba emular el comportamiento amplio de un dispositivo movil, barriendo multiples canales con una potencia moderada. Esta estrategia fue adecuada para validar que las antenas pudieran detectar probes en varios canales, pero genero una discrepancia metodologica respecto de la inferencia posterior, que ya no funcionaria sobre un barrido multicanal sino sobre un entorno controlado por canal fijo.")
    add_code_block(doc, "Fragmento 25.4. Firmware original del TAG con barrido de canales", OLD_TAG_CODE, "El emisor transmitia sobre un conjunto de canales consecutivos y con una potencia menor que la utilizada en la etapa final. En consecuencia, la distribucion espacial del RSSI durante las primeras campanas no era directamente equivalente a la que luego se observo en inferencia. Esta diferencia explica parte de las inconsistencias detectadas al comparar entrenamiento e inferencia en las corridas previas a la campana final.")

    add_heading(doc, "25.5 Firmware nuevo del emisor objetivo con canal fijo", level=2)
    add_body_paragraph(doc, "La solucion adoptada consistio en reemplazar el barrido multicanal por una emision fija sobre el mismo canal del punto de acceso. Ademas, se incremento la densidad de probes por ciclo y la potencia de transmision, con el objetivo de mejorar la probabilidad de recepcion simultanea por parte de las cuatro antenas ESP32.")
    add_code_block(doc, "Fragmento 25.5. Firmware final del TAG con canal fijo y rafagas densas", NEW_TAG_CODE, "Este bloque fija el canal del emisor, establece un periodo de noventa milisegundos y emite tres probes por ciclo con una separacion de doce milisegundos entre ellos. La potencia de transmision se eleva a aproximadamente diecinueve decibel-milivatios. El criterio de diseño fue simple: si la inferencia iba a exigir continuidad y simultaneidad de recepcion, el emisor objetivo debia adoptar una estrategia de transmision coherente con esa necesidad.")
    if "board_selection" in assets:
        add_picture_with_caption(doc, assets["board_selection"], "Figura 25.1. Captura del entorno de programacion durante la incorporacion de una nueva placa ESP32 para el emisor objetivo. La seleccion de la placa generica ESP32 Dev Module fue utilizada como criterio operativo para programar el hardware de reemplazo.", width=5.8)

    add_heading(doc, "25.6 Ajuste final del software de inferencia para evitar relleno artificial", level=2)
    add_body_paragraph(doc, "Durante la validacion de la inferencia en linea se identifico un problema metodologico adicional: cuando alguna antena ESP32 no completaba la ventana esperada de muestras, el pipeline completaba la fila de entrada del modelo con valores RSSI artificiales de -100 dBm. Aunque esta decision permitia sostener una dimension fija del vector, tambien sesgaba la huella de entrada y empujaba la prediccion hacia regiones del plano que no necesariamente correspondian a la posicion real del emisor objetivo.")
    add_code_block(doc, "Fragmento 25.6. Ajuste del pipeline para permitir inferencia estricta sin relleno artificial", LIVE_PADDING_CODE, "La incorporacion del parametro pad_missing y del selector strict_full_window permite que la interfaz decida si una prediccion debe ejecutarse solo cuando las antenas disponen de la ventana completa de muestras reales. Esta modificacion fortalece la consistencia metodologica del sistema, porque evita que el modelo reciba entradas que nunca formaron parte del espacio de entrenamiento verdadero.")
    add_body_paragraph(doc, "Desde un punto de vista ingenieril, este cambio cierra la principal brecha detectada entre el dataset y la operacion en linea: la dimension fija del vector de entrada ya no obliga a inventar mediciones inexistentes, sino que puede resolverse mediante una politica de espera hasta reunir la ventana real requerida por el modelo.")


def add_section_26(doc: Document, run_s2: dict, run_s4: dict, run_s5: dict, figures: dict[str, Path], assets: dict[str, Path]) -> None:
    add_heading(doc, "26. Resultados finales y discusion tecnica consolidada", level=1)
    add_body_paragraph(doc, "La lectura integrada de las campanas demo_s2, demo_s4 y demo_s5 permite formular una conclusion tecnica importante: el rendimiento de un sistema de localizacion indoor basado en fingerprinting RSSI no depende linealmente ni del numero de antenas ESP32 ni del volumen de muestras por punto. Ambos factores son relevantes, pero su efecto esta mediado por la estabilidad de la capa de enlace, la coherencia entre el firmware del emisor objetivo y el utilizado durante la inferencia, la simetria real del despliegue y la ausencia de artefactos metodologicos como el relleno artificial de muestras faltantes.")
    add_table_title(doc, "Tabla 26.1. Comparacion consolidada de las campanas con dataset y modelo")
    add_table(doc, ["Campana", "Antenas", "Muestras por antena", "Puntos", "Dimensiones de dataset", "Mejor modelo", "MAE [m]", "P95 [m]"], [
        ["demo_s2 / train_02", 3, 10, len(run_s2["points"]), f"{run_s2['dataset'].shape[0]} x {run_s2['dataset'].shape[1]}", run_s2["best_mae"]["model"], fmt_float(run_s2["best_mae"]["mae_eucl"], 3), fmt_float(run_s2["best_p95"]["p95"], 3)],
        ["demo_s4 / train_04", 4, 10, len(run_s4["points"]), f"{run_s4['dataset'].shape[0]} x {run_s4['dataset'].shape[1]}", run_s4["best_mae"]["model"], fmt_float(run_s4["best_mae"]["mae_eucl"], 3), fmt_float(run_s4["best_p95"]["p95"], 3)],
        ["demo_s5 / train_05", 4, 50, len(run_s5["points"]), f"{run_s5['dataset'].shape[0]} x {run_s5['dataset'].shape[1]}", run_s5["best_mae"]["model"], fmt_float(run_s5["best_mae"]["mae_eucl"], 3), fmt_float(run_s5["best_p95"]["p95"], 3)],
    ])
    if "campaign_compare" in figures:
        add_picture_with_caption(doc, figures["campaign_compare"], "Figura 26.1. Comparacion sintetica de MAE y P95 entre las campanas demo_s2, demo_s4 y demo_s5. La grafica muestra que el crecimiento del numero de antenas o del volumen muestral no se tradujo de manera automatica en una mejora monotona del error.", width=6.0)
    add_body_paragraph(doc, "El mejor error medio absoluto del proyecto se obtuvo en demo_s2 con tres antenas ESP32. Esta observacion, lejos de invalidar la evolucion posterior, ayuda a comprender la naturaleza del problema. demo_s2 se beneficio de un hardware mas simple y de una menor complejidad operativa, pero opero con una topologia menos rica y con un volumen muestral mas reducido. demo_s4 y demo_s5 aumentaron la complejidad del sistema y, con ello, introdujeron nuevas fuentes de variabilidad que debieron ser absorbidas mediante ajustes de firmware, de red y de metodologia.")
    add_body_paragraph(doc, "En demo_s4 la incorporacion de la cuarta antena mejoro la cobertura espacial, pero el sistema todavia sufria inestabilidades de asociacion, diferencias entre el emisor utilizado para entrenamiento y el empleado en inferencia y una politica de captura que no estaba completamente adaptada al canal fijo del punto de acceso. demo_s5 resolvio buena parte de estas limitaciones mediante el nuevo firmware de entrenamiento y el aumento a cincuenta muestras por antena. Sin embargo, el propio incremento de la exigencia muestral introdujo tiempos de captura muy variables y, por ende, una heterogeneidad experimental que tambien impacta en el comportamiento del modelo.")
    add_body_paragraph(doc, "La validacion final de inferencia en linea posterior a la recuperacion de la placa permite, no obstante, cerrar el proyecto con una conclusion positiva. Aun cuando la precision absoluta todavia admite mejoras, el sistema final logro operar con cuatro antenas ESP32, un emisor objetivo de canal fijo, un servidor de adquisicion en FastAPI, una interfaz de operacion en Streamlit y un pipeline de entrenamiento e inferencia unificados. Desde una perspectiva de ingenieria aplicada, este resultado constituye una plataforma funcional y reproducible, sobre la cual pueden apoyarse iteraciones posteriores de mejora.")
    if "sheet_router" in assets:
        add_picture_with_caption(doc, assets["sheet_router"], "Figura 26.2. Capturas de la configuracion del punto de acceso utilizadas durante los ajustes finales de compatibilidad Wi-Fi. Se documentan el canal fijo, la seguridad WPA2-PSK, la desactivacion de filtros MAC y la parametrizacion de radio empleada para la fase final.", width=6.2)


def add_section_27(doc: Document) -> None:
    add_heading(doc, "27. Posibles mejoras", level=1)
    add_body_paragraph(doc, "El sistema desarrollado ya es funcional y reproducible, pero aun ofrece un margen amplio de evolucion. Las mejoras potenciales no se restringen al modelo de aprendizaje automatico; abarcan el hardware de radiofrecuencia, la capa de red, la metodologia de relevamiento, la persistencia de datos, la instrumentacion de la interfaz y la validacion cuantitativa del sistema.")
    improvements = [
        ("Infraestructura de radio", "Emplear un punto de acceso dedicado exclusivamente al experimento, con SSID unico, canal fijo, ancho de banda de 20 MHz y aislamiento total respecto de otros extensores o repetidores con el mismo nombre de red."),
        ("Antenas ESP32", "Unificar el hardware de las cuatro placas receptoras, utilizando el mismo modelo, misma revision de placa y, de ser posible, antenas externas calibradas para reducir variaciones intermodulares entre dispositivos."),
        ("Emisor objetivo", "Implementar una carcasa fija para el TAG, con orientacion mecanica controlada y fuente de alimentacion estable, para reducir la variabilidad derivada de la mano del operador o de la posicion de la placa durante cada punto."),
        ("Alimentacion", "Migrar a una alimentacion dedicada y estabilizada para cada ESP32, evitando puertos USB ruidosos o cables con caida de tension que puedan afectar la etapa de radiofrecuencia."),
        ("Firmware", "Agregar telemetria local mas detallada en cada antena, incluyendo contadores de paquetes filtrados, razones de descarte y tiempo acumulado en cada estado de la maquina de conexion."),
        ("Persistencia local", "Incorporar un buffer circular persistente en memoria flash o en SPIFFS para que la antena pueda reenviar lotes si se produce una perdida momentanea de conectividad con el servidor."),
        ("Reloj y sincronizacion", "Sincronizar todas las placas con una fuente horaria comun o con marcas de tiempo derivadas del servidor para facilitar analisis temporales finos y correlacion entre antenas."),
        ("Metodologia de entrenamiento", "Mantener la exigencia de canal fijo y evaluar si el valor optimo de muestras por antena es cincuenta o cien, equilibrando robustez estadistica con tiempos razonables de campana."),
        ("Diseño del dataset", "Explorar representaciones mas compactas y robustas, como percentiles, mediana, IQR, promedio recortado o descriptores temporales, en lugar de depender exclusivamente de ventanas largas de muestras crudas."),
        ("Modelado", "Evaluar modelos probabilisticos o ensambles con estimacion de incertidumbre, de modo que la interfaz no solo entregue una posicion estimada sino tambien un indicador de confiabilidad asociado."),
        ("Postprocesado temporal", "Aplicar filtros temporales sobre la trayectoria inferida, por ejemplo promedios moviles, medianas deslizantes o filtros de Kalman, para estabilizar visualmente la posicion en linea."),
        ("Validacion", "Separar explicitamente un conjunto ciego de puntos de prueba no utilizados durante el entrenamiento, con multiples repeticiones por punto, para obtener una estimacion mas representativa del error real de generalizacion."),
        ("Diagnostico en interfaz", "Mostrar en la interfaz, de forma visible, las ultimas muestras RSSI por antena y alertas cuando la ventana de inferencia se construya con datos incompletos o insuficientes."),
        ("Control de versiones", "Versionar formalmente firmware, configuracion del punto de acceso, dataset, modelo y documento tecnico como un paquete unico por campana, con el fin de maximizar la reproducibilidad futura."),
        ("Escalabilidad", "Evaluar una arquitectura orientada a mensajes para la adquisicion, donde las antenas publiquen lotes a una cola local o broker ligero antes de que el servidor los persista y los procese."),
        ("Calibracion espacial", "Medir con mayor precision la geometria real del recinto, la altura efectiva de cada antena y la ubicacion exacta del emisor objetivo en cada punto, minimizando errores de etiquetado espacial."),
        ("Pruebas de robustez", "Agregar ensayos con personas en movimiento, puertas abiertas o cerradas y variacion deliberada de orientacion del emisor, para cuantificar la sensibilidad del sistema a perturbaciones ambientales."),
        ("Marco comparativo", "Contrastar el rendimiento del fingerprinting RSSI con metodos geometricos simplificados, con fusion sensorial o con BLE dedicado, a fin de situar el sistema dentro de un espectro mas amplio de soluciones de localizacion indoor."),
    ]
    add_table_title(doc, "Tabla 27.1. Propuestas de mejora identificadas a partir de la experiencia experimental")
    add_table(doc, ["Area", "Mejora propuesta"], improvements)


def add_section_28(doc: Document) -> None:
    add_heading(doc, "28. Conclusiones finales", level=1)
    add_body_paragraph(doc, "El proyecto permitio diseñar, implementar y validar una plataforma integral de localizacion indoor basada en fingerprinting RSSI, utilizando hardware de bajo costo basado en ESP32, un servidor de adquisicion en FastAPI, una interfaz de operacion en Streamlit y un pipeline de entrenamiento e inferencia construido en Python. Desde la perspectiva de un trabajo de ingenieria, el principal resultado no es solamente el valor de error alcanzado por un modelo particular, sino la construccion de un sistema completo, reproducible y suficientemente documentado para servir como base de evolucion futura.")
    add_body_paragraph(doc, "El recorrido experimental dejo tres conclusiones centrales. En primer lugar, la calidad del posicionamiento no depende solo del algoritmo de aprendizaje, sino de la coherencia entre el firmware de adquisicion, la configuracion del punto de acceso, la estabilidad de las antenas ESP32 y la consistencia del emisor objetivo entre entrenamiento e inferencia. En segundo lugar, el aumento de antenas o de muestras por punto no garantiza por si mismo una mejora monotona del error, ya que el beneficio de una mayor densidad espacial o estadistica puede verse contrarrestado por inestabilidades operativas si el sistema de captura no acompaña ese aumento de complejidad. En tercer lugar, la trazabilidad detallada de cada campana resulto indispensable para interpretar correctamente los resultados y para distinguir entre errores atribuibles al modelo y errores atribuibles a la instrumentacion.")
    add_body_paragraph(doc, "La campana demo_s2 proporciono la primera evidencia cuantitativa solida de viabilidad, mientras que demo_s4 mostro la necesidad de endurecer el sistema frente a las dificultades reales de trabajar con cuatro antenas ESP32. Finalmente, demo_s5 integro las principales correcciones metodologicas y de firmware: canal fijo, bloqueo por BSSID, firmware especifico de entrenamiento, firmware continuo de inferencia y version final del emisor objetivo. La validacion posterior con la placa recuperada confirmo que el sistema podia sostener una operacion simultanea de las cuatro antenas sobre una campana ya entrenada.")
    add_body_paragraph(doc, "En sintesis, el proyecto cumplio con el objetivo de desarrollar una solucion funcional de localizacion indoor de bajo costo, documentar rigurosamente sus resultados y exponer con honestidad tecnica sus limites actuales. El marco teorico del trabajo, utilizado como referencia conceptual de base, encontro aqui una materializacion practica que permitio contrastar teoria y realidad experimental. El sistema NEXA resultante no debe considerarse un producto cerrado, sino una plataforma experimental madura, con arquitectura propia, evidencia empirica y una hoja de ruta clara de mejoras para futuras iteraciones.")


def add_section_29(doc: Document) -> None:
    add_heading(doc, "29. Referencias tecnicas complementarias", level=1)
    add_body_paragraph(doc, "Las referencias que se listan a continuacion complementan el marco teorico general del proyecto y respaldan las decisiones de implementacion de hardware y software adoptadas durante la parte practica.")
    for item in REFERENCE_TEXTS:
        add_body_paragraph(doc, item)


def main() -> None:
    ensure_dirs()
    if not SOURCE_DOC.exists():
        raise FileNotFoundError(f"No se encontro el documento base: {SOURCE_DOC}")

    run_s1 = load_run_summary(RUN_S1)
    run_s2 = load_run_summary(RUN_S2)
    run_s3 = load_run_summary(RUN_S3)
    run_s4 = load_run_summary(RUN_S4)
    run_s5 = load_run_summary(RUN_S5)

    figures = build_figures(run_s2, run_s4, run_s5)
    assets = build_visual_assets()

    doc = Document(SOURCE_DOC)
    remove_body_from_heading_to_end(doc, "24.")

    add_section_24(doc, run_s1, run_s2, run_s3, run_s4, run_s5, figures, assets)
    add_page_break(doc)
    add_section_25(doc, assets)
    add_page_break(doc)
    add_section_26(doc, run_s2, run_s4, run_s5, figures, assets)
    add_page_break(doc)
    add_section_27(doc)
    add_page_break(doc)
    add_section_28(doc)
    add_page_break(doc)
    add_section_29(doc)

    doc.save(OUTPUT_DOC)
    print(f"Documento generado en: {OUTPUT_DOC}")


if __name__ == "__main__":
    main()
